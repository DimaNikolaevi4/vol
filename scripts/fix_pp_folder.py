#!/usr/bin/env python3
"""
Исправление всех проблем в папке ПП:
1. 103: D-баланс B→D (swap Q4 B↔D), окончания .→;
2. 104: окончания .→;
3. 105: 4 опции без окончания → добавить ;
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'
FOLDER = 'профессиональной подготовки 15474 Оператор станков'

Q_RE = re.compile(r'^(\d+)\s*[.)]\s*(.*)')
OPT_RE = re.compile(r'^([ABCD])\)\s*(.*)')
ANS_RE = re.compile(r'Правильный ответ под номером:\s*([ABCD])')
TICKET_RE = re.compile(r'Билет\s+№?\s*(\d+)')


def parse_103_104(filepath, return_headers=False):
    doc = Document(filepath)
    questions = []
    header_paras = []
    in_questions = False
    cur_q = None; cur_text = None; cur_opts = {}
    for p in doc.paragraphs:
        txt_s = p.text.strip()
        if not txt_s:
            if not in_questions: header_paras.append('')
            continue
        m_ans = ANS_RE.search(txt_s)
        if m_ans and cur_q is not None:
            questions.append({'num': cur_q, 'text': cur_text, 'options': dict(cur_opts), 'answer': m_ans.group(1)})
            cur_q = None; cur_opts = {}
            continue
        m_q = Q_RE.match(txt_s)
        if m_q and cur_q is None:
            in_questions = True
            cur_q = int(m_q.group(1)); cur_text = m_q.group(2).strip(); cur_opts = {}
            continue
        if not in_questions:
            header_paras.append(p.text)
            continue
        m_opt = OPT_RE.match(txt_s)
        if m_opt and cur_q is not None:
            cur_opts[m_opt.group(1)] = m_opt.group(2).strip()
    if return_headers:
        return header_paras, questions
    return questions


def parse_105(filepath):
    doc = Document(filepath)
    tickets = {}; cur_ticket = None; cur_q = None; cur_text = None; cur_opts = {}
    for p in doc.paragraphs:
        txt_s = p.text.strip()
        if not txt_s: continue
        tm = TICKET_RE.search(txt_s)
        if tm:
            cur_ticket = int(tm.group(1)); tickets[cur_ticket] = []; cur_q = None; continue
        if cur_ticket is None: continue
        m_ans = ANS_RE.search(txt_s)
        if m_ans and cur_q is not None:
            tickets[cur_ticket].append({'num': cur_q, 'text': cur_text, 'options': dict(cur_opts), 'answer': m_ans.group(1)})
            cur_q = None; cur_opts = {}
            continue
        m_q = Q_RE.match(txt_s)
        if m_q and cur_q is None:
            cur_q = int(m_q.group(1)); cur_text = m_q.group(2).strip(); cur_opts = {}
            continue
        m_opt = OPT_RE.match(txt_s)
        if m_opt and cur_q is not None:
            cur_opts[m_opt.group(1)] = m_opt.group(2).strip()
    return tickets


def build_103_104(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if not s: continue
        if s.startswith('Фонд') or 'тестовых' in s.lower() or 'Тестовые' in s:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
    for q in questions:
        qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
        qp.paragraph_format.space_after = Pt(0)
        qp.paragraph_format.space_before = Pt(2)
        for letter in 'ABCD':
            if letter in q['options']:
                op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                op.paragraph_format.space_after = Pt(0)
                op.paragraph_format.space_before = Pt(0)
                op.paragraph_format.left_indent = Cm(1)
        ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
        ap.paragraph_format.space_after = Pt(4)
        ap.paragraph_format.space_before = Pt(0)
        for run in ap.runs:
            run.font.bold = True
    doc.save(output_path)


def build_105(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if not s: continue
        if (s.startswith('Фонд') or 'итоговой' in s or 'по основной' in s or
            s.startswith('(профессиональной') or 'по профессии' in s or
            s.startswith('«Оператор') or 'Тестовые' in s or
            '(20 вопросов' in s or TICKET_RE.match(s)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if TICKET_RE.match(s) else Pt(12)
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets: tickets[t] = []
        tickets[t].append(q)
    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph('Билет ' + str(t))
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        for q in tickets[t]:
            qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)
            ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            for run in ap.runs:
                run.font.bold = True
    doc.save(output_path)


def fix_endings_103_104(questions):
    """Unify all option endings to ';'."""
    n = 0
    for q in questions:
        for letter in 'ABCD':
            if letter in q['options']:
                opt = q['options'][letter]
                if opt and not opt.endswith(';'):
                    if opt.endswith('.'):
                        q['options'][letter] = opt[:-1] + ';'
                        n += 1
                    else:
                        q['options'][letter] = opt + ';'
                        n += 1
    return n


def fix_endings_105(tickets):
    """Unify all option endings to ';' in 105 format."""
    n = 0
    for tnum in tickets:
        for q in tickets[tnum]:
            for letter in 'ABCD':
                if letter in q['options']:
                    opt = q['options'][letter]
                    if opt and not opt.endswith(';'):
                        if opt.endswith('.'):
                            q['options'][letter] = opt[:-1] + ';'
                            n += 1
                        else:
                            q['options'][letter] = opt + ';'
                            n += 1
    return n


def main():
    print('=' * 70)
    print('ИСПРАВЛЕНИЕ ВСЕЙ ПАПКИ ПП')
    print('=' * 70)

    # ===== FILE 103 =====
    f103 = BASE + '/' + FOLDER + '/103. ПО_П_ФОС_Оператор станков_2-3_разр ОП.0.0.docx'
    print('\n[1] ПП 103')
    header103, questions103 = parse_103_104(f103, return_headers=True)

    # 1a. Fix D-balance: Q4 B→D (swap options B and D text)
    for q in questions103:
        if q['num'] == 4:
            print('  D-баланс: swap Q4 B<->D')
            print(f'    Was B: {q["options"]["B"][:50]}')
            print(f'    Was D: {q["options"]["D"][:50]}')
            q['options']['B'], q['options']['D'] = q['options']['D'], q['options']['B']
            q['answer'] = 'D'
            print(f'    Now B: {q["options"]["B"][:50]}')
            print(f'    Now D: {q["options"]["D"][:50]}')
            print(f'    Answer: D')
            break

    d = Counter(q['answer'] for q in questions103)
    print(f'  D-баланс после swap: A={d["A"]} B={d["B"]} C={d["C"]} D={d["D"]}')
    assert d['A'] == d['B'] == d['C'] == d['D'] == 5, 'D-баланс 103 нарушен!'

    # 1b. Fix endings
    n_end = fix_endings_103_104(questions103)
    print(f'  Окончаний исправлено: {n_end}')

    # 1c. Build
    build_103_104(header103, questions103, f103)
    print('  103 сохранён')

    # Verify
    vq103 = parse_103_104(f103)
    dv = Counter(q['answer'] for q in vq103)
    print(f'  Верификация: {len(vq103)} вопросов, D-баланс: A={dv["A"]} B={dv["B"]} C={dv["C"]} D={dv["D"]}')

    # ===== FILE 104 =====
    f104 = BASE + '/' + FOLDER + '/104. ПО_П_ФОС_Оператор станков_2-3_разр МДК01.01.docx'
    print('\n[2] ПП 104')
    header104, questions104 = parse_103_104(f104, return_headers=True)

    # 2a. Fix endings
    n_end = fix_endings_103_104(questions104)
    print(f'  Окончаний исправлено: {n_end}')

    # 2b. Build
    build_103_104(header104, questions104, f104)
    print('  104 сохранён')

    # Verify
    vq104 = parse_103_104(f104)
    dv = Counter(q['answer'] for q in vq104)
    print(f'  Верификация: {len(vq104)} вопросов, D-баланс: A={dv["A"]} B={dv["B"]} C={dv["C"]} D={dv["D"]}')

    # ===== FILE 105 =====
    f105 = BASE + '/' + FOLDER + '/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx'
    print('\n[3] ПП 105')
    doc105 = Document(f105)
    header105 = [p.text for p in doc105.paragraphs]
    tickets105 = parse_105(f105)

    # 3a. Fix endings
    n_end = fix_endings_105(tickets105)
    print(f'  Окончаний исправлено: {n_end}')

    # 3b. Build (convert tickets to flat list)
    all_q = []
    for tnum in sorted(tickets105.keys()):
        for q in tickets105[tnum]:
            q['ticket'] = tnum
            all_q.append(q)
    build_105(header105, all_q, f105)
    print('  105 сохранён')

    # Verify
    vt105 = parse_105(f105)
    total = sum(len(v) for v in vt105.values())
    d105 = Counter()
    for tnum in vt105:
        for q in vt105[tnum]:
            d105[q['answer']] += 1
    print(f'  Верификация: {total} вопросов, D-баланс: A={d105["A"]} B={d105["B"]} C={d105["C"]} D={d105["D"]}')

    print('\n' + '=' * 70)
    print('ВСЕ ФАЙЛЫ ОБРАБОТАНЫ')
    print('=' * 70)


if __name__ == '__main__':
    main()
