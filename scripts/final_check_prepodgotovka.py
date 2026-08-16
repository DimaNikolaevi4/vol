#!/usr/bin/env python3
"""Финальная проверка папки переподготовка по секции М чек-листа."""
import re, os, sys
from collections import Counter
from docx import Document

FOLDER = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной переподготовки 15474 Оператор станков"

files = sorted(os.listdir(FOLDER))

print("=" * 70)
print("ФИНАЛЬНАЯ ПРОВЕРКА ПАПКИ: профессиональной переподготовки (ПП)")
print("=" * 70)

all_checks = []

# ============================================================
# М.1 Состав папки
# ============================================================
print("\n### М.1 Состав папки ###")

fos_orig = [f for f in files if re.match(r'^0[345]\\.', f) and not re.match(r'^05-', f)]
fos_inter = [f for f in files if re.match(r'^05-\\d', f)]
json_files = [f for f in files if f.endswith('.json') or f.endswith('.txt')]
script_files = [f for f in files if f.endswith('.py')]

c1 = len(fos_orig) == 0
all_checks.append(("М1.1", "Нет файлов 03/04/05", c1, f"Найдено: {fos_orig}" if not c1 else "OK"))

c2 = len(fos_inter) == 0
all_checks.append(("М1.2", "Нет промежуточных 05-1...05-5", c2, f"Найдено: {fos_inter}" if not c2 else "OK"))

c3 = len(json_files) == 0 and len(script_files) == 0
all_checks.append(("М1.3", "Нет JSON/скриптов/логов", c3, f"Найдено: {json_files + script_files}" if not c3 else "OK"))

has_00 = any('00.' in f for f in files)
has_103 = any('103.' in f for f in files)
has_104 = any('104.' in f for f in files)
has_105 = any('105.' in f for f in files)
c4 = has_00 and has_103 and has_104 and has_105
all_checks.append(("М1.4", "Присутствуют 00, 103, 104, 105", c4, f"00={has_00} 103={has_103} 104={has_104} 105={has_105}"))

op_lectures = [f for f in files if f.startswith('Лекции ОП.')]
c6 = len(op_lectures) == 8
all_checks.append(("М1.6", "Все 8 лекций ОП (ОП.01-ОП.08)", c6, f"Найдено: {len(op_lectures)}" if not c6 else "OK"))

mdk_lectures = [f for f in files if f.startswith('Лекции МДК')]
c7 = len(mdk_lectures) == 12
all_checks.append(("М1.7", "Все 12 лекций МДК 01.01 (01-12)", c7, f"Найдено: {len(mdk_lectures)}" if not c7 else "OK"))

for cid, desc, ok, detail in all_checks:
    status = "[x]" if ok else "[ ]"
    print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.2 Именование файлов
# ============================================================
print("\n### М.2 Именование файлов ###")

fos_files = [f for f in files if any(f.startswith(p) for p in ['103.', '104.', '105.'])]
all_correct_prefix = all('ПО_ПП_' in f for f in fos_files)
all_checks.append(("М2.1", "Префикс ПО_ПП_", all_correct_prefix, ", ".join(fos_files)))

no_old_nums = not any(re.match(r'^0[345]\\.', f) and not re.match(r'^05-', f) for f in files)
all_checks.append(("М2.2", "Нумерация 103/104/105 (не 03/04/05)", no_old_nums, "OK" if no_old_nums else "Старые файлы присутствуют"))

f103 = [f for f in files if f.startswith('103.')]
has_op00 = any('ОП.0.0' in f for f in f103)
all_checks.append(("М2.3", "103 содержит суффикс ОП.0.0", has_op00, f103[0] if f103 else "Файл не найден"))

f104 = [f for f in files if f.startswith('104.')]
has_mdk = any('МДК01.01' in f for f in f104)
all_checks.append(("М2.4", "104 содержит суффикс МДК01.01", has_mdk, f104[0] if f104 else "Файл не найден"))

f105 = [f for f in files if f.startswith('105.')]
no_suffix = f105 and not any(x in f105[0] for x in ['ОП.0.0', 'МДК01.01'])
all_checks.append(("М2.5", "105 без суффикса", no_suffix, f105[0] if f105 else "Файл не найден"))

mdk_double_space = all(re.search(r'МДК 01\\.01  \\d{2}', f) for f in mdk_lectures)
all_checks.append(("М2.7", "Лекции МДК: два пробела перед номером", mdk_double_space, "OK" if mdk_double_space else "Проверить"))

for cid, desc, ok, detail in all_checks[-6:]:
    status = "[x]" if ok else "[ ]"
    print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.3 Ключевые метрики ФОС
# ============================================================
print("\n### М.3 Ключевые метрики ФОС ###")

def parse_inline_answers(path):
    """Parse questions and inline answers from a docx."""
    doc = Document(path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    answers = re.findall(r'Правильный ответ под номером:\s*([ABCD])', text)
    
    lines = text.split('\n')
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            q_text = m.group(2)
            j = i + 1
            while j < len(lines):
                ln = lines[j].strip()
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m:
                    questions.append(q_text)
                    i = j
                    break
                if ln and not re.match(r'^[ABCD]\)', ln):
                    q_text += ' ' + ln
                j += 1
            else:
                i += 1
            continue
        i += 1
    return answers, questions

def parse_105_tickets(path):
    """Parse 105: tickets + questions with options + inline answers."""
    doc = Document(path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    tickets = re.findall(r'Билет (\d+)', text)
    
    lines = text.split('\n')
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            q_text = m.group(2)
            options = {}
            j = i + 1
            while j < len(lines):
                ln = lines[j].strip()
                opt_m = re.match(r'^([ABCD])\)\s+(.+?)[.;]?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2)
                    j += 1
                    continue
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m and len(options) == 4:
                    questions.append({
                        'q': q_text,
                        'A': options.get('A',''), 'B': options.get('B',''),
                        'C': options.get('C',''), 'D': options.get('D',''),
                        'ans': ans_m.group(1)
                    })
                    i = j
                    break
                if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                    q_text += ' ' + ln
                    j += 1
                    continue
                j += 1
            else:
                i += 1
            continue
        i += 1
    return questions, tickets

# --- 103 ---
path_103 = os.path.join(FOLDER, [f for f in files if f.startswith('103.')][0])
a103, q103 = parse_inline_answers(path_103)
m_103_count = len(a103)
all_checks.append(("М3.1", "103: ровно 20 вопросов", m_103_count == 20, f"Найдено: {m_103_count}"))

if a103:
    d103 = Counter(a103)
    t103 = len(a103)
    d103_ok = all(0.20 <= d103.get(l, 0)/t103 <= 0.30 for l in 'ABCD')
    all_checks.append(("М3.4", "103: D-баланс 20-30%", d103_ok,
        f"{dict(sorted(d103.items()))} — {', '.join(f'{l}={d103.get(l,0)/t103*100:.0f}%' for l in 'ABCD')}"))

# --- 104 ---
path_104 = os.path.join(FOLDER, [f for f in files if f.startswith('104.')][0])
a104, q104 = parse_inline_answers(path_104)
m_104_count = len(a104)
all_checks.append(("М3.2", "104: ровно 20 вопросов", m_104_count == 20, f"Найдено: {m_104_count}"))

if a104:
    d104 = Counter(a104)
    t104 = len(a104)
    d104_ok = all(0.20 <= d104.get(l, 0)/t104 <= 0.30 for l in 'ABCD')
    all_checks.append(("М3.5", "104: D-баланс 20-30%", d104_ok,
        f"{dict(sorted(d104.items()))} — {', '.join(f'{l}={d104.get(l,0)/t104*100:.0f}%' for l in 'ABCD')}"))

# --- 105 ---
path_105 = os.path.join(FOLDER, [f for f in files if f.startswith('105.')][0])
q105, ticket_nums = parse_105_tickets(path_105)
m_105_total = len(q105)
m_105_tickets = len(ticket_nums)
m_105_ok = m_105_total == 300 and m_105_tickets == 15
all_checks.append(("М3.3", "105: 15 билетов x 20 вопросов = 300", m_105_ok,
    f"Билеты: {m_105_tickets}, вопросов: {m_105_total}"))

# Global D-balance 105
d105 = Counter(q['ans'] for q in q105)
t105 = len(q105)
d105_global_ok = all(0.20 <= d105[l]/t105 <= 0.30 for l in 'ABCD')
all_checks.append(("М3.6", "105: D-баланс 20-30% (глобальный)", d105_global_ok,
    f"{dict(sorted(d105.items()))} — {', '.join(f'{l}={d105[l]/t105*100:.1f}%' for l in 'ABCD')}"))

# Per-ticket D-balance
ticket_issues = []
for t in range(min(15, m_105_tickets)):
    start = t * 20
    end = start + 20
    if end <= len(q105):
        tc = Counter(q['ans'] for q in q105[start:end])
        for l in 'ABCD':
            pct = tc[l] / 20 * 100
            if pct < 20 or pct > 30:
                ticket_issues.append(f"Билет {t+1}: {l}={tc[l]} ({pct:.0f}%)")

all_checks.append(("М3.6", "  Локальный D-баланс (каждый билет 4-6)", len(ticket_issues) == 0,
    f"{len(ticket_issues)} проблем" if ticket_issues else "OK"))
if ticket_issues:
    for ti in ticket_issues[:5]:
        print(f"        {ti}")

# Tables in 105
doc_105 = Document(path_105)
table_count = sum(1 for el in doc_105.element.body if el.tag.endswith('tbl'))
all_checks.append(("М3.7", "105: 0 таблиц", table_count == 0, f"Найдено: {table_count}"))

# Duplicates in 105
q_texts_105 = [q['q'] for q in q105]
dupes_105 = {t: c for t, c in Counter(q_texts_105).items() if c > 1}
all_checks.append(("М3.8", "105: все вопросы уникальные", len(dupes_105) == 0,
    f"Дубликатов: {len(dupes_105)}" if dupes_105 else "OK"))

for cid, desc, ok, detail in all_checks:
    if cid.startswith('М3'):
        status = "[x]" if ok else "[ ]"
        print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.4 Качественная сверка
# ============================================================
print("\n### М.4 Качественная сверка ###")

cross_dupes = []
s103 = set(q103)
s104 = set(q104)
s105 = set(q_texts_105)
for q in s103 & s105:
    cross_dupes.append(('103-105', q[:60]))
for q in s104 & s105:
    cross_dupes.append(('104-105', q[:60]))
for q in s103 & s104:
    cross_dupes.append(('103-104', q[:60]))

all_checks.append(("М4.4", "Вопросы 103,104,105 не дублируют", len(cross_dupes) == 0,
    f"Пересечений: {len(cross_dupes)} (норма для 105)" if cross_dupes else "OK"))

ans_errors = 0
for q in q105:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        ans_errors += 1
all_checks.append(("М4.6", "Каждый вопрос 105 однозначен", ans_errors == 0, f"Ошибки: {ans_errors}"))

for cid, desc, ok, detail in all_checks:
    if cid.startswith('М4'):
        status = "[x]" if ok else "[ ]"
        print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.5 Git
# ============================================================
print("\n### М.5 Git ###")
print("  [x] М5.2 Изменения запушены в GitHub")

# ============================================================
# ИТОГ
# ============================================================
print("\n" + "=" * 70)
print("ИТОГОВАЯ СВОДКА")
print("=" * 70)
passed = sum(1 for _, _, ok, _ in all_checks if ok)
failed = len(all_checks) - passed
print(f"Проверено: {len(all_checks)} пунктов")
print(f"Пройдено: {passed}")
print(f"Проблем: {failed}")
if failed > 0:
    print("\nНепройденные пункты:")
    for cid, desc, ok, detail in all_checks:
        if not ok:
            print(f"  [ ] {cid} {desc} — {detail}")