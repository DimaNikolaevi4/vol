#!/usr/bin/env python3
"""
ПП 105 v3: исправление 3 проблем в версии 8
- Б3в6: опечатка 'внутрениго' -> 'внутреннего' (ответ B сохранён)
- Б2в2: замена вопроса (ответ A сохранён)
- Б13в20: замена вопроса (ответ A сохранён)
D-баланс: 75/75/75/75 (без изменений)
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = '/home/z/my-project/upload/105. ПО_П_ФОС_Оператор станков_2-3_разр (8).docx'
BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'
FOLDER = 'профессиональной подготовки 15474 Оператор станков'
DST = BASE + '/' + FOLDER + '/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx'

# ---- Замены (сохраняем тот же ключ ответа) ----
REPLACEMENTS = {
    (2, 2): {  # [A] было: ПТЭЭП периодичность «раз в год»
        'text': 'Согласно ПТЭЭП (Приказ Минэнерго от 12.08.2022 № 811), в каком документе фиксируются результаты проверки знаний норм работы в электроустановках?',
        'options': {
            'A': 'В журнале учёта проверки знаний правил работы в электроустановках;',
            'B': 'В паспорте электроустановки;',
            'C': 'В журнале оперативного обслуживания;',
            'D': 'В наряде-допуске;',
        },
        'answer': 'A',
    },
    (13, 20): {  # [A] было: ПОТ ЭЭ «при наличии диэлектрических перчаток»
        'text': 'Согласно ПОТ ЭЭ (Приказ Минтруда от 15.12.2020 № 903н), какая организационная форма требуется для выполнения работ в электроустановках напряжением выше 1000 В?',
        'options': {
            'A': 'Наряд-допуск;',
            'B': 'Распоряжение;',
            'C': 'Устное указание мастера;',
            'D': 'Перечень работ в порядке текущей эксплуатации;',
        },
        'answer': 'A',
    },
}

# ---- Только опечатки ----
TYPO_FIXES = {
    (3, 6): {'old': 'внутрениго', 'new': 'внутреннего'},
}

Q_RE = re.compile(r'^(\d+)\s*[.)]\s*(.*)')
OPT_RE = re.compile(r'^([ABCD])\)\s*(.*)')
ANS_RE = re.compile(r'Правильный ответ под номером:\s*([ABCD])')
TICKET_RE = re.compile(r'Билет\s+№?\s*(\d+)')


def parse_105(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    in_questions = False
    cur_ticket = None
    cur_q = None
    cur_text = None
    cur_opts = {}
    cur_ans = None

    for p in doc.paragraphs:
        txt = p.text
        txt_s = txt.strip()
        if not txt_s:
            header_paras.append(txt)
            continue
        tm = TICKET_RE.match(txt_s)
        if tm:
            in_questions = True
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_ticket = int(tm.group(1))
            cur_text = None
            header_paras.append(txt)
            continue
        if not in_questions:
            header_paras.append(txt)
            continue
        m = Q_RE.match(txt_s)
        if m:
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_q = int(m.group(1))
            cur_text = m.group(2)
            cur_opts = {}
            cur_ans = None
            continue
        m2 = OPT_RE.match(txt_s)
        if m2 and cur_text is not None:
            cur_opts[m2.group(1)] = m2.group(2)
            continue
        m3 = ANS_RE.search(txt_s)
        if m3 and cur_text is not None:
            cur_ans = m3.group(1)
    if cur_text is not None:
        questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
    return header_paras, questions


def apply_fixes(questions):
    n_repl = 0
    n_typo = 0
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in REPLACEMENTS:
            new = REPLACEMENTS[key]
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            print('  ЗАМЕНА ' + tag + ': ' + q['text'][:60] + '...')
            q['text'] = new['text']
            q['options'] = dict(new['options'])
            q['answer'] = new['answer']
            print('    -> ' + new['text'][:70] + '...')
            n_repl += 1
        if key in TYPO_FIXES:
            fix = TYPO_FIXES[key]
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            if fix['old'] in q['text']:
                print('  ОПЕЧАТКА ' + tag + ': "' + fix['old'] + '" -> "' + fix['new'] + '"')
                q['text'] = q['text'].replace(fix['old'], fix['new'])
                n_typo += 1
            else:
                print('  ОПЕЧАТКА ' + tag + ': НЕ НАЙДЕНА "' + fix['old'] + '" (возможно уже исправлена)')
    return n_repl, n_typo


def build_105(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if not s:
            continue
        if (s.startswith('Фонд') or 'итоговой' in s or 'по основной' in s or
            s.startswith('(профессиональной') or 'по профессии' in s or
            s.startswith('«Оператор') or 'Тестовые' in s or
            '(20 вопросов' in s or TICKET_RE.match(s)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if TICKET_RE.match(s) else Pt(12)
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)
    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph('Билет ' + str(t))
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        for q in tickets[t]:
            qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)
            ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            for run in ap.runs:
                run.font.bold = True
    doc.save(output_path)


def main():
    print('=' * 60)
    print('ПП 105 v3: 2 замены + 1 опечатка (версия 8)')
    print('=' * 60)

    print('\n[1] Парсинг исходного файла...')
    header_paras, questions = parse_105(SRC)
    print('  Вопросов: ' + str(len(questions)))
    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс до: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))

    print('\n[2] Исправления...')
    n_repl, n_typo = apply_fixes(questions)
    print('  Заменено: ' + str(n_repl) + ', опечаток: ' + str(n_typo))

    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс после: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))
    if gc['A'] == gc['B'] == gc['C'] == gc['D'] == 75:
        print('  D-баланс ОК')
    else:
        print('  D-БАЛАНС НАРУШЕН!')
        return

    print('\n[3] Генерация файла...')
    build_105(header_paras, questions, DST)
    print('  Сохранено: ' + DST)

    print('\n[4] Верификация...')
    _, vq = parse_105(DST)
    gc2 = Counter(q['answer'] for q in vq)
    print('  Вопросов: ' + str(len(vq)) + ', D-баланс: A=' + str(gc2['A']) + ' B=' + str(gc2['B']) + ' C=' + str(gc2['C']) + ' D=' + str(gc2['D']))

    ok = True
    for q in vq:
        key = (q['ticket'], q['num'])
        if key in REPLACEMENTS:
            exp = REPLACEMENTS[key]
            match = q['text'] == exp['text'] and q['answer'] == exp['answer']
            status = 'OK' if match else 'FAIL'
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            print('  ' + status + ' ' + tag + ': ' + q['text'][:60] + '... [' + q['answer'] + ']')
            if not match:
                ok = False
        if key in TYPO_FIXES:
            fix = TYPO_FIXES[key]
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            if fix['new'] in q['text']:
                print('  OK ' + tag + ': опечатка исправлена')
            else:
                print('  FAIL ' + tag + ': опечатка НЕ исправлена')
                ok = False

    if ok and gc2['A'] == gc2['B'] == gc2['C'] == gc2['D'] == 75:
        print('\n  ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('\n  ЕСТЬ ПРОБЛЕМЫ')
    print('=' * 60)


if __name__ == '__main__':
    main()