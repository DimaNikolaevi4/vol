#!/usr/bin/env python3
"""
Fix 3 issues in КИПиА 105 (already restored from git):
1. Кондуктометр: answer C->A
2. Vacuum gasket contradiction: change Резина->Медь 
3. "часто" -> reword
Then rebalance D and regenerate.
"""
import re, os, copy, random, sys
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = '/home/z/my-project/vol/Слесарь по контрольно-измерительным приборам и автоматике/профессиональной подготовки 18494 Слесарь КИПиА'
OUTPUT = os.path.join(FOLDER, '105. ПО_П_ФОС_слесарь КИПиА_2-3_разр.docx')
PROF = 'Слесарь по контрольно-измерительным приборам и автоматике'
random.seed(42)

# ============================================================
# PARSE 105
# ============================================================
doc = Document(OUTPUT)
lines = [p.text.strip() for p in doc.paragraphs]

# Parse into tickets: each "Билет N" starts a new ticket
tickets_raw = []
current = None
for line in lines:
    tm = re.match(r'^Билет\s+(\d+)$', line)
    if tm:
        if current is not None:
            tickets_raw.append(current)
        current = []
        continue
    if current is None:
        continue
    if line:
        current.append(line)
if current is not None:
    tickets_raw.append(current)

print(f'Raw tickets: {len(tickets_raw)}')
for i, t in enumerate(tickets_raw):
    print(f'  Ticket {i+1}: {len(t)} lines')

# Parse questions from each ticket
all_questions = []
for t_idx, t_lines in enumerate(tickets_raw):
    i = 0
    while i < len(t_lines):
        line = t_lines[i]
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            q_text = m.group(2)
            options = {}
            ans = None
            j = i + 1
            while j < len(t_lines):
                ln = t_lines[j]
                opt_m = re.match(r'^([ABCD])\)\s+(.+?)\s*;?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2).rstrip(';').strip()
                    j += 1
                    continue
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m:
                    ans = ans_m.group(1)
                    j += 1
                    break
                # Continuation of question text (no option letter, no answer)
                if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                    q_text += ' ' + ln
                    j += 1
                    continue
                j += 1
            if len(options) == 4 and ans:
                all_questions.append({
                    'question': q_text,
                    'A': options['A'], 'B': options['B'],
                    'C': options['C'], 'D': options['D'],
                    'ans': ans, 'orig_ticket': t_idx + 1
                })
            else:
                if len(options) != 4:
                    print(f'  SKIP Ticket {t_idx+1} Q at line {i}: {len(options)} opts, ans={ans}')
            i = j
        else:
            i += 1

print(f'\nParsed: {len(all_questions)} questions')
counts_before = Counter(q['ans'] for q in all_questions)
print(f'Before fix D-balance: {dict(sorted(counts_before.items()))}')

# ============================================================
# APPLY FIXES
# ============================================================
fix_count = 0

# Fix 1: Кондуктометр answer
for q in all_questions:
    if 'электропроводность' in q['question'].lower() and 'жидкост' in q['question'].lower():
        old_ans = q['ans']
        for letter in 'ABCD':
            if 'Кондуктометр' in q[letter] or 'кондуктометр' in q[letter]:
                q['ans'] = letter
                fix_count += 1
                print(f'\nFix 1: Кондуктометр (Ticket {q["orig_ticket"]})')
                print(f'  Q: {q["question"][:70]}')
                print(f'  Old: {old_ans} -> New: {letter} ({q[letter]})')
                break
        break

# Fix 2: Vacuum gasket contradiction
vacuum_qs = [(i, q) for i, q in enumerate(all_questions) 
            if 'прокладк' in q['question'].lower() and 'вакуум' in q['question'].lower()]
print(f'\nVacuum gasket questions: {len(vacuum_qs)}')
for idx, (i, q) in enumerate(vacuum_qs):
    print(f'  [{idx}] Ticket {q["orig_ticket"]}: answer={q["ans"]}')
    print(f'       A={q["A"][:40]}, B={q["B"][:40]}, C={q["C"][:40]}, D={q["D"][:40]}')

# Find the one where answer is Резина (not metal-reinforced) and change to Медь
for i, q in vacuum_qs:
    ans_letter = q['ans']
    ans_text = q[ans_letter]
    # If answer is Резина without армирование
    if 'Резина' in ans_text and 'армирован' not in ans_text.lower():
        # Find which letter has Медь
        for target in 'ABCD':
            if 'Медь' in q[target]:
                old_ans = q['ans']
                q['ans'] = target
                fix_count += 1
                print(f'\nFix 2: Vacuum gasket (Ticket {q["orig_ticket"]})')
                print(f'  Old: {old_ans} ({ans_text}) -> New: {target} ({q[target]})')
                break
        break

# Fix 3: "часто" in M12 connector question
for q in all_questions:
    if 'часто' in q['question'].lower() and 'разъ' in q['question'].lower():
        old_q = q['question']
        # Replace: remove "часто", make definitive
        q['question'] = q['question'].replace('часто используется', 'применяется')
        # Also try other patterns
        q['question'] = re.sub(r'\bчасто\s+', '', q['question'], flags=re.IGNORECASE)
        q['question'] = re.sub(r'\s{2,}', ' ', q['question']).strip()
        # Ensure readable
        if 'предназначен' not in q['question']:
            q['question'] = q['question'].replace('тип электрического разъема используется для',
                                                    'тип электрического разъема предназначен для')
        fix_count += 1
        print(f'\nFix 3: Reworded (Ticket {q["orig_ticket"]})')
        print(f'  Old: {old_q}')
        print(f'  New: {q["question"]}')
        break

print(f'\nTotal fixes: {fix_count}')

# ============================================================
# D-BALANCE
# ============================================================
def swap_answer(q, target_letter):
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q

target_per_option = 75
indices = list(range(len(all_questions)))
MAX_ROUNDS = 500
for round_num in range(MAX_ROUNDS):
    counts = Counter(q['ans'] for q in all_questions)
    over = {l: counts[l] - target_per_option for l in 'ABCD' if counts[l] > target_per_option}
    under = {l: target_per_option - counts[l] for l in 'ABCD' if counts[l] < target_per_option}
    if not over and not under:
        print(f'\nD-balance in round {round_num + 1}: {dict(sorted(counts.items()))}')
        break
    if not over or not under:
        print(f'\nCannot balance: over={over}, under={under}')
        break
    over_letter = max(over, key=over.get)
    under_letter = max(under, key=under.get)
    random.shuffle(indices)
    for i in indices:
        if all_questions[i]['ans'] == over_letter:
            all_questions[i] = swap_answer(copy.deepcopy(all_questions[i]), under_letter)
            break

counts_after = Counter(q['ans'] for q in all_questions)
print(f'After D-balance: {dict(sorted(counts_after.items()))}')

# Verify
errors = 0
for q in all_questions:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        errors += 1
        print(f'  ERROR dup opts: {q["question"][:50]}')
print(f'Duplicate options: {errors}')

q_texts = [q['question'] for q in all_questions]
dupes = {t: c for t, c in Counter(q_texts).items() if c > 1}
if dupes:
    print(f'WARNING: {len(dupes)} duplicate questions!')
    for t, c in dupes.items():
        print(f'  "{t[:60]}" x{c}')
else:
    print('No duplicate questions.')

# ============================================================
# DISTRIBUTE INTO 15 TICKETS
# ============================================================
TICKETS = 15
by_letter = {letter: [] for letter in 'ABCD'}
for q in all_questions:
    by_letter[q['ans']].append(q)
for letter in 'ABCD':
    random.shuffle(by_letter[letter])

tickets = [[] for _ in range(TICKETS)]
for letter in 'ABCD':
    for i, q in enumerate(by_letter[letter]):
        tickets[i % TICKETS].append(q)
for t in tickets:
    random.shuffle(t)

print(f'\nPer-ticket D-balance:')
ticket_ok = True
for t_idx, ticket in enumerate(tickets):
    tc = Counter(q['ans'] for q in ticket)
    line = f'  Билет {t_idx+1:2d}: ' + ' '.join(f'{l}={tc[l]}' for l in 'ABCD')
    if not all(4 <= tc[l] <= 6 for l in 'ABCD'):
        line += ' ***'
        ticket_ok = False
    print(line)
print(f'All tickets OK: {ticket_ok}')

# ============================================================
# GENERATE DOCX
# ============================================================
doc_out = Document()
style = doc_out.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
for section in doc_out.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

title_lines = [
    ('Фонд оценочных средств', True, 16),
    ('итоговой аттестации', True, 16),
    ('', False, 14),
    ('по основной программе профессионального обучения', False, 14),
    ('(профессиональной подготовки)', False, 14),
    ('', False, 14),
    ('по профессии', False, 14),
    (f'\u00AB{PROF}\u00BB', True, 14),
    ('', False, 14),
    ('Тестовые задания с выбором одного правильного ответа', False, 14),
    ('(20 вопросов)', False, 14),
]
for text, bold, size in title_lines:
    p = doc_out.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
doc_out.add_page_break()

for t_idx, ticket in enumerate(tickets):
    p = doc_out.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Билет {t_idx + 1}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    doc_out.add_paragraph()
    
    for q_idx, q in enumerate(ticket, 1):
        p = doc_out.add_paragraph()
        run = p.add_run(f'{q_idx}. {q["question"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        for letter in 'ABCD':
            p = doc_out.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'{letter}) {q[letter]};')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        p = doc_out.add_paragraph()
        run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    if t_idx < len(tickets) - 1:
        doc_out.add_page_break()

doc_out.save(OUTPUT)
print(f'\nSaved: {OUTPUT}')
print(f'\n=== SUMMARY ===')
print(f'Questions: {len(all_questions)}, Fixes: {fix_count}')
print(f'D-balance: {dict(sorted(counts_after.items()))}')
print(f'All tickets OK: {ticket_ok}')
