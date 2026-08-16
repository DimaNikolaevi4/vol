#!/usr/bin/env python3
"""
Convert КИПиА 03->103 and 04->104.
Preserve original question text and options.
- а)б)в)г) -> A)B)C)D)
- Table answer keys -> inline 'Правильный ответ под номером: X'
- Fix D-balance by swapping option TEXT (preserves all original text)
- 04 has missing Q11 - will be noted
"""
import re, os, copy, random
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = '/home/z/my-project/vol/Слесарь по контрольно-измерительным приборам и автоматике/профессиональной подготовки 18494 Слесарь КИПиА'

random.seed(42)

def parse_old_format(path):
    """Parse old 03/04 format: questions in paragraphs with embedded options, answers in table."""
    doc = Document(path)
    
    # Get answer keys from table
    answer_table = doc.tables[1]
    answer_map = {}
    for row in answer_table.rows[1:]:
        q_num = int(row.cells[0].text.strip())
        ans_letter = row.cells[1].text.strip().lower()
        cyr_to_lat = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
        answer_map[q_num] = cyr_to_lat.get(ans_letter, ans_letter.upper())
    
    # Parse questions from paragraphs
    questions = []
    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'^(\d+)\.\s+(.+)', text, re.DOTALL)
        if not m:
            continue
        q_num = int(m.group(1))
        full_text = m.group(2)
        
        lines = full_text.split('\n')
        q_text = lines[0].strip()
        
        options = {}
        for line in lines[1:]:
            line = line.strip()
            opt_m = re.match(r'^([а-гА-Г])\)\s*(.+)', line)
            if opt_m:
                letter_cyr = opt_m.group(1).lower()
                letter_map = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
                options[letter_map[letter_cyr]] = opt_m.group(2).strip().rstrip(' ;.')
        
        if len(options) == 4 and q_num in answer_map:
            questions.append({
                'num': q_num, 'question': q_text,
                'A': options.get('A', ''), 'B': options.get('B', ''),
                'C': options.get('C', ''), 'D': options.get('D', ''),
                'ans': answer_map[q_num]
            })
        elif len(options) != 4:
            print(f'  WARNING Q{q_num}: {len(options)} options')
        elif q_num not in answer_map:
            print(f'  WARNING Q{q_num}: no answer in table')
    
    return questions

def swap_answer(q, target_letter):
    """Swap option TEXT to move answer to target letter (preserves all text)."""
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q

def fix_d_balance(questions):
    """Fix D-balance: each option 4-6 of 20."""
    target = 5
    qs = copy.deepcopy(questions)
    
    MAX_ROUNDS = 200
    for round_num in range(MAX_ROUNDS):
        counts = Counter(q['ans'] for q in qs)
        over = {l: counts[l] - target for l in 'ABCD' if counts[l] > target}
        under = {l: target - counts[l] for l in 'ABCD' if counts[l] < target}
        if not over and not under:
            print(f'  D-balance in round {round_num + 1}: {dict(sorted(counts.items()))}')
            break
        over_letter = next(iter(over)) if over else None
        under_letter = next(iter(under)) if under else None
        if not over_letter or not under_letter:
            break
        over_letter = max(over, key=over.get)
        under_letter = max(under, key=under.get)
        indices = list(range(len(qs)))
        random.shuffle(indices)
        for i in indices:
            if qs[i]['ans'] == over_letter:
                qs[i] = swap_answer(qs[i], under_letter)
                break
    
    return qs

def verify_answers(questions):
    errors = 0
    for q in questions:
        opts = [q['A'], q['B'], q['C'], q['D']]
        if len(opts) != len(set(opts)):
            print(f'  ERROR Q{q["num"]}: duplicate options!')
            errors += 1
    return errors

def create_docx(questions, output_path, title_subtitle, prof_name):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    
    title_lines = [
        ('Фонд оценочных средств', True, 16),
        (title_subtitle, True, 14),
        ('', False, 14),
        ('по основной программе профессионального обучения', False, 14),
        ('(профессиональной подготовки)', False, 14),
        ('', False, 14),
        ('по профессии', False, 14),
        (f'\u00AB{prof_name}\u00BB', True, 14),
        ('', False, 14),
        ('Тестовые задания с выбором одного правильного ответа (20 вопросов)', False, 14),
    ]
    
    for text, bold, size in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    
    doc.add_paragraph()
    
    for q_idx, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{q_idx}. {q["question"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        for letter in 'ABCD':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'{letter}) {q[letter]};')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    print(f'  Saved: {output_path}')

PROF = 'Слесарь по контрольно-измерительным приборам и автоматике'

# ============================================================
# 03 -> 103
# ============================================================
print(f'{"="*60}')
print('Converting 03 -> 103 (ОП.0.0)')
print(f'{"="*60}')

src = [f for f in os.listdir(FOLDER) if f.startswith('03.')][0]
questions_103 = parse_old_format(os.path.join(FOLDER, src))
print(f'  Parsed: {len(questions_103)} questions')

counts_before = Counter(q['ans'] for q in questions_103)
print(f'  Before D-balance: {dict(sorted(counts_before.items()))}')

questions_103 = fix_d_balance(questions_103)
errors = verify_answers(questions_103)
print(f'  Verification: {"OK" if errors == 0 else f"{errors} errors"}')

dst = '103. ПО_П_ФОС_слесарь КИПиА_2-3_разр ОП.0.0.docx'
create_docx(questions_103, os.path.join(FOLDER, dst),
           'промежуточной аттестации по общепрофессиональному циклу (ОП.0.0)', PROF)

# ============================================================
# 04 -> 104
# ============================================================
print(f'\n{"="*60}')
print('Converting 04 -> 104 (МДК01.01)')
print(f'{"="*60}')

src = [f for f in os.listdir(FOLDER) if f.startswith('04.')][0]
questions_104 = parse_old_format(os.path.join(FOLDER, src))
print(f'  Parsed: {len(questions_104)} questions')

# Check which question number is missing
found_nums = {q['num'] for q in questions_104}
missing = set(range(1, 21)) - found_nums
if missing:
    print(f'  MISSING question numbers: {sorted(missing)}')

# Get the answer for missing question from table
doc_04 = Document(os.path.join(FOLDER, src))
table_04 = doc_04.tables[1]
answer_map_04 = {}
for row in table_04.rows[1:]:
    q_num = int(row.cells[0].text.strip())
    ans = row.cells[1].text.strip().lower()
    cyr_to_lat = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
    answer_map_04[q_num] = cyr_to_lat.get(ans, ans.upper())

for m in sorted(missing):
    print(f'  Answer for Q{m} in table: {answer_map_04.get(m, "?")}')

counts_before = Counter(q['ans'] for q in questions_104)
print(f'  Before D-balance: {dict(sorted(counts_before.items()))}')

questions_104 = fix_d_balance(questions_104)
errors = verify_answers(questions_104)
print(f'  Verification: {"OK" if errors == 0 else f"{errors} errors"}')

dst = '104. ПО_П_ФОС_слесарь КИПиА_2-3_разр МДК01.01.docx'
create_docx(questions_104, os.path.join(FOLDER, dst),
           'промежуточной аттестации (МДК01.01) по профессиональному циклу', PROF)

print(f'\nDone!')
