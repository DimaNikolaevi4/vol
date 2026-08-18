#!/usr/bin/env python3
"""
Исправление ГОСТ 19265 → ГОСТ 3882-74 во всех ФОС
ГОСТ 19265-73 = быстрорежущая сталь (НЕ твёрдые сплавы)
ГОСТ 3882-74 = Сплавы твёрдые спечённые. Марки (ВК, ТК, ТТК, ТН)
Затрагивает 3 файла, 9 вопросов. Ключи НЕ меняются.
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'

FILES = {
    'П 105': BASE + '/профессиональной подготовки 15474 Оператор станков/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx',
    'ПП 103': BASE + '/профессиональной переподготовки 15474 Оператор станков/103. ПО_ПП_ФОС_Оператор станков_2-3_разр ОП.0.0.docx',
    'ПП 105': BASE + '/профессиональной переподготовки 15474 Оператор станков/105. ПО_ПП_ФОС_Оператор станков_2-3_разр.docx',
}

OLD = 'ГОСТ 19265'
NEW = 'ГОСТ 3882-74'

Q_RE = re.compile(r'^(\d+)\s*[.)]\s*(.*)')
OPT_RE = re.compile(r'^([ABCD])\)\s*(.*)')
ANS_RE = re.compile(r'Правильный ответ под номером:\s*([ABCD])')
TICKET_RE = re.compile(r'Билет\s+№?\s*(\d+)')
HEADER_WORDS = ('Тестовые', 'Фонд', 'итоговой')


def parse_105(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    in_questions = False
    cur_ticket = None
    cur_q = None
    cur_text = None
    cur_opts = {}
    cur_ans = None

    for p in doc.paragraphs:
        txt = p.text
        txt_s = txt.strip()
        if not txt_s:
            header_paras.append(txt)
            continue
        tm = TICKET_RE.match(txt_s)
        if tm:
            in_questions = True
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_ticket = int(tm.group(1))
            cur_text = None
            header_paras.append(txt)
            continue
        if not in_questions:
            header_paras.append(txt)
            continue
        m = Q_RE.match(txt_s)
        if m:
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_q = int(m.group(1))
            cur_text = m.group(2)
            cur_opts = {}
            cur_ans = None
            continue
        m2 = OPT_RE.match(txt_s)
        if m2 and cur_text is not None:
            cur_opts[m2.group(1)] = m2.group(2)
            continue
        m3 = ANS_RE.search(txt_s)
        if m3 and cur_text is not None:
            cur_ans = m3.group(1)
    if cur_text is not None:
        questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
    return header_paras, questions


def parse_103(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    in_questions = False
    cur_q = None
    cur_text = None
    cur_opts = {}
    cur_ans = None

    for p in doc.paragraphs:
        txt_s = p.text.strip()
        if not txt_s:
            header_paras.append('')
            continue
        if not in_questions:
            if any(w in txt_s for w in HEADER_WORDS):
                header_paras.append(p.text)
                continue
            m = Q_RE.match(txt_s)
            if m:
                in_questions = True
                cur_q = int(m.group(1))
                cur_text = m.group(2)
                cur_opts = {}
                cur_ans = None
                continue
            header_paras.append(p.text)
            continue
        m = Q_RE.match(txt_s)
        if m:
            if cur_text is not None:
                questions.append({'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_q = int(m.group(1))
            cur_text = m.group(2)
            cur_opts = {}
            cur_ans = None
            continue
        m2 = OPT_RE.match(txt_s)
        if m2 and cur_text is not None:
            cur_opts[m2.group(1)] = m2.group(2)
            continue
        m3 = ANS_RE.search(txt_s)
        if m3 and cur_text is not None:
            cur_ans = m3.group(1)
    if cur_text is not None:
        questions.append({'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
    return header_paras, questions


def build_105(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if not s:
            continue
        if (s.startswith('Фонд') or 'итоговой' in s or 'по основной' in s or
            s.startswith('(профессиональной') or 'по профессии' in s or
            s.startswith('«Оператор') or 'Тестовые' in s or
            '(20 вопросов' in s or TICKET_RE.match(s)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if TICKET_RE.match(s) else Pt(12)
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)
    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph('Билет ' + str(t))
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        for q in tickets[t]:
            qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)
            ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            for run in ap.runs:
                run.font.bold = True
    doc.save(output_path)


def build_103(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if s and ('Фонд' in s or 'итоговой' in s or 'по основной' in s or
                  'по профессии' in s or s.startswith('«Оператор') or 'Тестовые' in s):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
    for q in questions:
        qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
        qp.paragraph_format.space_after = Pt(0)
        qp.paragraph_format.space_before = Pt(2)
        for letter in 'ABCD':
            if letter in q['options']:
                op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                op.paragraph_format.space_after = Pt(0)
                op.paragraph_format.space_before = Pt(0)
                op.paragraph_format.left_indent = Cm(1)
        ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
        ap.paragraph_format.space_after = Pt(4)
        ap.paragraph_format.space_before = Pt(0)
        for run in ap.runs:
            run.font.bold = True
    doc.save(output_path)


def main():
    print('=' * 60)
    print('ГОСТ 19265 -> ГОСТ 3882-74 во всех ФОС')
    print('=' * 60)

    total_fixed = 0

    for label, path in FILES.items():
        print('\n--- ' + label + ' ---')
        is_105 = '105' in label
        parse_fn = parse_105 if is_105 else parse_103
        build_fn = build_105 if is_105 else build_103

        header_paras, questions = parse_fn(path)
        gc_before = Counter(q['answer'] for q in questions)

        n = 0
        for q in questions:
            if OLD in q['text']:
                q['text'] = q['text'].replace(OLD, NEW)
                n += 1
                tag = 'Б' + str(q['ticket']) + 'в' + str(q['num']) if is_105 else 'В' + str(q['num'])
                print('  ' + tag + ': ' + q['text'][:80])

        gc_after = Counter(q['answer'] for q in questions)
        if gc_before != gc_after:
            print('  WARNING: D-баланс изменился!')
            print('  До: ' + str(dict(gc_before)))
            print('  После: ' + str(dict(gc_after)))
        else:
            if is_105:
                print('  D-баланс: A=' + str(gc_after['A']) + ' B=' + str(gc_after['B']) + ' C=' + str(gc_after['C']) + ' D=' + str(gc_after['D']) + ' (без изменений)')
            else:
                print('  D-баланс: ' + str(dict(gc_after)) + ' (без изменений)')

        build_fn(header_paras, questions, path)
        print('  Файл сохранён. Исправлено: ' + str(n))
        total_fixed += n

    # Final verification
    print('\n--- Верификация ---')
    all_ok = True
    for label, path in FILES.items():
        doc = Document(path)
        count = sum(1 for p in doc.paragraphs if OLD in p.text)
        count_new = sum(1 for p in doc.paragraphs if NEW in p.text)
        status = 'OK' if count == 0 else 'FAIL (' + str(count) + ' remaining)'
        print('  ' + label + ': ' + OLD + ' -> 0, ' + NEW + ' -> ' + str(count_new) + '  [' + status + ']')
        if count > 0:
            all_ok = False

    print('\nВсего исправлено: ' + str(total_fixed) + ' вопросов')
    if all_ok:
        print('ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('ЕСТЬ ПРОБЛЕМЫ')
    print('=' * 60)


if __name__ == '__main__':
    main()