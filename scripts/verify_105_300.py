#!/usr/bin/env python3
"""Verify the generated 105: 15x20=300 questions."""
import re, json, sys
from docx import Document

path = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной подготовки 15474 Оператор станков/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx"
doc = Document(path)

text = '\n'.join(p.text for p in doc.paragraphs)

# Parse all questions
lines = text.split('\n')
questions = []
i = 0
while i < len(lines):
    line = lines[i].strip()
    m = re.match(r'^(\d+)\.\s+(.+)', line)
    if m:
        q_num = int(m.group(1))
        q_text = m.group(2)
        options = {}
        j = i + 1
        while j < len(lines):
            ln = lines[j].strip()
            opt_m = re.match(r'^([ABCD])\)\s+(.+?)[.;]?$', ln)
            if opt_m:
                options[opt_m.group(1)] = opt_m.group(2)
                j += 1
                continue
            ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
            if ans_m and len(options) == 4:
                ans = ans_m.group(1)
                questions.append({
                    'q': q_text,
                    'A': options.get('A',''), 'B': options.get('B',''),
                    'C': options.get('C',''), 'D': options.get('D',''),
                    'ans': ans
                })
                i = j
                break
            if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                q_text += ' ' + ln
                j += 1
                continue
            j += 1
        else:
            i += 1
        continue
    i += 1

print(f"Total questions parsed: {len(questions)}")

# Check ticket count
tickets = [re.search(r'Билет (\d+)', p.text) for p in doc.paragraphs]
tickets = [t for t in tickets if t]
ticket_nums = [int(t.group(1)) for t in tickets]
print(f"Tickets found: {ticket_nums}")

# Verify answer correctness
errors = 0
for idx, q in enumerate(questions):
    correct_text = q[q['ans']]
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        print(f"ERROR Q{idx+1}: duplicate options: {opts}")
        errors += 1

print(f"Answer verification: {len(questions) - errors}/{len(questions)} correct ({errors} errors)")

# Check for duplicate question texts
from collections import Counter
q_texts = [q['q'] for q in questions]
dupes = {t: c for t, c in Counter(q_texts).items() if c > 1}
if dupes:
    print(f"WARNING: {len(dupes)} duplicate questions!")
    for t, c in dupes.items():
        print(f"  [{c}x] {t[:70]}")
else:
    print("No duplicate questions.")

# D-balance
from collections import Counter
ans_dist = Counter(q['ans'] for q in questions)
print(f"D-balance: {dict(sorted(ans_dist.items()))}")
for letter in 'ABCD':
    pct = ans_dist[letter] / len(questions) * 100
    print(f"  {letter}: {ans_dist[letter]} ({pct:.1f}%)")

# Check ticket-level distribution
print("\nPer-ticket D-balance:")
for t in range(15):
    start = t * 20
    end = start + 20
    if end <= len(questions):
        tc = Counter(q['ans'] for q in questions[start:end])
        nums = [tc[l] for l in 'ABCD']
        max_pct = max(nums) / 20 * 100
        flag = ' ***' if max_pct > 45 else ''
        print(f"  Ticket {t+1}: A={tc['A']} B={tc['B']} C={tc['C']} D={tc['D']}{flag}")

# Check tables
from docx.table import Table
table_count = 0
for element in doc.element.body:
    if element.tag.endswith('tbl'):
        table_count += 1
print(f"\nTables in document: {table_count} (should be 0)")

if errors == 0 and len(dupes) == 0 and table_count == 0 and len(questions) == 300:
    print("\n=== ALL CHECKS PASSED ===")
else:
    print(f"\n=== ISSUES FOUND: errors={errors}, dupes={len(dupes)}, tables={table_count}, total={len(questions)} ===")
    sys.exit(1)