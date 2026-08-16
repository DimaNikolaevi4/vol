#!/usr/bin/env python3
"""Fix lowercase option letters in 103/104: а)б)в)г) -> A)B)C)D)"""
import os, re, sys
from docx import Document

FOLDER = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной переподготовки 15474 Оператор станков"

for prefix in ['103', '104']:
    fname = [f for f in os.listdir(FOLDER) if f.startswith(prefix + '.')][0]
    path = os.path.join(FOLDER, fname)
    doc = Document(path)
    
    fixed = 0
    for p in doc.paragraphs:
        original = p.text
        # Replace lowercase Cyrillic а) б) в) г) with uppercase Latin A) B) C) D)
        new_text = p.text
        new_text = re.sub(r'^([аa])\)', 'A)', new_text)
        new_text = re.sub(r'^([бb])\)', 'B)', new_text)
        new_text = re.sub(r'^([вv])\)', 'C)', new_text)
        new_text = re.sub(r'^([гg])\)', 'D)', new_text)
        
        if new_text != original:
            # Need to replace in runs
            # Clear all runs and set new text while preserving formatting
            if p.runs:
                # Keep first run's formatting, clear the rest
                for i, run in enumerate(p.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
            else:
                p.text = new_text
            fixed += 1
    
    doc.save(path)
    print(f'{prefix}: исправлено {fixed} строк')
    
    # Verify
    doc2 = Document(path)
    lower_a = sum(1 for p in doc2.paragraphs if re.match(r'^[аa]\)', p.text.strip()))
    upper_a = sum(1 for p in doc2.paragraphs if re.match(r'^A\)', p.text.strip()))
    print(f'  Проверка: строчных {lower_a}, заглавных {upper_a}')
