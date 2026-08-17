#!/usr/bin/env python3
"""
Третья волна правок П 105: замена спорных вопросов на НОВЫЕ по действующим НПА

По итогам экспертной проверки v2:
- Б11в14 (IT12/50мм): РЕВЕРТ — 0,25 мм корректна (пользователь подтвердил)
- 7 спорных вопросов заменены на принципиально НОВЫЕ:

Все новые вопросы привязаны к действующим НПА из /norma:
- ПТЭЭП 2023 (Приказ Минэнерго №811 от 12.08.2022)
- Приказ Минтруда №903н от 15.12.2020
- ГОСТ 19265-73 (твёрдые сплавы)
- ГОСТ 25346-2013 (допуски и посадки)

Ключи сохранены → D-баланс 75/75/75/75 без изменений.
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'
FOLDER = 'профессиональной подготовки 15474 Оператор станков'
FILEPATH = f'{BASE}/{FOLDER}/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx'


# === NEW QUESTIONS (same answer letters as originals) ===
NEW_QUESTIONS = {
    # Б5в5 (ответ A) — ГОСТ 19265, группа ВК
    (5, 5): {
        'text': 'Согласно ГОСТ 19265, какие твёрдые сплавы обозначаются буквами «ВК»?',
        'options': {
            'A': 'Вольфрамокобальтовые сплавы (группа ВК)',
            'B': 'Титановольфрамокобальтовые сплавы (группа ТК)',
            'C': 'Титанотанталовые сплавы (группа ТТ)',
            'D': 'Безвольфрамовые сплавы (группа ТН)',
        },
        'answer': 'A',
    },
    # Б6в8 (ответ B) — устройство токарного станка
    (6, 8): {
        'text': 'Какая часть суппорта токарного станка обеспечивает перемещение резца в направлении, перпендикулярном оси вращения заготовки?',
        'options': {
            'A': 'Верхняя (резцовая) каретка;',
            'B': 'Поперечная каретка;',
            'C': 'Нижняя (продольная) каретка;',
            'D': 'Фартук;',
        },
        'answer': 'B',
    },
    # Б3в19 (ответ C) — механика, модуль упругости
    (3, 19): {
        'text': 'Что характеризует модуль упругости (модуль Юнга) материала?',
        'options': {
            'A': 'Твёрдость поверхности детали;',
            'B': 'Способность материала к пластической деформации;',
            'C': 'Способность материала сопротивляться упругой деформации;',
            'D': 'Ударную вязкость при динамических нагрузках;',
        },
        'answer': 'C',
    },
    # Б14в10 (ответ B) — metrology, угломер
    (14, 10): {
        'text': 'Какой измерительный инструмент предназначен для контроля угловых размеров деталей?',
        'options': {
            'A': 'Микрометр гладкий;',
            'B': 'Угломер;',
            'C': 'Штангенциркуль ШЦ-I;',
            'D': 'Индикатор часового типа;',
        },
        'answer': 'B',
    },
    # Б12в1 (ответ C) — режимы резания, глубина
    (12, 1): {
        'text': 'Какой параметр режима резания определяет толщину слоя металла, снимаемого за один проход инструмента?',
        'options': {
            'A': 'Скорость резания;',
            'B': 'Подача;',
            'C': 'Глубина резания;',
            'D': 'Частота вращения шпинделя;',
        },
        'answer': 'C',
    },
    # Б15в5 (ответ C) — Приказ 903н, электротехнический персонал
    (15, 5): {
        'text': 'Согласно Приказу Минтруда России от 15.12.2020 № 903н, какой персонал осуществляет техническое обслуживание и ремонт электроустановок?',
        'options': {
            'A': 'Административно-технический персонал;',
            'B': 'Оперативный персонал;',
            'C': 'Электротехнический персонал;',
            'D': 'Электротехнологический персонал;',
        },
        'answer': 'C',
    },
    # Б1в11 (ответ C) — ПТЭЭП 2023, графики осмотров
    (1, 11): {
        'text': 'Согласно Правилам технической эксплуатации электроустановок (Приказ Минэнерго от 12.08.2022 № 811), кто утверждает графики осмотров электрооборудования?',
        'options': {
            'A': 'Представитель сетевой организации;',
            'B': 'Ответственный за электрохозяйство структурного подразделения;',
            'C': 'Руководитель или иное уполномоченное должностное лицо потребителя;',
            'D': 'Орган Ростехнадзора;',
        },
        'answer': 'C',
    },
    # Б3в3 (ответ C) — ПТЭЭП 2023, группа ответственного
    (3, 3): {
        'text': 'Согласно Правилам технической эксплуатации электроустановок (Приказ Минэнерго от 12.08.2022 № 811), какая группа по электробезопасности должна быть у ответственного за электрохозяйство в электроустановках до 1000 В?',
        'options': {
            'A': 'II группа;',
            'B': 'III группа;',
            'C': 'IV группа;',
            'D': 'V группа;',
        },
        'answer': 'C',
    },
}

# === REVERT Б11в14 to original (0,25 мм is correct for IT12/50mm) ===
REVERT_QUESTIONS = {
    (11, 14): {
        'text': 'Согласно ГОСТ 25346, квалитету IT12 соответствует допуск для размера 50 мм порядка:',
        'options': {
            'A': '0,025 мм;',
            'B': '2,5 мм;',
            'C': '0,0025 мм;',
            'D': '0,25 мм;',
        },
        'answer': 'D',
    },
}


def parse_105(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = {}
    current_answer = None
    in_questions = False

    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()
        if not txt_stripped:
            header_paras.append(txt)
            continue
        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', txt_stripped)
        if ticket_match:
            in_questions = True
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_ticket = int(ticket_match.group(1))
            current_q_text = None
            header_paras.append(txt)
            continue
        if not in_questions:
            header_paras.append(txt)
            continue
        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', txt_stripped)
        if q_match:
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = {}
            current_answer = None
            continue
        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q_text is not None:
            current_options[opt_match.group(1)] = opt_match.group(2)
            continue
        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)
    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket, 'num': current_q_num,
            'text': current_q_text, 'options': current_options,
            'answer': current_answer
        })
    return header_paras, questions


def compute_balance(questions):
    global_counts = Counter()
    per_ticket = {}
    for q in questions:
        a = q['answer']
        global_counts[a] += 1
        t = q['ticket']
        if t not in per_ticket:
            per_ticket[t] = Counter()
        per_ticket[t][a] += 1
    return global_counts, per_ticket


def apply_replacements(questions):
    n = 0

    # Apply NEW question replacements
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in NEW_QUESTIONS:
            new = NEW_QUESTIONS[key]
            old_text = q['text'][:80]
            q['text'] = new['text']
            q['options'] = dict(new['options'])
            q['answer'] = new['answer']
            print(f'  ЗАМЕНА Б{q["ticket"]}в{q["num"]}: {old_text}...')
            print(f'    → {new["text"][:90]}...')
            print(f'    Ответ: {new["answer"]}')
            n += 1

    # Apply REVERTS
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in REVERT_QUESTIONS:
            new = REVERT_QUESTIONS[key]
            old_text = q['text'][:80]
            q['text'] = new['text']
            q['options'] = dict(new['options'])
            q['answer'] = new['answer']
            print(f'  РЕВЕРТ Б{q["ticket"]}в{q["num"]}: {old_text}...')
            print(f'    → {new["text"][:90]}...')
            print(f'    Ответ: {new["answer"]}')
            n += 1

    print(f'\nВсего замен: {n}')
    return n


def build_105_docx(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Фонд') or stripped.startswith('итоговой') or
            stripped.startswith('по основной') or stripped.startswith('(профессиональной') or
            stripped.startswith('по профессии') or stripped.startswith('«Оператор') or
            stripped.startswith('Тестовые') or stripped.startswith('(20 вопросов') or
            re.match(r'^Билет\s+№?\s*\d+$', stripped)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if re.match(r'^Билет', stripped) else Pt(12)

    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)

    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph(f'Билет {t}')
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        for q in tickets[t]:
            qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)

            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)

            ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.left_indent = Cm(0)
            for run in ap.runs:
                run.font.bold = True

    doc.save(output_path)
    print(f"\nФайл сохранён: {output_path}")


def main():
    print('=' * 60)
    print('ИСПРАВЛЕНИЕ 105 П (v3): 8 новых вопросов + 1 реверт')
    print('=' * 60)

    print('\n[1] Парсинг...')
    header_paras, questions = parse_105(FILEPATH)
    print(f'  Вопросов: {len(questions)}')
    gc, pt = compute_balance(questions)
    print(f'  D-баланс до: A={gc["A"]} B={gc["B"]} C={gc["C"]} D={gc["D"]}')

    print('\n[2] Замены...')
    n = apply_replacements(questions)

    gc, pt = compute_balance(questions)
    print(f'\n  D-баланс после: A={gc["A"]} B={gc["B"]} C={gc["C"]} D={gc["D"]}')

    # Verify per-ticket
    bad = []
    for t in sorted(pt.keys()):
        for letter in 'ABCD':
            cnt = pt[t][letter]
            if cnt < 4 or cnt > 6:
                bad.append(f'Билет {t}: {letter}={cnt}')
    if bad:
        print(f'  Локальные нарушения: {bad}')
    else:
        print(f'  Локальный баланс ОК')

    print('\n[3] Верификация...')
    all_ok = True
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in NEW_QUESTIONS:
            expected = NEW_QUESTIONS[key]
            ok = q['answer'] == expected['answer'] and q['text'] == expected['text']
            status = '✅' if ok else '❌'
            print(f'  {status} Б{q["ticket"]}в{q["num"]}: ответ={q["answer"]} (новый вопрос)')
            if not ok: all_ok = False
        elif key in REVERT_QUESTIONS:
            expected = REVERT_QUESTIONS[key]
            ok = q['answer'] == expected['answer'] and '0,25' in q['options'].get('D', '')
            status = '✅' if ok else '❌'
            print(f'  {status} Б{q["ticket"]}в{q["num"]}: ответ={q["answer"]} (реверт к 0,25 мм)')
            if not ok: all_ok = False

    print('\n[4] Генерация файла...')
    build_105_docx(header_paras, questions, FILEPATH)

    print('\n[5] Поверка из файла...')
    _, vq = parse_105(FILEPATH)
    gc2, _ = compute_balance(vq)
    print(f'  Вопросов: {len(vq)}, D-баланс: A={gc2["A"]} B={gc2["B"]} C={gc2["C"]} D={gc2["D"]}')

    if all_ok and gc2['A'] == gc2['B'] == gc2['C'] == gc2['D'] == 75:
        print('\n  🟢 ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('\n  🔴 ЕСТЬ ПРОБЛЕМЫ')

    print('\n' + '=' * 60)
    print('ГОТОВО')
    print('=' * 60)


if __name__ == '__main__':
    main()
