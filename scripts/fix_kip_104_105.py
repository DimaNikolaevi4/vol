#!/usr/bin/env python3
"""
Fix issues found in КИПиА подготовка 104 and 105.
104 Q10: no correct answer -> replace with proper question
105: fix Кондуктометр answer, fix vacuum gasket contradiction, fix "часто" ambiguity
"""
import re, os, copy, random, sys
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = '/home/z/my-project/vol/Слесарь по контрольно-измерительным приборам и автоматике/профессиональной подготовки 18494 Слесарь КИПиА'
PROF = 'Слесарь по контрольно-измерительным приборам и автоматике'
random.seed(42)

# ============================================================
# PARSE HELPERS
# ============================================================
def parse_103_104(path):
    """Parse 103/104 format: numbered questions with inline answers."""
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text.strip())
    
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i]
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            q_num = int(m.group(1))
            q_text = m.group(2)
            options = {}
            ans = None
            j = i + 1
            while j < len(lines):
                ln = lines[j]
                opt_m = re.match(r'^([ABCD])\)\s+(.+?);?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2).rstrip(';')
                    j += 1
                    continue
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m:
                    ans = ans_m.group(1)
                    j += 1
                    break
                # Continuation of question text
                if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                    q_text += ' ' + ln
                    j += 1
                    continue
                j += 1
            if len(options) == 4 and ans:
                questions.append({
                    'num': q_num, 'question': q_text,
                    'A': options['A'], 'B': options['B'],
                    'C': options['C'], 'D': options['D'],
                    'ans': ans
                })
            i = j
        else:
            i += 1
    return questions


def parse_105(path):
    """Parse 105 format: 15 tickets x 20 questions with inline answers."""
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text.strip())
    
    tickets = []
    current_ticket = []
    
    for line in lines:
        ticket_m = re.match(r'^Билет\s+(\d+)$', line)
        if ticket_m:
            if current_ticket:
                tickets.append(current_ticket)
            current_ticket = []
            continue
        
        if not current_ticket:
            continue
        
        current_ticket.append(line)
    
    if current_ticket:
        tickets.append(current_ticket)
    
    # Parse each ticket's questions
    all_questions = []  # flat list of all questions with ticket info
    for t_idx, t_lines in enumerate(tickets):
        questions = []
        i = 0
        while i < len(t_lines):
            line = t_lines[i]
            m = re.match(r'^(\d+)\.\s+(.+)', line)
            if m:
                q_text = m.group(2)
                options = {}
                ans = None
                j = i + 1
                while j < len(t_lines):
                    ln = t_lines[j]
                    opt_m = re.match(r'^([ABCD])\)\s+(.+?);?$', ln)
                    if opt_m:
                        options[opt_m.group(1)] = opt_m.group(2).rstrip(';')
                        j += 1
                        continue
                    ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                    if ans_m:
                        ans = ans_m.group(1)
                        j += 1
                        break
                    if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                        q_text += ' ' + ln
                        j += 1
                        continue
                    j += 1
                if len(options) == 4 and ans:
                    questions.append({
                        'question': q_text,
                        'A': options['A'], 'B': options['B'],
                        'C': options['C'], 'D': options['D'],
                        'ans': ans
                    })
                i = j
            else:
                i += 1
        all_questions.extend(questions)
    
    return all_questions, tickets


def swap_answer(q, target_letter):
    """Swap option TEXT to move answer to target_letter."""
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q


def fix_d_balance_20(questions):
    """Fix D-balance for 20 questions: 4-6 per option."""
    qs = copy.deepcopy(questions)
    target = 5
    indices = list(range(len(qs)))
    MAX_ROUNDS = 200
    for round_num in range(MAX_ROUNDS):
        counts = Counter(q['ans'] for q in qs)
        over = {l: counts[l] - target for l in 'ABCD' if counts[l] > target}
        under = {l: target - counts[l] for l in 'ABCD' if counts[l] < target}
        if not over and not under:
            print(f'  D-balance in round {round_num + 1}: {dict(sorted(counts.items()))}')
            return qs
        over_letter = max(over, key=over.get)
        under_letter = max(under, key=under.get)
        random.shuffle(indices)
        for i in indices:
            if qs[i]['ans'] == over_letter:
                qs[i] = swap_answer(qs[i], under_letter)
                break
    counts = Counter(q['ans'] for q in qs)
    print(f'  D-balance after {MAX_ROUNDS} rounds: {dict(sorted(counts.items()))}')
    return qs


def fix_d_balance_300(questions):
    """Fix D-balance for 300 questions: 75 per option."""
    qs = copy.deepcopy(questions)
    target = 75
    indices = list(range(len(qs)))
    MAX_ROUNDS = 500
    for round_num in range(MAX_ROUNDS):
        counts = Counter(q['ans'] for q in qs)
        over = {l: counts[l] - target for l in 'ABCD' if counts[l] > target}
        under = {l: target - counts[l] for l in 'ABCD' if counts[l] < target}
        if not over and not under:
            print(f'  D-balance in round {round_num + 1}: {dict(sorted(counts.items()))}')
            return qs
        if not over or not under:
            print(f'  Cannot balance further: over={over}, under={under}')
            break
        over_letter = max(over, key=over.get)
        under_letter = max(under, key=under.get)
        random.shuffle(indices)
        for i in indices:
            if qs[i]['ans'] == over_letter:
                qs[i] = swap_answer(qs[i], under_letter)
                break
    counts = Counter(q['ans'] for q in qs)
    print(f'  D-balance after {MAX_ROUNDS} rounds: {dict(sorted(counts.items()))}')
    return qs


# ============================================================
# GENERATE DOCX HELPERS
# ============================================================
def create_103_104_docx(questions, output_path, title_subtitle, prof_name):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    
    title_lines = [
        ('Фонд оценочных средств', True, 16),
        (title_subtitle, True, 14),
        ('', False, 14),
        ('по основной программе профессионального обучения', False, 14),
        ('(профессиональной подготовки)', False, 14),
        ('', False, 14),
        ('по профессии', False, 14),
        (f'\u00AB{PROF}\u00BB', True, 14),
        ('', False, 14),
        ('Тестовые задания с выбором одного правильного ответа (20 вопросов)', False, 14),
    ]
    
    for text, bold, size in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    
    doc.add_paragraph()
    
    for q_idx, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{q_idx}. {q["question"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        for letter in 'ABCD':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'{letter}) {q[letter]};')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    print(f'  Saved: {output_path}')


def create_105_docx(tickets, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    title_lines = [
        ('Фонд оценочных средств', True, 16),
        ('итоговой аттестации', True, 16),
        ('', False, 14),
        ('по основной программе профессионального обучения', False, 14),
        ('(профессиональной подготовки)', False, 14),
        ('', False, 14),
        ('по профессии', False, 14),
        (f'\u00AB{PROF}\u00BB', True, 14),
        ('', False, 14),
        ('Тестовые задания с выбором одного правильного ответа', False, 14),
        ('(20 вопросов)', False, 14),
    ]
    for text, bold, size in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    doc.add_page_break()

    for t_idx, ticket in enumerate(tickets):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Билет {t_idx + 1}')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        doc.add_paragraph()
        
        for q_idx, q in enumerate(ticket, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{q_idx}. {q["question"]}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            for letter in 'ABCD':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(f'{letter}) {q[letter]};')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            p = doc.add_paragraph()
            run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        if t_idx < len(tickets) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f'  Saved: {output_path}')


# ============================================================
# FIX 104
# ============================================================
print('=' * 60)
print('FIXING 104')
print('=' * 60)

path_104 = os.path.join(FOLDER, '104. ПО_П_ФОС_слесарь КИПиА_2-3_разр МДК01.01.docx')
questions_104 = parse_103_104(path_104)
print(f'Parsed: {len(questions_104)} questions')

# Show current Q10
q10 = questions_104[9]
print(f'\nOLD Q10: {q10["question"]}')
print(f'  A: {q10["A"]}')
print(f'  B: {q10["B"]}')
print(f'  C: {q10["C"]}')
print(f'  D: {q10["D"]}')
print(f'  Answer: {q10["ans"]}')

counts_before = Counter(q['ans'] for q in questions_104)
print(f'\nBefore fix D-balance: {dict(sorted(counts_before.items()))}')

# Replace Q10 with correct question about height work safety
# Per Приказ Минтруда РФ от 16.11.2020 № 782н:
# Средства защиты от падения обязательны при работе на высоте 1,8 м и более
questions_104[9] = {
    'num': 10,
    'question': 'Согласно Правилам по охране труда при работе на высоте, средства индивидуальной защиты от падения с высоты обязательно применяются при выполнении работ на высоте:',
    'A': '1,0 м и более',
    'B': '1,3 м и более',
    'C': '1,8 м и более',
    'D': '2,5 м и более',
    'ans': 'C'
}

print(f'\nNEW Q10: {questions_104[9]["question"]}')
print(f'  A: {questions_104[9]["A"]}')
print(f'  B: {questions_104[9]["B"]}')
print(f'  C: {questions_104[9]["C"]}')
print(f'  D: {questions_104[9]["D"]}')
print(f'  Answer: {questions_104[9]["ans"]}')

# Rebalance D
print('\nRebalancing D for 104...')
questions_104 = fix_d_balance_20(questions_104)

# Verify
errors = 0
for q in questions_104:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        print(f'  ERROR Q{q["num"]}: duplicate options!')
        errors += 1
counts_after = Counter(q['ans'] for q in questions_104)
print(f'After fix D-balance: {dict(sorted(counts_after.items()))}')
print(f'Verification: {"OK" if errors == 0 else f"{errors} errors"}')

# Regenerate 104
create_103_104_docx(questions_104, path_104,
    'промежуточной аттестации (МДК01.01) по профессиональному циклу', PROF)


# ============================================================
# FIX 105
# ============================================================
print(f'\n{"="*60}')
print('FIXING 105')
print('=' * 60)

path_105 = os.path.join(FOLDER, '105. ПО_П_ФОС_слесарь КИПиА_2-3_разр.docx')
all_q105, ticket_lines = parse_105(path_105)
print(f'Parsed: {len(all_q105)} questions')

counts_105_before = Counter(q['ans'] for q in all_q105)
print(f'Before fix D-balance: {dict(sorted(counts_105_before.items()))}')

# Find and fix the 3 issues
fix_count = 0

# Fix 1: Кондуктометр (answer should be A, not C)
for i, q in enumerate(all_q105):
    if 'удельную электропроводность' in q['question'] or 'электропроводность жидкости' in q['question']:
        old_ans = q['ans']
        # Find which option is Кондуктометр
        for letter in 'ABCD':
            if 'Кондуктометр' in q[letter] or 'кондуктометр' in q[letter]:
                q['ans'] = letter
                fix_count += 1
                print(f'\nFix 1: Кондуктометр question')
                print(f'  Q: {q["question"][:60]}...')
                print(f'  Old answer: {old_ans} -> New answer: {letter} ({q[letter]})')
                break
        break

# Fix 2: Vacuum gasket contradiction
# Билет 4 Q9: answer D (Резина) should be C (Медь)
# Билет 13 Q12: answer A (Медь) - this is correct, keep it
# We need to find both questions and make consistent
vacuum_questions = []
for i, q in enumerate(all_q105):
    if ('прокладк' in q['question'].lower() and 'вакуум' in q['question'].lower()):
        vacuum_questions.append((i, q))

print(f'\nFound {len(vacuum_questions)} vacuum gasket questions:')
for idx, (i, q) in enumerate(vacuum_questions):
    print(f'  [{idx}] Q at pos {i}: {q["question"][:70]}...')
    print(f'       A: {q["A"]}')
    print(f'       B: {q["B"]}')
    print(f'       C: {q["C"]}')
    print(f'       D: {q["D"]}')
    print(f'       Answer: {q["ans"]}')

# Change the one with answer D(Резина) to answer C(Медь)
for i, q in vacuum_questions:
    # Find which letter has 'Резина' (without metal reinforcement)
    for letter in 'ABCD':
        text = q[letter]
        if 'Резина' in text and 'армирован' not in text.lower():
            # This is the question to fix - change answer to Медь
            for target_letter in 'ABCD':
                if 'Медь' in q[target_letter]:
                    old_ans = q['ans']
                    q['ans'] = target_letter
                    fix_count += 1
                    print(f'\nFix 2: Vacuum gasket - changed answer from {old_ans} to {target_letter} (Медь)')
                    break
            break

# Fix 3: "часто" -> reword to be unambiguous
for i, q in enumerate(all_q105):
    if 'часто' in q['question'].lower() and 'разъ' in q['question'].lower():
        old_q = q['question']
        q['question'] = q['question'].replace('часто используется', 'применяется').replace('часто используется', 'применяется')
        # More thorough replacement
        q['question'] = re.sub(r'\bчасто\b', '', q['question'], flags=re.IGNORECASE)
        q['question'] = re.sub(r'\s{2,}', ' ', q['question']).strip()
        # Ensure it reads well
        q['question'] = q['question'].replace('тип электрического разъема используется для подключения',
                                                'тип электрического разъема предназначен для подключения')
        fix_count += 1
        print(f'\nFix 3: Reworded ambiguous question')
        print(f'  Old: {old_q}')
        print(f'  New: {q["question"]}')
        break

print(f'\nTotal fixes applied: {fix_count}')

# Rebalance D for 105
print('\nRebalancing D for 105...')
all_q105 = fix_d_balance_300(all_q105)

counts_105_after = Counter(q['ans'] for q in all_q105)
print(f'After fix D-balance: {dict(sorted(counts_105_after.items()))}')

# Verify no duplicate options
errors = 0
for q in all_q105:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        errors += 1
        print(f'  ERROR: duplicate options in: {q["question"][:50]}...')
print(f'Verification: {"OK" if errors == 0 else f"{errors} errors"}')

# Verify no duplicate questions
q_texts = [q['question'] for q in all_q105]
dupes = {t: c for t, c in Counter(q_texts).items() if c > 1}
if dupes:
    print(f'WARNING: {len(dupes)} duplicate questions found!')
    for t, c in dupes.items():
        print(f'  "{t[:60]}..." x{c}')
else:
    print('No duplicate questions.')

# Redistribute into 15 tickets with per-ticket balance
print('\nRedistributing into 15 tickets...')
TICKETS = 15
by_letter = {letter: [] for letter in 'ABCD'}
for q in all_q105:
    by_letter[q['ans']].append(q)
for letter in 'ABCD':
    random.shuffle(by_letter[letter])

tickets_105 = [[] for _ in range(TICKETS)]
for letter in 'ABCD':
    for i, q in enumerate(by_letter[letter]):
        tickets_105[i % TICKETS].append(q)
for t in tickets_105:
    random.shuffle(t)

print('Per-ticket D-balance:')
ticket_ok = True
for t_idx, ticket in enumerate(tickets_105):
    tc = Counter(q['ans'] for q in ticket)
    line = f'  Билет {t_idx+1:2d}: ' + ' '.join(f'{l}={tc[l]}' for l in 'ABCD')
    if not all(4 <= tc[l] <= 6 for l in 'ABCD'):
        line += ' ***'
        ticket_ok = False
    print(line)
print(f'All tickets OK: {ticket_ok}')

# Regenerate 105
create_105_docx(tickets_105, path_105)

print(f'\n=== DONE ===')
print(f'104: replaced Q10, rebalanced, regenerated')
print(f'105: fixed {fix_count} issues, rebalanced, regenerated')
