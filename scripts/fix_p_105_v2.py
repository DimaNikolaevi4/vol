#!/usr/bin/env python3
"""
Вторая волна правок П 105: 9 текстовых уточнений (без изменения ключей)

Все исправления — текстовые (формулировки вопросов и варианты ответов).
D-баланс не меняется: остаётся 75/75/75/75.

1. Б5в5:   Вариант A — «долях процента» → «карбида титана в процентах»
2. Б6в8:   Вариант B — «подачи» → «компенсация погрешностей базирования»
             (+ «основная» в вопросе)
3. Б11в14: Полная замена вопроса про IT12/50мм (0,25 → 0,39 мм)
4. Б3в19:  Уточнение «сдвиг» — «взаимным смещением параллельных слоёв»
5. Б14в10: Добавлен «диапазон измерения до 25 мм» (устраняет двусмысленность)
6. Б12в1:  Добавлено «типично» + «при точении стальных заготовок»
7. Б1в11:  «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»
8. Б3в3:   «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»
9. Б15в5:  Добавлены условия «минимальное» + «в трубе до 30 м»
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"
FOLDER = "профессиональной подготовки 15474 Оператор станков"
FILEPATH = f"{BASE}/{FOLDER}/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx"


def parse_105(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = {}
    current_answer = None
    in_questions = False

    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()
        if not txt_stripped:
            header_paras.append(txt)
            continue
        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', txt_stripped)
        if ticket_match:
            in_questions = True
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_ticket = int(ticket_match.group(1))
            current_q_text = None
            header_paras.append(txt)
            continue
        if not in_questions:
            header_paras.append(txt)
            continue
        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', txt_stripped)
        if q_match:
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = {}
            current_answer = None
            continue
        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q_text is not None:
            current_options[opt_match.group(1)] = opt_match.group(2)
            continue
        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)
    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket, 'num': current_q_num,
            'text': current_q_text, 'options': current_options,
            'answer': current_answer
        })
    return header_paras, questions


def compute_balance(questions):
    global_counts = Counter()
    per_ticket = {}
    for q in questions:
        a = q['answer']
        global_counts[a] += 1
        t = q['ticket']
        if t not in per_ticket:
            per_ticket[t] = Counter()
        per_ticket[t][a] += 1
    return global_counts, per_ticket


def apply_fixes(questions):
    n = 0

    for q in questions:
        # 1. Б5в5: вариант A — уточнить «карбида титана в процентах»
        if q['ticket'] == 5 and q['num'] == 5:
            q['options']['A'] = 'Содержание карбида титана в процентах'
            print(f'  Б5в5:  вариант A → «Содержание карбида титана в процентах»')
            n += 1

        # 2. Б6в8: вариант B — «компенсация погрешностей базирования» + «основная»
        elif q['ticket'] == 6 and q['num'] == 8:
            q['text'] = q['text'].replace('функция', 'основная функция')
            q['options']['B'] = 'Компенсация погрешностей базирования заготовки (самоустановка)'
            print(f'  Б6в8:  вариант B → «Компенсация погрешностей базирования...» + «основная»')
            n += 1

        # 3. Б11в14: полная замена вопроса про IT12/50мм
        elif q['ticket'] == 11 and q['num'] == 14:
            q['text'] = 'Согласно ГОСТ 25346, каково ориентировочное значение допуска IT12 для номинального размера 50 мм (интервал 30–50 мм)?'
            q['options'] = {
                'A': '0,025 мм',
                'B': '0,10 мм',
                'C': '0,25 мм',
                'D': '0,39 мм'
            }
            q['answer'] = 'D'
            print(f'  Б11в14: вопрос + варианты заменены (IT12/50мм: 0,39 мм, ответ D)')
            n += 1

        # 4. Б3в19: уточнение «сдвиг»
        elif q['ticket'] == 3 and q['num'] == 19:
            q['text'] = 'Какой вид деформации характеризуется взаимным смещением параллельных слоёв тела без изменения объёма?'
            print(f'  Б3в19: формулировка уточнена → «взаимным смещением параллельных слоёв»')
            n += 1

        # 5. Б14в10: добавить «диапазон до 25 мм»
        elif q['ticket'] == 14 and q['num'] == 10:
            q['text'] = 'Какой измерительный инструмент применяется для контроля наружных размеров с точностью 0,01 мм и диапазоном измерения до 25 мм?'
            print(f'  Б14в10: добавлен «диапазон измерения до 25 мм»')
            n += 1

        # 6. Б12в1: добавить «типично» + «при точении стальных заготовок»
        elif q['ticket'] == 12 and q['num'] == 1:
            q['text'] = 'Какой вид обработки типично обеспечивает шероховатость поверхности Ra 12,5 при точении стальных заготовок?'
            print(f'  Б12в1: добавлено «типично» + «при точении стальных заготовок»')
            n += 1

        # 7. Б1в11: «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»
        elif q['ticket'] == 1 and q['num'] == 11:
            q['text'] = q['text'].replace('Согласно ПТЭЭП', 'Согласно ранее действовавшим ПТЭЭП (утв. 2003 г.)')
            print(f'  Б1в11: «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»')
            n += 1

        # 8. Б3в3: «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»
        elif q['ticket'] == 3 and q['num'] == 3:
            q['text'] = q['text'].replace('Согласно ПТЭЭП', 'Согласно ранее действовавшим ПТЭЭП (утв. 2003 г.)')
            print(f'  Б3в3: «ПТЭЭП» → «ранее действовавшим ПТЭЭП (утв. 2003 г.)»')
            n += 1

        # 9. Б15в5: добавить «минимальное» + «в трубе до 30 м»
        elif q['ticket'] == 15 and q['num'] == 5:
            q['text'] = 'Согласно ПУЭ, какое минимальное сечение медного провода следует выбрать для питания станка мощностью 5 кВт при напряжении 380 В при условии прокладки в трубе длиной до 30 м?'
            print(f'  Б15в5: добавлены условия «минимальное» + «в трубе до 30 м»')
            n += 1

    print(f'\nВсего исправлений: {n}')
    return n


def build_105_docx(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Фонд') or stripped.startswith('итоговой') or
            stripped.startswith('по основной') or stripped.startswith('(профессиональной') or
            stripped.startswith('по профессии') or stripped.startswith('«Оператор') or
            stripped.startswith('Тестовые') or stripped.startswith('(20 вопросов') or
            re.match(r'^Билет\s+№?\s*\d+$', stripped)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if re.match(r'^Билет', stripped) else Pt(12)

    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)

    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph(f'Билет {t}')
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        for q in tickets[t]:
            qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)

            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)

            ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.left_indent = Cm(0)
            for run in ap.runs:
                run.font.bold = True

    doc.save(output_path)
    print(f"\nФайл сохранён: {output_path}")


def main():
    print('=' * 60)
    print('ИСПРАВЛЕНИЕ 105 П (v2): 9 текстовых уточнений')
    print('=' * 60)

    print('\n[1] Парсинг...')
    header_paras, questions = parse_105(FILEPATH)
    print(f'  Вопросов: {len(questions)}')
    gc, pt = compute_balance(questions)
    print(f'  D-баланс до: A={gc["A"]} B={gc["B"]} C={gc["C"]} D={gc["D"]}')

    print('\n[2] Применение исправлений...')
    n = apply_fixes(questions)

    gc, pt = compute_balance(questions)
    print(f'\n  D-баланс после: A={gc["A"]} B={gc["B"]} C={gc["C"]} D={gc["D"]}')

    # Verify
    print('\n[3] Верификация...')
    checks = {
        (5, 5):   ('opt_contains', 'A', 'карбида титана'),
        (6, 8):   ('opt_contains', 'B', 'Компенсация погрешностей базирования'),
        (11, 14): ('opt_contains', 'D', '0,39'),
        (3, 19):  ('text_contains', 'взаимным смещением параллельных слоёв'),
        (14, 10): ('text_contains', 'до 25 мм'),
        (12, 1):  ('text_contains', 'при точении стальных заготовок'),
        (1, 11):  ('text_contains', 'ранее действовавшим ПТЭЭП'),
        (3, 3):   ('text_contains', 'ранее действовавшим ПТЭЭП'),
        (15, 5):  ('text_contains', 'в трубе'),
    }

    all_ok = True
    for q in questions:
        key = (q['ticket'], q['num'])
        if key not in checks:
            continue
        ct, *args = checks[key]
        if ct == 'opt_contains':
            letter, substr = args
            ok = substr in q['options'].get(letter, '')
        elif ct == 'text_contains':
            ok = args[0] in q['text']
        else:
            ok = True
        status = '✅' if ok else '❌'
        if not ok:
            all_ok = False
        print(f'  {status} Б{q["ticket"]}в{q["num"]}')

    # Build
    print('\n[4] Генерация файла...')
    build_105_docx(header_paras, questions, FILEPATH)

    # Re-verify
    print('\n[5] Поверка из файла...')
    _, vq = parse_105(FILEPATH)
    gc2, _ = compute_balance(vq)
    print(f'  Вопросов: {len(vq)}, D-баланс: A={gc2["A"]} B={gc2["B"]} C={gc2["C"]} D={gc2["D"]}')

    if all_ok and gc2['A'] == gc2['B'] == gc2['C'] == gc2['D'] == 75:
        print('\n  🟢 ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('\n  🔴 ЕСТЬ ПРОБЛЕМЫ')

    print('\n' + '=' * 60)
    print('ГОТОВО')
    print('=' * 60)


if __name__ == '__main__':
    main()
