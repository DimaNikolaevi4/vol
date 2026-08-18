#!/usr/bin/env python3
"""
Исправление ошибок в 105. ПО_П_ФОС_Оператор станков_2-3_разр.docx

По результатам экспертной проверки 300 вопросов:

ЯВНЫЕ ОШИБКИ (9):
  Q29  (Б2в9):  ГОСТ 868 → ГОСТ 1050-2013 (текст вопроса, ответ A остаётся)
  Q51  (Б3в11): ключ A→B (разомкнутая линия = для сечений, ГОСТ 2.303)
  Q69  (Б4в9):  ГОСТ 868 → ГОСТ 1050-2013 (текст вопроса, ответ B остаётся)
  Q101 (Б6в1):  ключ C→B (допуск 0,035мм при 100мм = IT7, не IT8, ГОСТ 25346)
  Q108 (Б6в8):  «плавающего суппорта» → «плавающего резцедержателя» + исправление варианта B
  Q158 (Б8в18): ключ B→C (Rz = наибольшая высота профиля, ГОСТ 2789)
  Q159 (Б8в19): значение 5272 Вт → 5265 Вт (√3×380×10×0,8)
  Q171 (Б9в11): ключ B→A (межоперационный припуск = за одну операцию)
  Q298 (Б15в18): ГОСТ 868 → ГОСТ 1050-2013 (текст вопроса, ответ D остаётся)

МЕЛКИЕ НЕТОЧНОСТИ (3):
  Q49  (Б3в9):  добавить «прямым зубчатым зацеплением» (устраняет двусмысленность с цепной)
  Q262 (Б14в2): «площадь формата» → «размеры формата» (варианты дают линейные размеры)
  Q296 (Б15в16): «пересекающимися» → «перекрещивающимися» (червячная = перекрещивающиеся, не пересекающиеся)

D-баланс:
  4 ключевых изменения взаимно компенсируются: A→B, C→B, B→C, B→A
  Итоговый D-баланс остаётся 75/75/75/75 — ребалансировка не требуется
"""

import re
import copy
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"
FOLDER = "профессиональной подготовки 15474 Оператор станков"
FILEPATH = f"{BASE}/{FOLDER}/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx"


def parse_105(filepath):
    """Parse 105 file into structured data."""
    doc = Document(filepath)

    header_paras = []
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = {}
    current_answer = None
    in_questions = False

    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()

        if not txt_stripped:
            header_paras.append(txt)
            continue

        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', txt_stripped)
        if ticket_match:
            in_questions = True
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_ticket = int(ticket_match.group(1))
            current_q_text = None
            header_paras.append(txt)
            continue

        if not in_questions:
            header_paras.append(txt)
            continue

        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', txt_stripped)
        if q_match:
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = {}
            current_answer = None
            continue

        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q_text is not None:
            current_options[opt_match.group(1)] = opt_match.group(2)
            continue

        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)

    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket, 'num': current_q_num,
            'text': current_q_text, 'options': current_options,
            'answer': current_answer
        })

    return header_paras, questions


def compute_balance(questions):
    """Compute D-balance globally and per-ticket."""
    global_counts = Counter()
    per_ticket = {}
    for q in questions:
        a = q['answer']
        global_counts[a] += 1
        t = q['ticket']
        if t not in per_ticket:
            per_ticket[t] = Counter()
        per_ticket[t][a] += 1
    return global_counts, per_ticket


def apply_fixes(questions):
    """Apply all 12 fixes: 4 key changes + 8 text changes."""
    fixes_applied = []

    # === KEY CHANGES (4) ===
    key_fixes = {
        (3, 11): ('A', 'B'),   # Q51:  разомкнутая линия = для сечений (ГОСТ 2.303)
        (6, 1):  ('C', 'B'),   # Q101: допуск 0,035мм/100мм = IT7 (ГОСТ 25346)
        (8, 18): ('B', 'C'),   # Q158: Rz = наибольшая высота профиля (ГОСТ 2789)
        (9, 11): ('B', 'A'),   # Q171: межоперационный припуск = за одну операцию
    }

    for q in questions:
        key = (q['ticket'], q['num'])
        if key in key_fixes:
            old_ans, new_ans = key_fixes[key]
            if q['answer'] == old_ans:
                print(f"  КЛЮЧ Б{q['ticket']}в{q['num']}: {old_ans}→{new_ans}  |  {q['text'][:90]}...")
                q['answer'] = new_ans
                fixes_applied.append(f"Б{q['ticket']}в{q['num']}: ключ {old_ans}→{new_ans}")
            elif q['answer'] == new_ans:
                print(f"  Б{q['ticket']}в{q['num']}: ключ уже {new_ans}, пропускаем")
            else:
                print(f"  ⚠️ Б{q['ticket']}в{q['num']}: ожидаемый ключ {old_ans}, найден {q['answer']}")

    # === TEXT FIXES (8) ===

    # Q29 (Б2в9): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 2 and q['num'] == 9:
            old = q['text']
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б2в9: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б2в9: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Q69 (Б4в9): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 4 and q['num'] == 9:
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б4в9: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б4в9: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Q298 (Б15в18): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 15 and q['num'] == 18:
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б15в18: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б15в18: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Q108 (Б6в8): «плавающего суппорта» → «плавающего резцедержателя» + вариант B
    for q in questions:
        if q['ticket'] == 6 and q['num'] == 8:
            q['text'] = q['text'].replace('плавающего суппорта токарного станка',
                                          'плавающего резцедержателя при обработке центровых отверстий')
            if q['options'].get('B') == 'Автоматическое поддержание постоянной подачи':
                q['options']['B'] = 'Компенсация погрешностей центрирования инструмента'
            print(f"  ТЕКСТ Б6в8: «плавающего суппорта» → «плавающего резцедержателя», вариант B исправлен")
            fixes_applied.append("Б6в8: плавающий суппорт → плавающий резцедержатель + вариант B")
            break

    # Q159 (Б8в19): 5272 Вт → 5265 Вт (√3×380×10×0,8 ≈ 5265)
    for q in questions:
        if q['ticket'] == 8 and q['num'] == 19:
            for letter in q['options']:
                if '5272' in q['options'][letter]:
                    q['options'][letter] = q['options'][letter].replace('5272', '5265')
                    print(f"  ЗНАЧЕНИЕ Б8в19: 5272 Вт → 5265 Вт (вариант {letter})")
                    fixes_applied.append(f"Б8в19: 5272→5265 Вт")
                    break
            break

    # Q49 (Б3в9): добавить «прямым зубчатым зацеплением» для устранения двусмысленности
    for q in questions:
        if q['ticket'] == 3 and q['num'] == 9:
            q['text'] = q['text'].replace('параллельными валами',
                                          'параллельными валами прямым зубчатым зацеплением')
            print(f"  ТЕКСТ Б3в9: добавлено «прямым зубчатым зацеплением»")
            fixes_applied.append("Б3в9: +«прямым зубчатым зацеплением»")
            break

    # Q262 (Б14в2): «площадь формата» → «размеры формата»
    for q in questions:
        if q['ticket'] == 14 and q['num'] == 2:
            q['text'] = q['text'].replace('площадь формата А1', 'размеры (длина и ширина) формата А1')
            print(f"  ТЕКСТ Б14в2: «площадь формата» → «размеры (длина и ширина) формата»")
            fixes_applied.append("Б14в2: площадь → размеры формата")
            break

    # Q296 (Б15в16): «пересекающимися» → «перекрещивающимися»
    for q in questions:
        if q['ticket'] == 15 and q['num'] == 16:
            q['text'] = q['text'].replace('пересекающимися', 'перекрещивающимися')
            print(f"  ТЕКСТ Б15в16: «пересекающимися» → «перекрещивающимися»")
            fixes_applied.append("Б15в16: пересекающимися → перекрещивающимися")
            break

    print(f"\nВсего исправлений: {len(fixes_applied)}")
    return len(fixes_applied)


def build_105_docx(header_paras, questions, output_path):
    """Build the 105 docx file with corrected content."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Write header (non-question paragraphs before tickets)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Фонд') or stripped.startswith('итоговой') or
            stripped.startswith('по основной') or stripped.startswith('(профессиональной') or
            stripped.startswith('по профессии') or stripped.startswith('«Оператор') or
            stripped.startswith('Тестовые') or stripped.startswith('(20 вопросов') or
            re.match(r'^Билет\s+№?\s*\d+$', stripped)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if re.match(r'^Билет', stripped) else Pt(12)

    # Group questions by ticket
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)

    # Write tickets
    for t in sorted(tickets.keys()):
        # Ticket header
        tp = doc.add_paragraph(f'Билет {t}')
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        for q in tickets[t]:
            # Question text
            qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)

            # Options
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)

            # Inline answer
            ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.left_indent = Cm(0)
            for run in ap.runs:
                run.font.bold = True

    doc.save(output_path)
    print(f"\nФайл сохранён: {output_path}")


def main():
    print("=" * 60)
    print("ИСПРАВЛЕНИЕ 105 П: 9 явных ошибок + 3 мелких неточности")
    print("=" * 60)

    # Step 1: Parse
    print("\n[1] Парсинг...")
    header_paras, questions = parse_105(FILEPATH)
    print(f"  Вопросов: {len(questions)}")

    gc, pt = compute_balance(questions)
    print(f"  Исходный D-баланс: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    # Step 2: Apply all fixes
    print("\n[2] Применение исправлений...")
    n_fixes = apply_fixes(questions)

    # Step 3: Verify D-balance
    print("\n[3] Проверка D-баланса...")
    gc, pt = compute_balance(questions)
    print(f"  D-баланс после исправлений: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    if gc['A'] == 75 and gc['B'] == 75 and gc['C'] == 75 and gc['D'] == 75:
        print("  Глобальный баланс: OK (75/75/75/75)")
    else:
        print("  ⚠️ Глобальный баланс нарушен — нужна ребалансировка!")

    # Check per-ticket
    bad_tickets = []
    for t in sorted(pt.keys()):
        for letter in 'ABCD':
            cnt = pt[t][letter]
            if cnt < 4 or cnt > 6:
                bad_tickets.append(f"Билет {t}: {letter}={cnt}")
    if bad_tickets:
        print(f"  ⚠️ Нарушения локального баланса: {bad_tickets}")
    else:
        print("  Локальный баланс: OK (4-6 в каждом билете)")

    # Step 4: Verify fixed questions
    print("\n[4] Верификация исправленных вопросов...")
    checks = {
        (2, 9):   ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
        (3, 9):   ('text_contains', 'прямым зубчатым зацеплением', None),
        (3, 11):  ('key', 'B'),
        (4, 9):   ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
        (6, 1):   ('key', 'B'),
        (6, 8):   ('text_contains', 'плавающего резцедержателя', 'плавающего суппорта'),
        (8, 18):  ('key', 'C'),
        (8, 19):  ('option_no_contains', '5272'),
        (9, 11):  ('key', 'A'),
        (14, 2):  ('text_contains', 'размеры', 'площадь'),
        (15, 16): ('text_contains', 'перекрещивающимися', 'пересекающимися'),
        (15, 18): ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
    }

    all_ok = True
    for q in questions:
        key = (q['ticket'], q['num'])
        if key not in checks:
            continue
        check = checks[key]
        check_type = check[0]

        if check_type == 'key':
            expected = check[1]
            ok = q['answer'] == expected
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: ключ={q['answer']} (ожидается {expected})")
            if not ok: all_ok = False

        elif check_type == 'text_contains':
            should_have = check[1]
            should_not = check[2]
            has_good = should_have in q['text']
            has_bad = should_not and should_not in q['text']
            ok = has_good and not has_bad
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: содержит «{should_have}»={'да' if has_good else 'НЕТ'}, «{should_not}»={'да' if has_bad else 'нет'}")
            if not ok: all_ok = False

        elif check_type == 'option_no_contains':
            bad_str = check[1]
            found_bad = any(bad_str in v for v in q['options'].values())
            ok = not found_bad
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: варианты {'не ' if ok else ''}содержат «{bad_str}»")
            if not ok: all_ok = False

    # Step 5: Build
    print("\n[5] Генерация файла...")
    build_105_docx(header_paras, questions, FILEPATH)

    # Step 6: Re-parse and verify
    print("\n[6] Повторная верификация из файла...")
    _, verify_qs = parse_105(FILEPATH)
    gc2, pt2 = compute_balance(verify_qs)
    print(f"  Вопросов из файла: {len(verify_qs)}")
    print(f"  D-баланс из файла: A={gc2['A']} B={gc2['B']} C={gc2['C']} D={gc2['D']}")

    no_answer = [q for q in verify_qs if not q['answer']]
    if no_answer:
        print(f"  ⚠️ Вопросов без ответа: {len(no_answer)}")
        all_ok = False
    else:
        print(f"  Все {len(verify_qs)} вопросов имеют ответы")

    if all_ok:
        print("\n  🟢 ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ")
    else:
        print("\n  🔴 ЕСТЬ ПРОБЛЕМЫ")

    print("\n" + "=" * 60)
    print("ГОТОВО")
    print("=" * 60)


if __name__ == '__main__':
    main()
