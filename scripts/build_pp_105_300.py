#!/usr/bin/env python3
"""
Build ПП 105: 15 tickets x 20 questions = 300 questions total.
74 existing + 150 new + 76 extra = 300 unique, D-balanced.
"""
import json, random, copy, sys, os
sys.path.insert(0, '/home/z/my-project/vol/scripts')
from new_150_questions import NEW_QUESTIONS
from extra_76_questions import EXTRA_QUESTIONS

random.seed(42)

FOLDER = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной переподготовки 15474 Оператор станков"
OUTPUT = os.path.join(FOLDER, "105. ПО_ПП_ФОС_Оператор станков_2-3_разр.docx")

# ============================================================
# 1. Load existing 74 unique questions
# ============================================================
with open('/home/z/my-project/vol/scripts/pp_existing_unique.json', 'r', encoding='utf-8') as f:
    existing = json.load(f)

print(f"Loaded {len(existing)} existing unique questions")
print(f"Loaded {len(NEW_QUESTIONS)} new questions")
print(f"Loaded {len(EXTRA_QUESTIONS)} extra questions")

# Build combined list
all_questions = []

for q in existing:
    all_questions.append({
        'topic': 'existing', 'question': q['question'],
        'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D'],
        'ans': q['ans']
    })

for topic, q_text, a, b, c, d, ans in NEW_QUESTIONS:
    all_questions.append({
        'topic': topic, 'question': q_text,
        'A': a, 'B': b, 'C': c, 'D': d, 'ans': ans
    })

for topic, q_text, a, b, c, d, ans in EXTRA_QUESTIONS:
    all_questions.append({
        'topic': topic, 'question': q_text,
        'A': a, 'B': b, 'C': c, 'D': d, 'ans': ans
    })

print(f"Total: {len(all_questions)} questions")

# Check duplicates
from collections import Counter
texts = [q['question'] for q in all_questions]
dupes = {t: c for t, c in Counter(texts).items() if c > 1}
if dupes:
    print(f"WARNING: {len(dupes)} duplicates!")
    for t, c in dupes.items():
        print(f"  [{c}x] {t[:60]}...")
    sys.exit(1)
else:
    print("No duplicate questions.")

# ============================================================
# 2. D-balance
# ============================================================
target_per_option = len(all_questions) // 4  # 75

def swap_answer(q, target_letter):
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q

counts = Counter(q['ans'] for q in all_questions)
print(f"Before D-balance: {dict(sorted(counts.items()))}")

indices = list(range(len(all_questions)))
random.shuffle(indices)

MAX_ROUNDS = 200
for round_num in range(MAX_ROUNDS):
    counts = Counter(q['ans'] for q in all_questions)
    over = {l: counts[l] - target_per_option for l in 'ABCD' if counts[l] > target_per_option}
    under = {l: target_per_option - counts[l] for l in 'ABCD' if counts[l] < target_per_option}
    if not over and not under:
        print(f"D-balance achieved in round {round_num + 1}")
        break
    over_letter = next(iter(over))
    under_letter = next(iter(under))
    for i in indices:
        if all_questions[i]['ans'] == over_letter:
            all_questions[i] = swap_answer(copy.deepcopy(all_questions[i]), under_letter)
            break

counts = Counter(q['ans'] for q in all_questions)
print(f"After D-balance: {dict(sorted(counts.items()))}")

# Verify
errors = 0
for q in all_questions:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        errors += 1
if errors:
    print(f"{errors} errors!")
    sys.exit(1)
else:
    print("All 300 answers verified.")

# ============================================================
# 3. Distribute into 15 tickets x 20 with ticket-level D-balance
# ============================================================
TICKETS = 15
Q_PER_TICKET = 20

by_letter = {letter: [] for letter in 'ABCD'}
for q in all_questions:
    by_letter[q['ans']].append(q)

for letter in 'ABCD':
    random.shuffle(by_letter[letter])

tickets = [[] for _ in range(TICKETS)]
for letter in 'ABCD':
    group = by_letter[letter]
    for i, q in enumerate(group):
        ticket_idx = i % TICKETS
        tickets[ticket_idx].append(q)

for t in tickets:
    random.shuffle(t)

print(f"Formed {len(tickets)} tickets x {len(tickets[0])} questions")

# ============================================================
# 4. Generate DOCX
# ============================================================
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

title_lines = [
    ("Фонд оценочных средств", True, 16),
    ("итоговой аттестации", True, 16),
    ("", False, 14),
    ("по основной программе профессионального обучения", False, 14),
    ("(профессиональной переподготовки)", False, 14),
    ("", False, 14),
    ("по профессии", False, 14),
    ("\u00ABОператор автоматических и полуавтоматических станков и линий станков\u00BB", True, 14),
    ("", False, 14),
    ("Тестовые задания с выбором одного правильного ответа", False, 14),
    ("(20 вопросов)", False, 14),
]

for text, bold, size in title_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

doc.add_page_break()

for t_idx, ticket in enumerate(tickets):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Билет {t_idx + 1}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    doc.add_paragraph()
    
    for q_idx, q in enumerate(ticket, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{q_idx}. {q['question']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        for letter in 'ABCD':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"{letter}) {q[letter]};")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        run = p.add_run(f"Правильный ответ под номером: {q['ans']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    if t_idx < len(tickets) - 1:
        doc.add_page_break()

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")

# Summary
print(f"\n=== SUMMARY ===")
print(f"Total: {len(all_questions)}, Tickets: {len(tickets)} x {Q_PER_TICKET}")
print(f"D-balance: {dict(sorted(counts.items()))}")
for l in 'ABCD':
    print(f"  {l}: {counts[l]} ({counts[l]/len(all_questions)*100:.1f}%)")

topic_counts = Counter(q['topic'] for q in all_questions)
for t, c in sorted(topic_counts.items()):
    print(f"  {t}: {c}")
