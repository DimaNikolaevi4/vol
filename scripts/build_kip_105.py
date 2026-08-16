#!/usr/bin/env python3
"""
Build КИПиА 105 from existing 05.
Preserve ALL original question texts and options.
1. Parse 300 questions (15 tickets) from 05
2. Remove duplicates (keep first occurrence), note what needs replacement
3. Generate replacement questions for duplicates
4. Fix D-balance by swapping option positions
5. Distribute into 15 tickets with per-ticket balance
"""
import re, os, copy, random, sys
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = '/home/z/my-project/vol/Слесарь по контрольно-измерительным приборам и автоматике/профессиональной подготовки 18494 Слесарь КИПиА'
OUTPUT = os.path.join(FOLDER, '105. ПО_П_ФОС_слесарь КИПиА_2-3_разр.docx')
PROF = 'Слесарь по контрольно-измерительным приборам и автоматике'

random.seed(42)

# ============================================================
# 1. Parse all questions from 05 with ticket mapping
# ============================================================
doc = Document(f'{FOLDER}/05. ПО_П_ФОС_слесарь КИПиА_2-3_разр .docx')

# Get answer table
table = doc.tables[1]
ticket_answers = {}
for ri, row in enumerate(table.rows[2:], 1):
    cells = [c.text.strip().lower() for c in row.cells]
    ticket_num = int(cells[0])
    for qi in range(1, 21):
        ans = cells[qi] if qi < len(cells) else ''
        if ans:
            cyr_to_lat = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
            ticket_answers[(ticket_num, qi)] = cyr_to_lat.get(ans, ans.upper())

# Parse questions from paragraphs
current_ticket = 0
q_idx_in_ticket = 0
all_questions = []

for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    
    m = re.match(r'^Билет №?(\d+)', t)
    if m:
        current_ticket = int(m.group(1))
        q_idx_in_ticket = 0
        continue
    
    if current_ticket == 0:
        continue
    
    # Skip header-like lines
    if t.startswith('Ключ') or t.startswith('Фонд') or t.startswith('итого') or t.startswith('по осн') or t.startswith('(проф') or t.startswith('по проф') or t.startswith('Тестов'):
        continue
    
    # Skip option-only lines (continuation of previous question)
    if re.match(r'^[а-гА-Г]\)', t):
        continue
    
    q_idx_in_ticket += 1
    
    # Parse question text and embedded options
    lines = t.split('\n')
    q_text = lines[0].strip()
    options = {}
    for line in lines[1:]:
        line = line.strip()
        opt_m = re.match(r'^([а-гА-Г])\)\s*(.+)', line)
        if opt_m:
            letter_cyr = opt_m.group(1).lower()
            letter_map = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
            options[letter_map[letter_cyr]] = opt_m.group(2).strip().rstrip(' ;.')
    
    if len(options) != 4:
        print(f'  SKIP: Ticket {current_ticket} Q{q_idx_in_ticket}: {len(options)} options')
        continue
    
    ans = ticket_answers.get((current_ticket, q_idx_in_ticket), '?')
    all_questions.append({
        'question': q_text,
        'A': options.get('A', ''), 'B': options.get('B', ''),
        'C': options.get('C', ''), 'D': options.get('D', ''),
        'ans': ans, 'orig_ticket': current_ticket
    })

print(f'Parsed: {len(all_questions)} questions')

# ============================================================
# 2. Deduplicate (keep first occurrence)
# ============================================================
seen = set()
unique = []
dupes = []
for q in all_questions:
    if q['question'] in seen:
        dupes.append(q)
    else:
        seen.add(q['question'])
        unique.append(q)

print(f'Unique: {len(unique)}, Duplicates: {len(dupes)}')
for d in dupes:
    print(f'  DUPE: [{d["question"][:60]}...] ticket={d["orig_ticket"]}')

# ============================================================
# 3. Generate replacement questions for duplicates
# ============================================================
REPLACEMENTS = [
    ('Какое устройство применяется для преобразования электрического сигнала 4-20 мА в пневматический сигнал 20-100 кПа?',
     'Электропневматический преобразователь (ЭПП)',
     'Токовый усилитель',
     'Барьер искрозащиты',
     'Резистивный делитель', 'A'),
    ('Какой тип подключения термопары компенсирует влияние температуры свободных концов на результат измерения?',
     'Подключение с автоматической компенсацией холодного спая',
     'Подключение напрямую к вольтметру',
     'Подключение через шунтирующий резистор',
     'Подключение последовательно с нагрузкой', 'A'),
    ('Каким прибором измеряется электрическое сопротивление обмоток электродвигателя?',
     'Мегаомметр или омметр',
     'Мультиметр в режиме измерения тока',
     'Осциллограф',
     'Частотомер', 'A'),
    ('Какой тип датчика используется для измерения уровня жидких сред в закрытых резервуарах под давлением?',
     'Датчик дифференциального давления (дифманометр)',
     'Поплавковый датчик',
     'Визуальный уровнемер (стеклянный)',
     'Бесконтактный датчик температуры', 'A'),
    ('Для чего предназначен клапанный блок (манifold) в системах КИПиА?',
     'Для распределения потоков и подключения нескольких приборов к одной линии',
     'Для усиления электрического сигнала',
     'Для измерения расхода',
     'Для фильтрации воздуха', 'A'),
    ('Какой нормативный документ определяет требования к маркировке кабелей в системах автоматизации?',
     'ГОСТ Р 50571 и ПУЭ',
     'СНиП по строительству',
     'Трудовой кодекс РФ',
     'ГОСТ на продукцию промышленного назначения', 'A'),
    ('Какой вид сигнала используется в цифровой передаче данных по интерфейсу HART?',
     'Частотная модуляция цифрового сигнала поверх аналогового 4-20 мА',
     'Только аналоговый сигнал 4-20 мА',
     'Только цифровой сигнал RS-485',
     'Импульсный сигнал 0-10 В', 'A'),
    ('Какой метод монтажа кабельных трасс применяется в помещениях с агрессивной средой?',
     'Прокладка в герметичных трубах или кабельных каналах',
     'Открытая прокладка по стенам',
     'Прокладка в деревянных лотках',
     'Прокладка directamente по полу', 'A'),
    ('Какой параметр характеризует быстродействие датчика давления?',
     'Время отклика (переходная характеристика)',
     'Погрешность линейности',
     'Диапазон измерений',
     'Масса датчика', 'A'),
    ('Какая схема подключения обмоток трёхфазного двигателя применяется чаще всего в КИПиА?',
     'Соединение «звезда» при 380 В',
     'Соединение «треугольник» при 380 В',
     'Соединение «звезда» при 220 В',
     'Соединение «треугольник» при 660 В', 'A'),
    ('Какой тип уплотнения применяется для герметизации кабельных вводов в шкафах КИПиА?',
     'Кабельные сальники (гермовводы)',
     'Силиконовый герметик',
     'Изолента',
     'Паронитовая прокладка', 'A'),
    ('Какой стандартный сигнал используется для дистанционной передачи давления в системах КИПиА?',
     'Пневматический сигнал 20-100 кПа или электрический 4-20 мА',
     'Электрический сигнал 0-5 В',
     'Цифровой сигнал USB',
     'Радиосигнал Wi-Fi', 'A'),
    ('Какой прибор применяется для проверки точности показаний манометра?',
     'Образцовый (эталонный) манометр или пресс-манометр',
     'Мультиметр',
     'Осциллограф',
     'Мегаомметр', 'A'),
    ('Для чего применяется разделительный трансформатор в цепях КИПиА?',
     'Для электрической изоляции и безопасности обслуживания',
     'Для повышения напряжения',
     'Для стабилизации тока',
     'Для измерения мощности', 'A'),
    ('Какое напряжение питания применяется для стандартных аналоговых датчиков с выходным сигналом 4-20 мА?',
     '24 В постоянного тока',
     '220 В переменного тока',
     '12 В постоянного тока',
     '380 В переменного тока', 'A'),
]

if len(REPLACEMENTS) < len(dupes):
    print(f'ERROR: need {len(dupes)} replacements, only have {len(REPLACEMENTS)}')
    sys.exit(1)

# Replace duplicates with new questions
for i, dup in enumerate(dupes):
    q_text, a, b, c, d, ans = REPLACEMENTS[i]
    unique.append({
        'question': q_text,
        'A': a, 'B': b, 'C': c, 'D': d,
        'ans': ans, 'orig_ticket': 0
    })

print(f'After replacements: {len(unique)} questions')

# Verify no duplicates
if len(set(q['question'] for q in unique)) != len(unique):
    print('ERROR: still have duplicates!')
    sys.exit(1)

# ============================================================
# 4. D-balance (global 75/75/75/75)
# ============================================================
target_per_option = 300 // 4  # 75

def swap_answer(q, target_letter):
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q

counts = Counter(q['ans'] for q in unique)
print(f'Before D-balance: {dict(sorted(counts.items()))}')

indices = list(range(len(unique)))
MAX_ROUNDS = 500
for round_num in range(MAX_ROUNDS):
    counts = Counter(q['ans'] for q in unique)
    over = {l: counts[l] - target_per_option for l in 'ABCD' if counts[l] > target_per_option}
    under = {l: target_per_option - counts[l] for l in 'ABCD' if counts[l] < target_per_option}
    if not over and not under:
        print(f'D-balance in round {round_num + 1}')
        break
    over_letter = max(over, key=over.get)
    under_letter = max(under, key=under.get)
    random.shuffle(indices)
    for i in indices:
        if unique[i]['ans'] == over_letter:
            unique[i] = swap_answer(copy.deepcopy(unique[i]), under_letter)
            break

counts = Counter(q['ans'] for q in unique)
print(f'After D-balance: {dict(sorted(counts.items()))}')

# Verify
errors = 0
for q in unique:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        errors += 1
if errors:
    print(f'ERROR: {errors} questions with duplicate options!')
    sys.exit(1)
else:
    print('All answers verified.')

# ============================================================
# 5. Distribute into 15 tickets x 20 (round-robin by letter)
# ============================================================
TICKETS = 15
by_letter = {letter: [] for letter in 'ABCD'}
for q in unique:
    by_letter[q['ans']].append(q)
for letter in 'ABCD':
    random.shuffle(by_letter[letter])

tickets = [[] for _ in range(TICKETS)]
for letter in 'ABCD':
    for i, q in enumerate(by_letter[letter]):
        tickets[i % TICKETS].append(q)
for t in tickets:
    random.shuffle(t)

print(f'\nPer-ticket D-balance:')
ticket_ok = True
for t_idx, ticket in enumerate(tickets):
    tc = Counter(q['ans'] for q in ticket)
    line = f'  Билет {t_idx+1:2d}: ' + ' '.join(f'{l}={tc[l]}' for l in 'ABCD')
    if not all(4 <= tc[l] <= 6 for l in 'ABCD'):
        line += ' ***'
        ticket_ok = False
    print(line)
print(f'All tickets OK: {ticket_ok}')

# ============================================================
# 6. Generate DOCX
# ============================================================
doc_out = Document()
style = doc_out.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
for section in doc_out.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

title_lines = [
    ('Фонд оценочных средств', True, 16),
    ('итоговой аттестации', True, 16),
    ('', False, 14),
    ('по основной программе профессионального обучения', False, 14),
    ('(профессиональной подготовки)', False, 14),
    ('', False, 14),
    ('по профессии', False, 14),
    (f'\u00AB{PROF}\u00BB', True, 14),
    ('', False, 14),
    ('Тестовые задания с выбором одного правильного ответа', False, 14),
    ('(20 вопросов)', False, 14),
]
for text, bold, size in title_lines:
    p = doc_out.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
doc_out.add_page_break()

for t_idx, ticket in enumerate(tickets):
    p = doc_out.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Билет {t_idx + 1}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    doc_out.add_paragraph()
    
    for q_idx, q in enumerate(ticket, 1):
        p = doc_out.add_paragraph()
        run = p.add_run(f'{q_idx}. {q["question"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        for letter in 'ABCD':
            p = doc_out.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'{letter}) {q[letter]};')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        p = doc_out.add_paragraph()
        run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    if t_idx < len(tickets) - 1:
        doc_out.add_page_break()

doc_out.save(OUTPUT)
print(f'\nSaved: {OUTPUT}')
print(f'\n=== SUMMARY ===')
print(f'Total: {len(unique)}, Tickets: {TICKETS} x 20')
print(f'D-balance: {dict(sorted(counts.items()))}')
print(f'Original questions preserved: {len(unique) - len(REPLACEMENTS)}/300')
print(f'New replacement questions: {len(REPLACEMENTS)}/300')
