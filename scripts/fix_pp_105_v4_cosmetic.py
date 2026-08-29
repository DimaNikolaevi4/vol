#!/usr/bin/env python3
"""
ПП 105 v4: 3 косметических текст-фикса (без замены вопросов, без изменения ключей)
- Б4в19: +"с плоских поверхностей"
- Б11в1: D дистрактор "если руки сухие" -> "(без прикосновения)"
- Б2в16, Б5в4: ГОСТ 25347 -> ГОСТ 25346
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'
FOLDER = 'профессиональной подготовки 15474 Оператор станков'
FILEPATH = BASE + '/' + FOLDER + '/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx'

Q_RE = re.compile(r'^(\d+)\s*[.)]\s*(.*)')
OPT_RE = re.compile(r'^([ABCD])\)\s*(.*)')
ANS_RE = re.compile(r'Правильный ответ под номером:\s*([ABCD])')
TICKET_RE = re.compile(r'Билет\s+№?\s*(\d+)')

# Текст-фиксы: {(билет, вопрос): [(old, new), ...]}
TEXT_FIXES = {
    (4, 19): [
        ('при снятии большого припуска?', 'при снятии большого припуска с плоских поверхностей?'),
    ],
    (11, 1): [
        ('Не нормируется, если руки сухие', 'Не нормируется (без прикосновения)'),
    ],
    (2, 16): [
        ('Согласно ГОСТ 25347', 'Согласно ГОСТ 25346'),
    ],
    (5, 4): [
        ('Согласно ГОСТ 25347', 'Согласно ГОСТ 25346'),
    ],
}


def parse_105(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    in_questions = False
    cur_ticket = None
    cur_q = None
    cur_text = None
    cur_opts = {}
    cur_ans = None

    for p in doc.paragraphs:
        txt = p.text
        txt_s = txt.strip()
        if not txt_s:
            header_paras.append(txt)
            continue
        tm = TICKET_RE.match(txt_s)
        if tm:
            in_questions = True
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_ticket = int(tm.group(1))
            cur_text = None
            header_paras.append(txt)
            continue
        if not in_questions:
            header_paras.append(txt)
            continue
        m = Q_RE.match(txt_s)
        if m:
            if cur_text is not None:
                questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_q = int(m.group(1))
            cur_text = m.group(2)
            cur_opts = {}
            cur_ans = None
            continue
        m2 = OPT_RE.match(txt_s)
        if m2 and cur_text is not None:
            cur_opts[m2.group(1)] = m2.group(2)
            continue
        m3 = ANS_RE.search(txt_s)
        if m3 and cur_text is not None:
            cur_ans = m3.group(1)
    if cur_text is not None:
        questions.append({'ticket': cur_ticket, 'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
    return header_paras, questions


def apply_text_fixes(questions):
    n_fixes = 0
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in TEXT_FIXES:
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            for old, new in TEXT_FIXES[key]:
                # Check question text
                if old in q['text']:
                    q['text'] = q['text'].replace(old, new)
                    print('  FIX ' + tag + ' Q: "' + old[:50] + '" -> "' + new[:50] + '"')
                    n_fixes += 1
                # Check options
                for letter in 'ABCD':
                    if letter in q['options'] and old in q['options'][letter]:
                        q['options'][letter] = q['options'][letter].replace(old, new)
                        print('  FIX ' + tag + ' opt ' + letter + ': "' + old[:50] + '" -> "' + new[:50] + '"')
                        n_fixes += 1
    return n_fixes


def build_105(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if not s:
            continue
        if (s.startswith('Фонд') or 'итоговой' in s or 'по основной' in s or
            s.startswith('(профессиональной') or 'по профессии' in s or
            s.startswith('«Оператор') or 'Тестовые' in s or
            '(20 вопросов' in s or TICKET_RE.match(s)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if TICKET_RE.match(s) else Pt(12)
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)
    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph('Билет ' + str(t))
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        for q in tickets[t]:
            qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)
            ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            for run in ap.runs:
                run.font.bold = True
    doc.save(output_path)


def main():
    print('=' * 60)
    print('ПП 105 v4: 3 косметических текст-фикса')
    print('=' * 60)

    print('\n[1] Парсинг...')
    header_paras, questions = parse_105(FILEPATH)
    print('  Вопросов: ' + str(len(questions)))
    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс до: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))

    print('\n[2] Текст-фиксы...')
    n = apply_text_fixes(questions)
    print('  Применено: ' + str(n))

    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс после: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))
    if gc['A'] == gc['B'] == gc['C'] == gc['D'] == 75:
        print('  D-баланс ОК')
    else:
        print('  D-БАЛАНС НАРУШЕН!')
        return

    print('\n[3] Генерация файла...')
    build_105(header_paras, questions, FILEPATH)

    print('\n[4] Верификация...')
    _, vq = parse_105(FILEPATH)
    gc2 = Counter(q['answer'] for q in vq)
    print('  Вопросов: ' + str(len(vq)) + ', D-баланс: A=' + str(gc2['A']) + ' B=' + str(gc2['B']) + ' C=' + str(gc2['C']) + ' D=' + str(gc2['D']))

    ok = True
    for q in vq:
        key = (q['ticket'], q['num'])
        if key in TEXT_FIXES:
            tag = 'Б' + str(q['ticket']) + 'в' + str(q['num'])
            for old, new in TEXT_FIXES[key]:
                found_in_q = new in q['text']
                found_in_opts = any(new in q['options'].get(l, '') for l in 'ABCD')
                if found_in_q or found_in_opts:
                    print('  OK ' + tag + ': фикс применён')
                else:
                    print('  FAIL ' + tag + ': фикс НЕ найден')
                    ok = False

    if ok and gc2['A'] == gc2['B'] == gc2['C'] == gc2['D'] == 75:
        print('\n  ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('\n  ЕСТЬ ПРОБЛЕМЫ')
    print('=' * 60)


if __name__ == '__main__':
    main()