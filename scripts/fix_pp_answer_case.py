#!/usr/bin/env python3
"""Fix answer lines in ПП 103/104: а/б/в/г -> A/B/C/D"""
import re, os
from docx import Document

FOLDER = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной переподготовки 15474 Оператор станков'

cyr_to_lat = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}

for prefix in ['103', '104']:
    fname = [f for f in os.listdir(FOLDER) if f.startswith(prefix + '.')][0]
    path = os.path.join(FOLDER, fname)
    doc = Document(path)
    
    fixed = 0
    for p in doc.paragraphs:
        m = re.match(r'^(Правильный ответ под номером: )([абвг])$', p.text.strip())
        if m:
            new_letter = cyr_to_lat[m.group(2)]
            new_text = f'Правильный ответ под номером: {new_letter}'
            if p.runs:
                for i, run in enumerate(p.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
            else:
                p.text = new_text
            fixed += 1
    
    doc.save(path)
    print(f'{prefix}: исправлено {fixed} ответов')
    
    # Verify
    doc2 = Document(path)
    ans = [p.text.strip() for p in doc2.paragraphs if 'Правильный ответ' in p.text]
    bad = [a for a in ans if not re.match(r'^Правильный ответ под номером: [ABCD]$', a)]
    print(f'  Проверка: {len(ans)} ответов, {len(bad)} плохих')

print('Done!')
