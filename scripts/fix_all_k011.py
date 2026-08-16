#!/usr/bin/env python3
"""
Fix К0.11 stop-words in all 9 ФОС files.
Only changes question TEXT, not options or answers → D-balance preserved.
"""
import re, os, copy, random
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = 'Оператор автоматических и полуавтоматических станков и линий станков'
PROF = 'Оператор автоматических и полуавтоматических станков и линий станков'
random.seed(42)

# Exact replacements: (old_substring, new_substring)
# Apply to question text only
RULES = [
    # 1. наилучшее сочетание
    ('обеспечивает наилучшее сочетание прочности и вязкости после термообработки',
     'образуется при улучшении (закалка + высокий отпуск) и обеспечивает высокую прочность и вязкость'),
    # 2. наилучшую шероховатость
    ('обеспечивает наилучшую шероховатость обработки',
     'обеспечивает минимальную шероховатость обработки'),
    # 3. наиболее опасным
    ('является наиболее опасным при изгибе и кручении',
     'определяется как опасное при совместном действии изгиба и кручения'),
    # 4. наиболее рационален
    ('наиболее рационален для деталей типа вал в массовом производстве',
     'рекомендуется для деталей типа вал при массовом производстве'),
    # 5. наиболее производительным
    ('является наиболее производительным для получения плоских поверхностей',
     'выполняется многозубым инструментом при вращательном главном движении для получения плоских поверхностей'),
    # 6. наиболее точным
    ('является наиболее точным из стандартных',
     'имеет наименьшее значение допускаемой основной погрешности из перечисленных'),
    # 7. наиболее широко
    ('наиболее широко применяется для привода станков',
     'является основным типом, применяемым для привода станков'),
    # 8. наиболее влияет
    ('Какой фактор наиболее влияет на стойкость шлифовального круга?',
     'Какие факторы определяют стойкость шлифовального круга?'),
    # 9. около
    ('содержание углерода около 0,45%',
     'среднее содержание углерода 0,45%'),
    # 10. примерно 0,06 мм
    ('допуск примерно 0,06 мм для размера 50 мм',
     'допуск для номинального размера 50 мм составляет 0,062 мм'),
    # 11. обычно
    ('обычно обеспечивается при чистовом шлифовании',
     'обеспечивается при чистовом шлифовании'),
    # 12. порядка 0,035 мм
    ('допуск для размера 100 мм порядка 0,035 мм',
     'допуск для номинального размера 100 мм составляет 0,035 мм'),
    # 13. как часто (ПТЭЭП)
    ('как часто проводится визуальный осмотр электрооборудования станка',
     'с какой периодичностью проводится визуальный осмотр электрооборудования станка'),
]

def fix_text(q_text):
    for old, new in RULES:
        if old in q_text:
            return q_text.replace(old, new), old[:50]
    return q_text, None

# ============================================================
# PARSERS
# ============================================================
def parse_103_104(path):
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs]
    questions = []
    i = 0
    while i < len(lines):
        m = re.match(r'^(\d+)\.\s+(.+)', lines[i])
        if m:
            q_text = m.group(2)
            options = {}; ans = None
            j = i + 1
            while j < len(lines):
                ln = lines[j]
                opt_m = re.match(r'^([ABCD])\)\s+(.+?);?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2).rstrip(';')
                    j += 1; continue
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m:
                    ans = ans_m.group(1); j += 1; break
                if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                    q_text += ' ' + ln; j += 1; continue
                j += 1
            if len(options) == 4 and ans:
                questions.append({'question': q_text, 'A': options['A'], 'B': options['B'],
                                 'C': options['C'], 'D': options['D'], 'ans': ans})
            i = j
        else:
            i += 1
    return questions

def parse_105(path):
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tickets_raw = []; current = None
    for line in lines:
        if re.match(r'^Билет\s+(\d+)$', line):
            if current is not None: tickets_raw.append(current)
            current = []; continue
        if current is None: continue
        if line: current.append(line)
    if current is not None: tickets_raw.append(current)

    all_q = []
    for t_lines in tickets_raw:
        i = 0
        while i < len(t_lines):
            m = re.match(r'^(\d+)\.\s+(.+)', t_lines[i])
            if m:
                q_text = m.group(2)
                options = {}; ans = None
                j = i + 1
                while j < len(t_lines):
                    ln = t_lines[j]
                    opt_m = re.match(r'^([ABCD])\)\s+(.+?)\s*;?$', ln)
                    if opt_m:
                        options[opt_m.group(1)] = opt_m.group(2).rstrip(';').strip()
                        j += 1; continue
                    ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                    if ans_m:
                        ans = ans_m.group(1); j += 1; break
                    if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                        q_text += ' ' + ln; j += 1; continue
                    j += 1
                if len(options) == 4 and ans:
                    all_q.append({'question': q_text, 'A': options['A'], 'B': options['B'],
                              'C': options['C'], 'D': options['D'], 'ans': ans})
                i = j
            else:
                i += 1
    return all_q

# ============================================================
# GENERATORS
# ============================================================
def gen_103_104(questions, path, title_sub):
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(14)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(1.5)
    for text, bold, sz in [('Фонд оценочных средств', True, 16), (title_sub, True, 14),
        ('', False, 14), ('по основной программе профессионального обучения', False, 14)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold; r.font.name = 'Times New Roman'; r.font.size = Pt(sz)
    doc.add_paragraph()
    for qi, q in enumerate(questions, 1):
        p = doc.add_paragraph(); r = p.add_run(f'{qi}. {q["question"]}')
        r.font.name = 'Times New Roman'; r.font.size = Pt(14)
        for letter in 'ABCD':
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(f'{letter}) {q[letter]};'); r.font.name = 'Times New Roman'; r.font.size = Pt(14)
        p = doc.add_paragraph(); r = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        r.font.name = 'Times New Roman'; r.font.size = Pt(14)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    doc.save(path)

def gen_105(questions, path, lt):
    TICKETS = 15
    by_letter = {l: [] for l in 'ABCD'}
    for q in questions: by_letter[q['ans']].append(q)
    for l in 'ABCD': random.shuffle(by_letter[l])
    tickets = [[] for _ in range(TICKETS)]
    for l in 'ABCD':
        for i, q in enumerate(by_letter[l]): tickets[i % TICKETS].append(q)
    for t in tickets: random.shuffle(t)

    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(14)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(1.5)
    lt_map = {'П': '(профессиональной подготовки)', 'ПП': '(переподготовки)', 'КПК': '(повышения квалификации)'}
    for text, bold, sz in [('Фонд оценочных средств', True, 16), ('итоговой аттестации', True, 16),
        ('', False, 14), ('по основной программе профессионального обучения', False, 14),
        (lt_map.get(lt, ''), False, 14), ('', False, 14), ('по профессии', False, 14),
        (f'\u00AB{PROF}\u00BB', True, 14), ('', False, 14),
        ('Тестовые задания с выбором одного правильного ответа', False, 14), ('(20 вопросов)', False, 14)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold; r.font.name = 'Times New Roman'; r.font.size = Pt(sz)
    doc.add_page_break()
    for ti, ticket in enumerate(tickets):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'Билет {ti+1}'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(16)
        doc.add_paragraph()
        for qi, q in enumerate(ticket, 1):
            p = doc.add_paragraph(); r = p.add_run(f'{qi}. {q["question"]}')
            r.font.name = 'Times New Roman'; r.font.size = Pt(14)
            for letter in 'ABCD':
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
                r = p.add_run(f'{letter}) {q[letter]};'); r.font.name = 'Times New Roman'; r.font.size = Pt(14)
            p = doc.add_paragraph(); r = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
            r.font.name = 'Times New Roman'; r.font.size = Pt(14)
        if ti < len(tickets) - 1: doc.add_page_break()
    doc.save(path)
    return tickets

# ============================================================
# MAIN
# ============================================================
sf_map = {
    'профессиональной подготовки': ('ПО_П_', 'П'),
    'профессиональной переподготовки': ('ПО_ПП_', 'ПП'),
    'повышения квалификации': ('ПО_КПК_', 'КПК'),
}

total_fixed = 0
for sf_name, (prefix, lt) in sorted(sf_map.items()):
    # Find actual folder containing sf_name
    actual = [d for d in os.listdir(BASE) if sf_name in d and os.path.isdir(os.path.join(BASE, d))]
    if not actual: continue
    folder = os.path.join(BASE, actual[0])
    print(f'\n{"="*60}\n{lt}: {sf_name}\n{"="*60}')

    for ftype in ['103', '104', '105']:
        files = [f for f in os.listdir(folder) if f.startswith(f'{ftype}.')]
        if not files: continue
        path = os.path.join(folder, files[0])

        if ftype in ('103', '104'):
            qs = parse_103_104(path)
            fixed = 0
            for q in qs:
                new_text, rule = fix_text(q['question'])
                if new_text != q['question']:
                    print(f'  {ftype}: "{rule}..."')
                    print(f'    → {new_text[:90]}')
                    q['question'] = new_text; fixed += 1
            if fixed:
                title_sub = files[0].split(prefix)[1].split('.docx')[0].strip()
                gen_103_104(qs, path, title_sub)
                print(f'  {ftype}: {fixed} fixes, saved')
                total_fixed += fixed
        else:
            qs = parse_105(path)
            fixed = 0
            for q in qs:
                new_text, rule = fix_text(q['question'])
                if new_text != q['question']:
                    print(f'  105: "{rule}..."')
                    print(f'    → {new_text[:90]}')
                    q['question'] = new_text; fixed += 1
            if fixed:
                tickets = gen_105(qs, path, lt)
                d = Counter(q['ans'] for q in qs)
                print(f'  105: {fixed} fixes, D-balance: {dict(sorted(d.items()))}')
                total_fixed += fixed

print(f'\n{"="*60}\nИТОГО: {total_fixed} исправлений по К0.11\n{"="*60}')
