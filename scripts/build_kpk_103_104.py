#!/usr/bin/env python3
"""
Convert КПК 03 -> 103 and 04 -> 104.
- а)б)в)г) -> A)B)C)D)
- Table answer keys -> inline 'Правильный ответ под номером: X'
- Fix D-balance to 20-30% per option
- Remove tables
- 20 questions each
"""
import re, os, copy, random
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/повышения квалификации 15474 Оператор станков"

random.seed(42)

def parse_old_format(path):
    """Parse old 03/04 format: all options in same paragraph (separated by \n), answers in table."""
    doc = Document(path)
    
    # Get answer keys from table (second table, skip header row)
    answer_table = doc.tables[1] if len(doc.tables) >= 2 else None
    answer_map = {}
    if answer_table:
        for row in answer_table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0].isdigit():
                q_num = int(cells[0])
                ans_letter = cells[1].strip().lower()
                cyr_to_lat = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
                answer_map[q_num] = cyr_to_lat.get(ans_letter, ans_letter.upper())
    
    # Parse questions from paragraphs (each question is ONE paragraph with \n inside)
    questions = []
    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'^(\d+)\.\s+(.+)', text, re.DOTALL)
        if not m:
            continue
        q_num = int(m.group(1))
        full_text = m.group(2)  # everything after "N. "
        
        # Split by \n to find question and options
        lines = full_text.split('\n')
        # First line (or up to first \n) is the question text
        q_text = lines[0].strip()
        
        # Remaining lines are options
        options = {}
        for line in lines[1:]:
            line = line.strip()
            opt_m = re.match(r'^([а-гА-Г])\)\s*(.+)', line)
            if opt_m:
                letter_cyr = opt_m.group(1).lower()
                letter_map = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
                options[letter_map[letter_cyr]] = opt_m.group(2).strip().rstrip(' ;.')
        
        if len(options) == 4 and q_num in answer_map:
            questions.append({
                'num': q_num,
                'question': q_text,
                'A': options.get('A', ''),
                'B': options.get('B', ''),
                'C': options.get('C', ''),
                'D': options.get('D', ''),
                'ans': answer_map[q_num]
            })
        elif len(options) != 4:
            print(f'  WARNING Q{q_num}: {len(options)} options: {list(options.keys())}')
        elif q_num not in answer_map:
            print(f'  WARNING Q{q_num}: no answer in table')
    
    return questions

def swap_answer(q, target_letter):
    """Swap answer to target_letter by swapping option text."""
    current = q['ans']
    if current == target_letter:
        return q
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    ci = letters.index(current)
    ti = letters.index(target_letter)
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    q['A'] = opts['A']; q['B'] = opts['B']; q['C'] = opts['C']; q['D'] = opts['D']
    q['ans'] = target_letter
    return q

def fix_d_balance(questions):
    """Fix D-balance: each option 20-30% of 20 = 4-6 answers."""
    target = 5  # 25% of 20
    qs = copy.deepcopy(questions)
    
    MAX_ROUNDS = 200
    for round_num in range(MAX_ROUNDS):
        counts = Counter(q['ans'] for q in qs)
        over = {l: counts[l] - target for l in 'ABCD' if counts[l] > target}
        under = {l: target - counts[l] for l in 'ABCD' if counts[l] < target}
        if not over and not under:
            print(f'  D-balance achieved in round {round_num + 1}: {dict(sorted(counts.items()))}')
            break
        over_letter = next(iter(over))
        under_letter = next(iter(under))
        indices = list(range(len(qs)))
        random.shuffle(indices)
        for i in indices:
            if qs[i]['ans'] == over_letter:
                qs[i] = swap_answer(qs[i], under_letter)
                break
    
    return qs

def verify_answers(questions):
    """Verify each question's answer matches its options."""
    errors = 0
    for q in questions:
        opts = [q['A'], q['B'], q['C'], q['D']]
        if len(opts) != len(set(opts)):
            print(f'  ERROR Q{q["num"]}: duplicate options!')
            errors += 1
        ans_text = q[q['ans']]
        if not ans_text:
            print(f'  ERROR Q{q["num"]}: answer {q["ans"]} has empty text!')
            errors += 1
    return errors

def create_103_104_docx(questions, output_path, title_subtitle):
    """Create 103/104 docx with inline answers, no tables."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    
    # Title block
    title_lines = [
        ('Фонд оценочных средств', True, 16),
        (title_subtitle, True, 14),
        ('', False, 14),
        ('по основной программе профессионального обучения', False, 14),
        ('(повышения квалификации)', False, 14),
        ('', False, 14),
        ('по профессии', False, 14),
        ('\u00ABОператор автоматических и полуавтоматических станков и линий станков\u00BB', True, 14),
        ('', False, 14),
        ('Тестовые задания с выбором одного правильного ответа (20 вопросов)', False, 14),
    ]
    
    for text, bold, size in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    
    doc.add_paragraph()
    
    for q_idx, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{q_idx}. {q["question"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        for letter in 'ABCD':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'{letter}) {q[letter]};')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        run = p.add_run(f'Правильный ответ под номером: {q["ans"]}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    
    doc.save(output_path)
    print(f'  Saved: {output_path}')

# ============================================================
# MAIN
# ============================================================

for src_num, dst_num, suffix, title_sub in [
    ('03', '103', 'ОП.0.0', 'промежуточной аттестации по общепрофессиональному циклу (ОП.0.0)'),
    ('04', '104', 'МДК01.01', 'промежуточной аттестации (МДК01.01) по профессиональному циклу'),
]:
    print(f'\n{"="*60}')
    print(f'Converting {src_num} -> {dst_num} ({suffix})')
    print(f'{"="*60}')
    
    src_file = [f for f in os.listdir(FOLDER) if f.startswith(src_num + '.')][0]
    src_path = os.path.join(FOLDER, src_file)
    
    questions = parse_old_format(src_path)
    print(f'  Parsed {len(questions)} questions')
    
    if len(questions) != 20:
        print(f'  ERROR: Expected 20, got {len(questions)}!')
        continue
    
    counts_before = Counter(q['ans'] for q in questions)
    print(f'  Before D-balance: {dict(sorted(counts_before.items()))}')
    for l in 'ABCD':
        print(f'    {l}: {counts_before.get(l, 0)}/20 ({counts_before.get(l, 0)/20*100:.0f}%)')
    
    questions = fix_d_balance(questions)
    
    errors = verify_answers(questions)
    if errors:
        print(f'  ERROR: {errors} verification errors!')
        continue
    else:
        print(f'  All answers verified.')
    
    counts_after = Counter(q['ans'] for q in questions)
    print(f'  After D-balance: {dict(sorted(counts_after.items()))}')
    
    dst_file = f'{dst_num}. ПО_КПК_ФОС_Оператор станков_4_разр {suffix}.docx'
    dst_path = os.path.join(FOLDER, dst_file)
    create_103_104_docx(questions, dst_path, title_sub)
    
    # Final verify from saved file
    doc_check = Document(dst_path)
    text_check = '\n'.join(p.text for p in doc_check.paragraphs)
    ans_check = re.findall(r'Правильный ответ под номером:\s*([ABCD])', text_check)
    table_check = sum(1 for el in doc_check.element.body if el.tag.endswith('tbl'))
    lower_check = sum(1 for p in doc_check.paragraphs if re.match(r'^[а-г]\)', p.text.strip()))
    print(f'  Verification:')
    print(f'    Answers: {len(ans_check)}')
    print(f'    Tables: {table_check}')
    print(f'    Lowercase: {lower_check}')
    counts_v = Counter(ans_check)
    print(f'    D-balance: {dict(sorted(counts_v.items()))}')
    print(f'  {"OK" if len(ans_check) == 20 and table_check == 0 and lower_check == 0 else "ISSUE!"}')

print('\nDone!')
