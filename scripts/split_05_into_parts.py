#!/usr/bin/env python3
"""
Разбивает оригинальный файл 05 (15 билетов x 20 вопросов)
на 5 автономных файлов 05-1 ... 05-5 (по 3 билета каждый).

Вход: audits/05_parsed_data.json
Выход: 05-N.docx в той же папке, что и оригинальный 05
"""
import json
import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === ПУТИ ===
BASE = "/home/z/my-project/vol"
AUDITS_DIR = f"{BASE}/audits"
OUTPUT_DIR = (
    f"{BASE}/Оператор автоматических и полуавтоматических станков и линий станков"
    f"/профессиональной подготовки 15474 Оператор станков"
)

# Пары (часть_номер, [билеты])
PARTS = [
    (1, ["1", "2", "3"]),
    (2, ["4", "5", "6"]),
    (3, ["7", "8", "9"]),
    (4, ["10", "11", "12"]),
    (5, ["13", "14", "15"]),
]

def load_data():
    with open(f"{AUDITS_DIR}/05_parsed_data.json", encoding="utf-8") as f:
        return json.load(f)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    # Убрать отступы ячейки
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Вертикальное выравнивание по центру
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

def add_header_table(doc):
    """Таблица утверждения ПРИНЯТО / УТВЕРЖДАЮ"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Установить ширину таблицы
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    # Объединить ячейки
    cell_a = table.cell(0, 0)
    cell_b = table.cell(0, 1)
    cell_c = table.cell(0, 2)
    cell_a.merge(cell_b)
    merged = cell_a
    set_cell_text(merged, (
        "ПРИНЯТО и рекомендовано к утверждению решением "
        "педагогического совета"
    ), bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_c, (
        "УТВЕРЖДАЮ:\nДиректор ООО УЦ 'ПрофСтрой'\n___________"
    ), bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()  # отступ
def add_title(doc, part_num):
    """Заголовок документа"""
    lines = [
        "Фонд оценочных средств",
        "итоговой аттестации (часть " + str(part_num) + " из 5)",
        "по основной программе профессионального обучения",
        "(профессиональной подготовки)",
        "по профессии 'Оператор автоматических и полуавтоматических станков и линий станков'",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()  # отступ

def add_subtitle(doc, part_num):
    """Подзаголовок"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = f"Тестовые задания с выбором одного правильного ответа (20 вопросов, часть {part_num} из 5)"
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()

def add_key_table(doc, pre_section):
    """Таблица 2: все 20 уникальных вопросов с ключами"""
    # Заголовок таблицы
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Таблица 2. Ключи к тестовым заданиям")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=21, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Заголовок
    headers = ["No", "Текст вопроса", "Ответ"]
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, size=11)
        set_cell_shading(table.cell(0, j), "D9E2F3")

    for i, q in enumerate(pre_section):
        row_idx = i + 1
        set_cell_text(table.cell(row_idx, 0), str(q["num"]), size=10)
        # Текст вопроса
        cell_q = table.cell(row_idx, 1)
        cell_q.text = ""
        p = cell_q.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(q["text"])
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        # Ключ
        set_cell_text(table.cell(row_idx, 2), q["key"], size=11)

    doc.add_paragraph()  # отступ

def add_question_text(doc, q_num, q_text):
    """Добавляет текст вопроса (без ключа)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{q_num}. {q_text}")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def add_variant(doc, letter, text, is_last=False):
    """Добавляет вариант ответа"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    # Форматирование variants - List Paragraph стиль
    p.style = doc.styles['List Paragraph']
    separator = "." if is_last else ";"
    run = p.add_run(f"{letter}) {text}{separator}")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def add_ticket(doc, ticket_num, questions):
    """Добавляет один билет (20 вопросов без ключей)"""
    # Заголовок билета
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Билет No {ticket_num}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()

    for i, q in enumerate(questions):
        add_question_text(doc, i + 1, q["text"])
        for j, v in enumerate(q["variants"]):
            is_last = (j == len(q["variants"]) - 1)
            add_variant(doc, v["letter"], v["text"], is_last=is_last)
        # НЕТ строки "Правильный ответ" — ключи только в матрице
        doc.add_paragraph()  # отступ между вопросами

def add_matrix(doc, ticket_nums, matrix_data):
    """Матрица ответов (N билетов x 20 вопросов)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Матрица ответов")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()

    n_tickets = len(ticket_nums)
    table = doc.add_table(rows=n_tickets + 1, cols=21)  # +1 для заголовка
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Заголовок: "Билет" | 1 | 2 | ... | 20
    set_cell_text(table.cell(0, 0), "Билет", bold=True, size=10)
    set_cell_shading(table.cell(0, 0), "D9E2F3")
    for q in range(1, 21):
        set_cell_text(table.cell(0, q), str(q), bold=True, size=10)
        set_cell_shading(table.cell(0, q), "D9E2F3")

    # Данные
    for i, tn in enumerate(ticket_nums):
        row_idx = i + 1
        set_cell_text(table.cell(row_idx, 0), str(tn), bold=True, size=10)
        row_data = matrix_data[str(tn)]
        for q in range(20):
            set_cell_text(table.cell(row_idx, q + 1), row_data[q], size=10)

def create_part_doc(part_num, ticket_nums, data):
    """Создаёт один файл 05-N.docx"""
    doc = Document()

    # Установить поля
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # 1. Шапка
    add_header_table(doc)

    # 2. Заголовок
    add_title(doc, part_num)

    # 3. Подзаголовок
    add_subtitle(doc, part_num)

    # 4. Таблица ключей (все 20 вопросов)
    add_key_table(doc, data["pre_section"])

    # 5. Билеты
    for tn in ticket_nums:
        add_ticket(doc, int(tn), data["tickets"][tn])

    # 6. Матрица ответов
    add_matrix(doc, [int(t) for t in ticket_nums], data["matrix"])

    # Сохранить
    out_path = f"{OUTPUT_DIR}/05-{part_num}. ПО_П_ФОС_Оператор станков_2-3_разр.docx"
    doc.save(out_path)
    print(f"  Created: 05-{part_num} (билеты {ticket_nums[0]}-{ticket_nums[-1]}) -> {out_path}")
    return out_path

def main():
    print("Loading parsed data...")
    data = load_data()
    print(f"  pre_section: {len(data['pre_section'])} questions")
    print(f"  tickets: {len(data['tickets'])} tickets")
    print(f"  matrix: {len(data['matrix'])} rows")

    print("\nGenerating 05-1 ... 05-5:")
    created = []
    for part_num, ticket_nums in PARTS:
        path = create_part_doc(part_num, ticket_nums, data)
        created.append(path)

    print(f"\nDone. Created {len(created)} files.")

if __name__ == "__main__":
    main()
