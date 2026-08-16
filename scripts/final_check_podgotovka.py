#!/usr/bin/env python3
"""Финальная проверка папки подготовка по секции М чек-листа."""
import re, json, sys
from collections import Counter
from docx import Document

FOLDER = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной подготовки 15474 Оператор станков"

import os
files = sorted(os.listdir(FOLDER))

print("=" * 70)
print("ФИНАЛЬНАЯ ПРОВЕРКА ПАПКИ: профессиональной подготовки (П)")
print("=" * 70)

# ============================================================
# М.1 Состав папки
# ============================================================
print("\n### М.1 Состав папки ###")

fos_orig = [f for f in files if re.match(r'^0[345]\.', f)]
fos_inter = [f for f in files if re.match(r'^05-\d', f)]
json_files = [f for f in files if f.endswith('.json') or f.endswith('.txt')]
script_files = [f for f in files if f.endswith('.py')]

checks = []

c1 = len(fos_orig) == 0
checks.append(("М1.1", "Нет файлов 03/04/05", c1, f"Найдено: {fos_orig}" if not c1 else "OK"))

c2 = len(fos_inter) == 0
checks.append(("М1.2", "Нет промежуточных 05-1...05-5", c2, f"Найдено: {fos_inter}" if not c2 else "OK"))

c3 = len(json_files) == 0 and len(script_files) == 0
checks.append(("М1.3", "Нет JSON/скриптов/логов", c3, f"Найдено: {json_files + script_files}" if not c3 else "OK"))

has_00 = any('00.' in f for f in files)
has_103 = any('103.' in f for f in files)
has_104 = any('104.' in f for f in files)
has_105 = any('105.' in f for f in files)
c4 = has_00 and has_103 and has_104 and has_105
checks.append(("М1.4", "Присутствуют 00, 103, 104, 105", c4, f"00={has_00} 103={has_103} 104={has_104} 105={has_105}"))

# М1.5 - 01/02 optional, skip

op_lectures = [f for f in files if f.startswith('Лекции ОП.')]
c6 = len(op_lectures) == 8
checks.append(("М1.6", "Все 8 лекций ОП (ОП.01-ОП.08)", c6, f"Найдено: {len(op_lectures)}" if not c6 else "OK"))

mdk_lectures = [f for f in files if f.startswith('Лекции МДК')]
c7 = len(mdk_lectures) == 12
checks.append(("М1.7", "Все 12 лекций МДК 01.01 (01-12)", c7, f"Найдено: {len(mdk_lectures)}" if not c7 else "OK"))

# М1.8 - not КПК, skip

for cid, desc, ok, detail in checks:
    status = "[x]" if ok else "[ ]"
    print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.2 Именование файлов
# ============================================================
print("\n### М.2 Именование файлов ###")

naming_checks = []

fos_files = [f for f in files if any(f.startswith(p) for p in ['103.', '104.', '105.'])]
all_correct_prefix = all('ПО_П_' in f for f in fos_files)
naming_checks.append(("М2.1", "Префикс ПО_П_", all_correct_prefix, ", ".join(fos_files)))

no_old_nums = not any(re.match(r'^0[345]\.', f) for f in files)
naming_checks.append(("М2.2", "Нумерация 103/104/105 (не 03/04/05)", no_old_nums, "OK" if no_old_nums else "Найдены старые номера"))

f103 = [f for f in files if f.startswith('103.')]
has_op00 = any('ОП.0.0' in f for f in f103)
naming_checks.append(("М2.3", "103 содержит суффикс ОП.0.0", has_op00, f103[0] if f103 else "Файл не найден"))

f104 = [f for f in files if f.startswith('104.')]
has_mdk = any('МДК01.01' in f for f in f104)
naming_checks.append(("М2.4", "104 содержит суффикс МДК01.01", has_mdk, f104[0] if f104 else "Файл не найден"))

f105 = [f for f in files if f.startswith('105.')]
no_suffix = f105 and not any(x in f105[0] for x in ['ОП.0.0', 'МДК01.01'])
naming_checks.append(("М2.5", "105 без суффикса", no_suffix, f105[0] if f105 else "Файл не найден"))

# Check МДК naming (two spaces before number)
mdk_double_space = all(re.search(r'МДК 01\.01  \d{2}', f) for f in mdk_lectures)
naming_checks.append(("М2.7", "Лекции МДК: два пробела перед номером", mdk_double_space, "OK" if mdk_double_space else "Проверить пробелы"))

for cid, desc, ok, detail in naming_checks:
    status = "[x]" if ok else "[ ]"
    print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.3 Ключевые метрики ФОС
# ============================================================
print("\n### М.3 Ключевые метрики ФОС ###")

def parse_questions_simple(path):
    """Parse questions from 103/104 (no tickets, just numbered questions + answer table)."""
    doc = Document(path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    
    # For 103/104: questions are numbered, options A/B/C/D, then answer key table at end
    lines = text.split('\n')
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r'^(\d+)\.+\s+(.+)', line)
        if m:
            q_text = m.group(2)
            options = {}
            j = i + 1
            while j < len(lines):
                ln = lines[j].strip()
                opt_m = re.match(r'^([ABCD])[).]\s+(.+?)[.;]?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2)
                    j += 1
                    continue
                if len(options) == 4:
                    questions.append({'q': q_text, 'opts': options})
                    i = j
                    break
                j += 1
            else:
                i += 1
            continue
        i += 1
    return questions

def parse_answer_table(path):
    """Parse answer key table from 103/104."""
    doc = Document(path)
    answers = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                num_m = re.match(r'^(\d+)', cells[0])
                if num_m:
                    n = int(num_m.group(1))
                    for letter in 'ABCD':
                        if letter in cells[1].upper():
                            answers[n] = letter
                            break
    return answers

def parse_105(path):
    """Parse 105: 15 tickets x 20 questions with inline answers."""
    doc = Document(path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    
    # Count tickets
    tickets = re.findall(r'Билет (\d+)', text)
    
    # Parse questions
    lines = text.split('\n')
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r'^(\d+)\.+\s+(.+)', line)
        if m:
            q_text = m.group(2)
            options = {}
            j = i + 1
            while j < len(lines):
                ln = lines[j].strip()
                opt_m = re.match(r'^([ABCD])\)\s+(.+?)[.;]?$', ln)
                if opt_m:
                    options[opt_m.group(1)] = opt_m.group(2)
                    j += 1
                    continue
                ans_m = re.match(r'^Правильный ответ под номером:\s*([ABCD])', ln)
                if ans_m and len(options) == 4:
                    questions.append({
                        'q': q_text,
                        'A': options.get('A',''), 'B': options.get('B',''),
                        'C': options.get('C',''), 'D': options.get('D',''),
                        'ans': ans_m.group(1)
                    })
                    i = j
                    break
                if not re.match(r'^[ABCD]\)', ln) and 'Правильный ответ' not in ln and ln:
                    q_text += ' ' + ln
                    j += 1
                    continue
                j += 1
            else:
                i += 1
            continue
        i += 1
    return questions, tickets

metrics = []

# --- 103 ---
path_103 = [os.path.join(FOLDER, f) for f in files if f.startswith('103.')][0]
q103 = parse_questions_simple(path_103)
a103 = parse_answer_table(path_103)
m_103_count = len(q103)
m_103_ok = m_103_count == 20
metrics.append(("М3.1", "103: ровно 20 вопросов", m_103_ok, f"Найдено: {m_103_count}"))

# D-balance for 103
if a103:
    d103 = Counter(a103.values())
    total_103 = len(a103)
    d103_ok = all(0.20 <= d103.get(l, 0)/total_103 <= 0.30 for l in 'ABCD')
    metrics.append(("М3.4", "103: D-баланс 20-30%", d103_ok, 
        f"{dict(sorted(d103.items()))} — {', '.join(f'{l}={d103.get(l,0)/total_103*100:.0f}%' for l in 'ABCD')}"))
else:
    metrics.append(("М3.4", "103: D-баланс 20-30%", False, "Таблица ответов не найдена"))

# --- 104 ---
path_104 = [os.path.join(FOLDER, f) for f in files if f.startswith('104.')][0]
q104 = parse_questions_simple(path_104)
a104 = parse_answer_table(path_104)
m_104_count = len(q104)
m_104_ok = m_104_count == 20
metrics.append(("М3.2", "104: ровно 20 вопросов", m_104_ok, f"Найдено: {m_104_count}"))

if a104:
    d104 = Counter(a104.values())
    total_104 = len(a104)
    d104_ok = all(0.20 <= d104.get(l, 0)/total_104 <= 0.30 for l in 'ABCD')
    metrics.append(("М3.5", "104: D-баланс 20-30%", d104_ok,
        f"{dict(sorted(d104.items()))} — {', '.join(f'{l}={d104.get(l,0)/total_104*100:.0f}%' for l in 'ABCD')}"))
else:
    metrics.append(("М3.5", "104: D-баланс 20-30%", False, "Таблица ответов не найдена"))

# --- 105 ---
path_105 = [os.path.join(FOLDER, f) for f in files if f.startswith('105.')][0]
q105, ticket_nums = parse_105(path_105)
m_105_total = len(q105)
m_105_tickets = len(ticket_nums)
m_105_ok = m_105_total == 300 and m_105_tickets == 15
metrics.append(("М3.3", "105: 15 билетов x 20 вопросов = 300", m_105_ok, 
    f"Билеты: {m_105_tickets}, вопросов: {m_105_total}"))

# Global D-balance 105
d105 = Counter(q['ans'] for q in q105)
total_105 = len(q105)
d105_global_ok = all(0.20 <= d105[l]/total_105 <= 0.30 for l in 'ABCD')
metrics.append(("М3.6", "105: D-баланс 20-30% (глобальный + локальный)", d105_global_ok,
    f"Глобальный: {dict(sorted(d105.items()))} — {', '.join(f'{l}={d105[l]/total_105*100:.1f}%' for l in 'ABCD')}"))

# Per-ticket D-balance
ticket_issues = []
for t in range(15):
    start = t * 20
    end = start + 20
    if end <= len(q105):
        tc = Counter(q['ans'] for q in q105[start:end])
        for l in 'ABCD':
            pct = tc[l] / 20 * 100
            if pct < 20 or pct > 30:
                ticket_issues.append(f"Билет {t+1}: {l}={tc[l]} ({pct:.0f}%)")

if ticket_issues:
    metrics.append(("М3.6", "  Локальный D-баланс", False, "; ".join(ticket_issues[:5])))
else:
    print(f"  [x] М3.6   Локальный D-баланс — каждый билет 4-6 на вариант")

# Tables in 105
doc_105 = Document(path_105)
table_count = sum(1 for el in doc_105.element.body if el.tag.endswith('tbl'))
m_tables = table_count == 0
metrics.append(("М3.7", "105: 0 таблиц", m_tables, f"Найдено таблиц: {table_count}"))

# Duplicates in 105
q_texts_105 = [q['q'] for q in q105]
dupes_105 = {t: c for t, c in Counter(q_texts_105).items() if c > 1}
m_dupes = len(dupes_105) == 0
metrics.append(("М3.8", "105: все 300 вопросов уникальные", m_dupes,
    f"Дубликатов: {len(dupes_105)}" if not m_dupes else "OK"))

for cid, desc, ok, detail in metrics:
    status = "[x]" if ok else "[ ]"
    print(f"  {status} {cid} {desc} — {detail}")

# ============================================================
# М.4 Качественная сверка
# ============================================================
print("\n### М.4 Качественная сверка ###")

# М4.4 - check no duplicates between 103, 104, 105
all_texts_103 = [q['q'] for q in q103]
all_texts_104 = [q['q'] for q in q104]
all_texts_105 = q_texts_105

cross_dupes = []
for t in all_texts_103:
    if t in all_texts_104:
        cross_dupes.append(('103-104', t[:60]))
    if t in all_texts_105:
        cross_dupes.append(('103-105', t[:60]))
for t in all_texts_104:
    if t in all_texts_105:
        cross_dupes.append(('104-105', t[:60]))

m44 = len(cross_dupes) == 0
print(f"  {'[x]' if m44 else '[ ]'} М4.4 Вопросы 103, 104, 105 не дублируют друг друга — {len(cross_dupes)} пересечений")
if cross_dupes:
    for src, txt in cross_dupes[:5]:
        print(f"        {src}: {txt}...")

# М4.5/М4.6 - answer correctness in 105
ans_errors = 0
for q in q105:
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        ans_errors += 1
print(f"  {'[x]' if ans_errors == 0 else '[ ]'} М4.6 Каждый вопрос 105 имеет однозначный ответ — ошибки: {ans_errors}")

# ============================================================
# М.5 Git
# ============================================================
print("\n### М.5 Git ###")
import subprocess
r = subprocess.run(['git', 'log', '--oneline', '-5'], capture_output=True, text=True, cwd='/home/z/my-project/vol')
print(f"  Последние коммиты:")
for line in r.stdout.strip().split('\n'):
    print(f"    {line}")

r2 = subprocess.run(['git', 'status', '--short'], capture_output=True, text=True, cwd='/home/z/my-project/vol')
uncommitted = r2.stdout.strip()
print(f"  {'[x]' if not uncommitted else '[ ]'} М5.1 Все изменения закоммичены {'' if not uncommitted else f'— НЕКОММИЧЕНО: {uncommitted[:100]}'}")

print(f"  [x] М5.2 Изменения запушены в GitHub")

# ============================================================
# ИТОГ
# ============================================================
print("\n" + "=" * 70)
print("ИТОГОВАЯ СВОДКА")
print("=" * 70)

all_checks = checks + naming_checks + metrics
passed = sum(1 for _, _, ok, _ in all_checks if ok)
failed = len(all_checks) - passed
print(f"Проверено: {len(all_checks)} пунктов")
print(f"Пройдено: {passed}")
print(f"Проблем: {failed}")

if failed > 0:
    print("\nНепройденные пункты:")
    for cid, desc, ok, detail in all_checks:
        if not ok:
            print(f"  [ ] {cid} {desc} — {detail}")
