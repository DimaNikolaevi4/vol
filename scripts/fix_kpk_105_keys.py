#!/usr/bin/env python3
"""
Исправление ключей в 105. ПО_КПК_ФОС_Оператор станков_4_разр.docx

10 ключевых исправлений по результатам ручной проверки:
  Б1в1:   C→B  (ГОСТ 25346: 100мм/0,035мм = IT7, не IT8)
  Б1в14:  C→B  (чистовое точение = Ra 6,3-3,2, не Ra 0,4)
  Б3в8:   C→B  (число проходов = шаг+материал+точность, не длина)
  Б8в8:   C→B  (конусность = конусный калибр по краске, не визуальный осмотр)
  Б8в14:  C→B  (производительность = фрезерование торцевой фрезой, не сверление)
  Б9в14:  C→B  (стойкость круга = зернистость+твёрдость+режим правки, не масса)
  Б11в18: C→B  (внутренние дефекты = УЗ дефектоскопия, не микрометр)
  Б12в2:  A→C  (твёрдость инструмента = закалка с отпуском, не нормализация)
  Б13в20: D→C  (соосность = индикатор на оправке, не калибр-пробка)
  Б14в13: C→B  (износ в автомате = контроль размера + нормативы, не прослушивание)

D-баланс до: A=75 B=75 C=75 D=75 (идеальный)
D-баланс после исправлений: A=74 B=83 C=69 D=74 (нужна ребалансировка)
Ребалансировка: 1× B→A, 6× B→C, 1× B→D  (swap option text)
D-баланс финал: A=75 B=75 C=75 D=75
"""

import re
import copy
import json
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"
FOLDER = "повышения квалификации 15474 Оператор станков"
FILEPATH = f"{BASE}/{FOLDER}/105. ПО_КПК_ФОС_Оператор станков_4_разр.docx"

def parse_105(filepath):
    """Parse 105 file into structured data."""
    doc = Document(filepath)
    
    # Extract header paragraphs (before first ticket)
    header_paras = []
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = {}
    current_answer = None
    in_questions = False
    
    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()
        
        if not txt_stripped:
            header_paras.append(txt)
            continue
        
        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', txt_stripped)
        if ticket_match:
            in_questions = True
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_ticket = int(ticket_match.group(1))
            current_q_text = None
            header_paras.append(txt)
            continue
        
        if not in_questions:
            header_paras.append(txt)
            continue
        
        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', txt_stripped)
        if q_match:
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = {}
            current_answer = None
            continue
        
        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q_text is not None:
            current_options[opt_match.group(1)] = opt_match.group(2)
            continue
        
        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)
    
    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket, 'num': current_q_num,
            'text': current_q_text, 'options': current_options,
            'answer': current_answer
        })
    
    return header_paras, questions


def compute_balance(questions):
    """Compute D-balance globally and per-ticket."""
    global_counts = Counter()
    per_ticket = {}
    for q in questions:
        a = q['answer']
        global_counts[a] += 1
        t = q['ticket']
        if t not in per_ticket:
            per_ticket[t] = Counter()
        per_ticket[t][a] += 1
    return global_counts, per_ticket


def apply_key_fixes(questions):
    """Apply the 10 key corrections."""
    fixes = {
        (1, 1): 'B',   # C→B
        (1, 14): 'B',  # C→B (also fixes contradiction with Q10)
        (3, 8): 'B',   # C→B
        (8, 8): 'B',   # C→B
        (8, 14): 'B',  # C→B
        (9, 14): 'B',  # C→B
        (11, 18): 'B', # C→B
        (12, 2): 'C',  # A→C
        (13, 20): 'C', # D→C
        (14, 13): 'B', # C→B
    }
    
    fixed = 0
    for q in questions:
        key = (q['ticket'], q['num'])
        if key in fixes:
            old = q['answer']
            new = fixes[key]
            if old != new:
                print(f"  Б{q['ticket']}в{q['num']}: {old}→{new}  |  {q['text'][:80]}...")
                q['answer'] = new
                fixed += 1
            else:
                print(f"  Б{q['ticket']}в{q['num']}: уже {new}, пропускаем")
    
    print(f"\nИсправлено ключей: {fixed}")
    return fixed


def rebalance(questions, target_per_letter=75):
    """Rebalance D-distribution by swapping option texts within questions."""
    global_counts, per_ticket = compute_balance(questions)
    
    print(f"\nБаланс до ребалансировки: A={global_counts['A']} B={global_counts['B']} C={global_counts['C']} D={global_counts['D']}")
    
    # Calculate needed changes
    over = {}
    under = {}
    for letter in 'ABCD':
        diff = global_counts[letter] - target_per_letter
        if diff > 0:
            over[letter] = diff
        elif diff < 0:
            under[letter] = -diff
    
    print(f"Избыток: {dict(over)}")
    print(f"Дефицит: {dict(under)}")
    
    if not over or not under:
        print("Ребалансировка не нужна!")
        return 0
    
    # Fixed questions that we MUST NOT touch (user-corrected)
    fixed_keys = {(1,1),(1,14),(3,8),(8,8),(8,14),(9,14),(11,18),(12,2),(13,20),(14,13)}
    used_indices = set()
    
    swaps = []
    
    # Build candidate pool: over-letter questions not yet used
    def get_candidates(source_letter, target_letter):
        result = []
        for i, q in enumerate(questions):
            if i in used_indices:
                continue
            if (q['ticket'], q['num']) in fixed_keys:
                continue
            if q['answer'] != source_letter:
                continue
            if target_letter not in q['options']:
                continue
            if q['options'][source_letter] == q['options'][target_letter]:
                continue
            # Per-ticket: source must be >4 (can afford to lose one), target must be <6
            t = q['ticket']
            tc = per_ticket[t]
            if tc[source_letter] <= 4 or tc[target_letter] >= 6:
                continue
            result.append((i, q))
        return result
    
    total_swaps = 0
    max_iterations = 50
    iteration = 0
    
    while (over or under) and iteration < max_iterations:
        iteration += 1
        made_swap = False
        
        # Recalculate over/under
        for letter in 'ABCD':
            diff = global_counts[letter] - target_per_letter
            if diff > 0:
                over[letter] = diff
            elif diff <= 0:
                over.pop(letter, None)
        for letter in 'ABCD':
            diff = global_counts[letter] - target_per_letter
            if diff < 0:
                under[letter] = -diff
            elif diff >= 0:
                under.pop(letter, None)
        
        if not over or not under:
            break
        
        # Try each (source, target) pair
        for source_letter in sorted(over.keys(), key=lambda x: -over[x]):
            if made_swap:
                break
            for target_letter in sorted(under.keys(), key=lambda x: -under[x]):
                if made_swap:
                    break
                if over.get(source_letter, 0) <= 0:
                    break
                if under.get(target_letter, 0) <= 0:
                    continue
                
                candidates = get_candidates(source_letter, target_letter)
                if not candidates:
                    continue
                
                # Pick best candidate (one that hurts local balance least)
                best = None
                best_score = 999
                for idx, q in candidates:
                    t = q['ticket']
                    # Score: prefer questions where local balance is closest to 5
                    tc = per_ticket[t]
                    src_after = tc[source_letter] - 1
                    tgt_after = tc[target_letter] + 1
                    # Penalize if result goes out of 4-6
                    score = 0
                    if src_after < 4: score += 100
                    if tgt_after > 6: score += 100
                    # Prefer tickets where both letters are close to 5
                    score += abs(src_after - 5) + abs(tgt_after - 5)
                    if score < best_score:
                        best_score = score
                        best = (idx, q)
                
                if best is None:
                    continue
                
                idx, q = best
                src_text = q['options'][source_letter]
                tgt_text = q['options'][target_letter]
                
                # Perform swap
                q['options'][source_letter] = tgt_text
                q['options'][target_letter] = src_text
                q['answer'] = target_letter
                used_indices.add(idx)
                
                # Update counters
                global_counts[source_letter] -= 1
                global_counts[target_letter] += 1
                over[source_letter] = over.get(source_letter, 0) - 1
                under[target_letter] = under.get(target_letter, 0) - 1
                if over.get(source_letter, 0) <= 0:
                    over.pop(source_letter, None)
                if under.get(target_letter, 0) <= 0:
                    under.pop(target_letter, None)
                
                # Update per-ticket
                t = q['ticket']
                per_ticket[t][source_letter] -= 1
                per_ticket[t][target_letter] += 1
                
                total_swaps += 1
                made_swap = True
                swaps.append(f"  Б{q['ticket']}в{q['num']}: {source_letter}→{target_letter}")
    
    print(f"\nСвопы для ребалансировки ({total_swaps}):")
    for s in swaps:
        print(s)
    
    print(f"\nБаланс после ребалансировки: A={global_counts['A']} B={global_counts['B']} C={global_counts['C']} D={global_counts['D']}")
    
    # Verify per-ticket
    bad_tickets = []
    for t in sorted(per_ticket.keys()):
        for letter in 'ABCD':
            cnt = per_ticket[t][letter]
            if cnt < 4 or cnt > 6:
                bad_tickets.append(f"Билет {t}: {letter}={cnt}")
    if bad_tickets:
        print(f"⚠️ Нарушения локального баланса: {bad_tickets}")
    else:
        print("Локальный баланс ОК (4-6 в каждом билете)")
    
    return total_swaps


def fix_local_balance(questions, fixed_keys):
    """Fix per-ticket balance violations with cross-ticket swaps (net-zero global impact)."""
    from collections import defaultdict
    gc, pt = compute_balance(questions)
    
    print(f"\n[локальная ребалансировка]")
    print(f"  Глобальный баланс перед: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")
    
    def get_violations():
        viols = []
        for t in sorted(pt.keys()):
            for letter in 'ABCD':
                cnt = pt[t][letter]
                if cnt < 4:
                    viols.append((t, letter, 'deficit', cnt))
                elif cnt > 6:
                    viols.append((t, letter, 'excess', cnt))
        return viols
    
    violations = get_violations()
    if not violations:
        print("  Локальный баланс уже ОК")
        return 0
    
    print(f"  Нарушения: {[(f'Б{v[0]}:{v[1]}={v[3]}') for v in violations]}")
    
    swaps_done = 0
    max_iter = 20
    
    for iteration in range(max_iter):
        gc, pt = compute_balance(questions)
        violations = get_violations()
        if not violations:
            break
        
        made_swap = False
        
        # Build excess/deficit maps per ticket
        for t1, l1, kind1, cnt1 in violations:
            if made_swap: break
            if kind1 != 'excess': continue
            
            # Find a deficit letter in the SAME ticket
            for t1b, l2, kind2, cnt2 in violations:
                if made_swap: break
                if t1b != t1: continue
                if kind2 != 'deficit': continue
                
                # T1 has excess of l1 and deficit of l2
                # Need: T2 with l1 can increase (l1<=5) and l2 can decrease (l2>=5)
                for t2 in range(1, 16):
                    if t2 == t1: continue
                    tc2 = pt[t2]
                    if tc2[l1] >= 6 or tc2[l2] <= 4:
                        continue
                    
                    # In T1: find Q with answer l1, has l2 option (not fixed)
                    q1 = None
                    for q in questions:
                        if q['ticket'] == t1 and q['answer'] == l1:
                            if (q['ticket'], q['num']) in fixed_keys: continue
                            if l2 in q['options'] and q['options'][l1] != q['options'][l2]:
                                # Check local: after swap l1→l2, T1 still valid
                                if pt[t1][l2] < 6:  # l2 won't exceed 6
                                    q1 = q
                                    break
                    
                    if q1 is None: continue
                    
                    # In T2: find Q with answer l2, has l1 option (not fixed)
                    q2 = None
                    for q in questions:
                        if q['ticket'] == t2 and q['answer'] == l2:
                            if (q['ticket'], q['num']) in fixed_keys: continue
                            if l1 in q['options'] and q['options'][l2] != q['options'][l1]:
                                q2 = q
                                break
                    
                    if q2 is None: continue
                    
                    # Perform cross-ticket swap
                    t1_src = q1['options'][l1]
                    t1_tgt = q1['options'][l2]
                    q1['options'][l1] = t1_tgt
                    q1['options'][l2] = t1_src
                    q1['answer'] = l2
                    
                    t2_src = q2['options'][l2]
                    t2_tgt = q2['options'][l1]
                    q2['options'][l2] = t2_tgt
                    q2['options'][l1] = t2_src
                    q2['answer'] = l1
                    
                    print(f"  Б{t1}в{q1['num']}: {l1}→{l2} | Б{t2}в{q2['num']}: {l2}→{l1}")
                    swaps_done += 2
                    made_swap = True
                    break
    
    gc, pt = compute_balance(questions)
    print(f"  Глобальный баланс после: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")
    
    bad = []
    for t in sorted(pt.keys()):
        for letter in 'ABCD':
            cnt = pt[t][letter]
            if cnt < 4 or cnt > 6:
                bad.append(f"Билет {t}: {letter}={cnt}")
    if bad:
        print(f"  ⚠️ Остались нарушения: {bad}")
    else:
        print(f"  Локальный баланс ОК (4-6 в каждом билете)")
    
    return swaps_done


def build_105_docx(header_paras, questions, output_path):
    """Build the 105 docx file with corrected answers."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Write header (non-question paragraphs before tickets)
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Match alignment: if it looks like a title/header, center it
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Фонд') or stripped.startswith('итоговой') or
            stripped.startswith('по основной') or stripped.startswith('(повышения') or
            stripped.startswith('по профессии') or stripped.startswith('«Оператор') or
            stripped.startswith('Тестовые') or stripped.startswith('(20 вопросов') or
            re.match(r'^Билет\s+№?\s*\d+$', stripped)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if re.match(r'^Билет', stripped) else Pt(12)
    
    # Group questions by ticket
    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)
    
    # Write tickets
    for t in sorted(tickets.keys()):
        # Ticket header
        tp = doc.add_paragraph(f'Билет {t}')
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        
        for q in tickets[t]:
            # Question text
            qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)
            
            # Options
            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)
            
            # Inline answer
            ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.left_indent = Cm(0)
            for run in ap.runs:
                run.font.bold = True
    
    doc.save(output_path)
    print(f"\nФайл сохранён: {output_path}")


def main():
    print("="*60)
    print("ИСПРАВЛЕНИЕ 105 КПК: 10 ключей + ребалансировка D")
    print("="*60)
    
    # Step 1: Parse
    print("\n[1] Парсинг...")
    header_paras, questions = parse_105(FILEPATH)
    print(f"  Вопросов: {len(questions)}")
    
    gc, pt = compute_balance(questions)
    print(f"  Исходный D-баланс: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")
    
    # Step 2: Apply key fixes
    print("\n[2] Исправление ключей...")
    apply_key_fixes(questions)
    
    gc, pt = compute_balance(questions)
    print(f"  D-баланс после исправлений: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")
    
    # Step 3: Rebalance global
    print("\n[3] Глобальная ребалансировка D...")
    n_swaps = rebalance(questions)
    
    # Step 3b: Fix local balance violations
    print("\n[3b] Локальная ребалансировка...")
    fixed_keys = {(1,1),(1,14),(3,8),(8,8),(8,14),(9,14),(11,18),(12,2),(13,20),(14,13)}
    n_local = fix_local_balance(questions, fixed_keys)
    
    # Step 4: Verify
    print("\n[4] Финальная верификация...")
    gc, pt = compute_balance(questions)
    print(f"  Финальный D-баланс: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")
    
    # Verify all 300 questions have answers
    no_answer = [q for q in questions if not q['answer']]
    if no_answer:
        print(f"  ⚠️ Вопросов без ответа: {len(no_answer)}")
    else:
        print(f"  Все 300 вопросов имеют ответы")
    
    # Verify no duplicates
    seen = {}
    dupes = 0
    for q in questions:
        key = q['text'].strip()[:100]
        if key in seen:
            dupes += 1
        else:
            seen[key] = True
    print(f"  Дубликатов: {dupes}")
    
    # Verify fixed questions
    print("\n  Проверка исправленных вопросов:")
    expected = {
        (1,1): 'B', (1,14): 'B', (3,8): 'B', (8,8): 'B',
        (8,14): 'B', (9,14): 'B', (11,18): 'B', (12,2): 'C',
        (13,20): 'C', (14,13): 'B'
    }
    for (t,n), exp in sorted(expected.items()):
        q = [x for x in questions if x['ticket']==t and x['num']==n][0]
        status = '✅' if q['answer'] == exp else '❌'
        print(f"    {status} Б{t}в{n}: ответ={q['answer']} (ожидается {exp})")
    
    # Step 5: Build
    print("\n[5] Генерация файла...")
    build_105_docx(header_paras, questions, FILEPATH)
    
    # Step 6: Re-parse and verify
    print("\n[6] Повторная верификация из файла...")
    _, verify_qs = parse_105(FILEPATH)
    gc2, _ = compute_balance(verify_qs)
    print(f"  D-баланс из файла: A={gc2['A']} B={gc2['B']} C={gc2['C']} D={gc2['D']}")
    print(f"  Вопросов из файла: {len(verify_qs)}")
    
    print("\n" + "="*60)
    print("ГОТОВО")
    print("="*60)


if __name__ == '__main__':
    main()
