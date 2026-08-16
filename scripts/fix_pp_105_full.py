#!/usr/bin/env python3
"""
Исправление ошибок в 105. ПО_ПП_ФОС_Оператор станков_2-3_разр.docx

По результатам экспертной проверки 300 вопросов (тот же набор, что и в П 105,
но вопросы расположены в других билетах).

Б9в18 (IT7) уже исправлен в коммите 8d9f863.

Остальные 11 исправлений:

КЛЮЧЕВЫЕ (3):
  Б6в20 (Q120): ключ B→C (Rz = наибольшая высота профиля, ГОСТ 2789)
  Б7в7  (Q127): ключ A→B (разомкнутая линия = для сечений, ГОСТ 2.303)
  Б9в4  (Q164): ключ B→A (межоперационный припуск = за одну операцию)

ТЕКСТОВЫЕ (8):
  Б6в9  (Q109):  ГОСТ 868 → ГОСТ 1050-2013
  Б6в11 (Q111):  «площадь формата» → «размеры формата»
  Б11в10(Q210):  ГОСТ 868 → ГОСТ 1050-2013
  Б12в4 (Q224):  5272 Вт → 5265 Вт (вариант D, ответ D не меняется)
  Б12в11(Q231):  «плавающего суппорта» → «плавающего резцедержателя» + вариант A
  Б15в7 (Q287):  ГОСТ 868 → ГОСТ 1050-2013
  Б15в16(Q296):  «пересекающимися» → «перекрещивающимися»

D-баланс:
  3 ключевых изменения: A(-1+1=0), B(-1+1-1=-1), C(+1), D(0)
  Результат: A=75, B=74, C=76, D=75 → нужен 1 своп C→B
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"
FOLDER = "профессиональной переподготовки 15474 Оператор станков"
FILEPATH = f"{BASE}/{FOLDER}/105. ПО_ПП_ФОС_Оператор станков_2-3_разр.docx"


def parse_105(filepath):
    """Parse 105 file into structured data."""
    doc = Document(filepath)
    header_paras = []
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = {}
    current_answer = None
    in_questions = False

    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()

        if not txt_stripped:
            header_paras.append(txt)
            continue

        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', txt_stripped)
        if ticket_match:
            in_questions = True
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_ticket = int(ticket_match.group(1))
            current_q_text = None
            header_paras.append(txt)
            continue

        if not in_questions:
            header_paras.append(txt)
            continue

        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', txt_stripped)
        if q_match:
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket, 'num': current_q_num,
                    'text': current_q_text, 'options': current_options,
                    'answer': current_answer
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = {}
            current_answer = None
            continue

        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q_text is not None:
            current_options[opt_match.group(1)] = opt_match.group(2)
            continue

        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)

    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket, 'num': current_q_num,
            'text': current_q_text, 'options': current_options,
            'answer': current_answer
        })

    return header_paras, questions


def compute_balance(questions):
    """Compute D-balance globally and per-ticket."""
    global_counts = Counter()
    per_ticket = {}
    for q in questions:
        a = q['answer']
        global_counts[a] += 1
        t = q['ticket']
        if t not in per_ticket:
            per_ticket[t] = Counter()
        per_ticket[t][a] += 1
    return global_counts, per_ticket


def apply_fixes(questions):
    """Apply all fixes: 3 key changes + 8 text changes."""
    fixes_applied = []

    # === KEY CHANGES (3) ===
    key_fixes = {
        (6, 20):  ('B', 'C'),   # Q120: Rz = наибольшая высота профиля (ГОСТ 2789)
        (7, 7):   ('A', 'B'),   # Q127: разомкнутая линия = для сечений (ГОСТ 2.303)
        (9, 4):   ('B', 'A'),   # Q164: межоперационный припуск = за одну операцию
    }

    for q in questions:
        key = (q['ticket'], q['num'])
        if key in key_fixes:
            old_ans, new_ans = key_fixes[key]
            if q['answer'] == old_ans:
                print(f"  КЛЮЧ Б{q['ticket']}в{q['num']}: {old_ans}→{new_ans}  |  {q['text'][:90]}...")
                q['answer'] = new_ans
                fixes_applied.append(f"Б{q['ticket']}в{q['num']}: ключ {old_ans}→{new_ans}")
            elif q['answer'] == new_ans:
                print(f"  Б{q['ticket']}в{q['num']}: ключ уже {new_ans}, пропускаем")
            else:
                print(f"  ⚠️ Б{q['ticket']}в{q['num']}: ожидаемый {old_ans}, найден {q['answer']}")

    # === TEXT FIXES (8) ===

    # Б6в9 (Q109): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 6 and q['num'] == 9:
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б6в9: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б6в9: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Б6в11 (Q111): «площадь формата» → «размеры формата»
    for q in questions:
        if q['ticket'] == 6 and q['num'] == 11:
            q['text'] = q['text'].replace('площадь формата А1', 'размеры (длина и ширина) формата А1')
            print(f"  ТЕКСТ Б6в11: «площадь формата» → «размеры формата»")
            fixes_applied.append("Б6в11: площадь → размеры формата")
            break

    # Б11в10 (Q210): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 11 and q['num'] == 10:
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б11в10: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б11в10: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Б12в4 (Q224): 5272 Вт → 5265 Вт (в варианте D, ответ D не меняется)
    for q in questions:
        if q['ticket'] == 12 and q['num'] == 4:
            for letter in q['options']:
                if '5272' in q['options'][letter]:
                    q['options'][letter] = q['options'][letter].replace('5272', '5265')
                    print(f"  ЗНАЧЕНИЕ Б12в4: 5272 Вт → 5265 Вт (вариант {letter})")
                    fixes_applied.append("Б12в4: 5272→5265 Вт")
                    break
            break

    # Б12в11 (Q231): «плавающего суппорта» → «плавающего резцедержателя» + вариант A
    for q in questions:
        if q['ticket'] == 12 and q['num'] == 11:
            q['text'] = q['text'].replace('плавающего суппорта токарного станка',
                                          'плавающего резцедержателя при обработке центровых отверстий')
            if q['options'].get('A') == 'Автоматическое поддержание постоянной подачи':
                q['options']['A'] = 'Компенсация погрешностей центрирования инструмента'
            print(f"  ТЕКСТ Б12в11: «плавающего суппорта» → «плавающего резцедержателя», вариант A исправлен")
            fixes_applied.append("Б12в11: плавающий суппорт → плавающий резцедержатель + вариант A")
            break

    # Б15в7 (Q287): ГОСТ 868 → ГОСТ 1050-2013
    for q in questions:
        if q['ticket'] == 15 and q['num'] == 7:
            q['text'] = q['text'].replace('ГОСТ 868', 'ГОСТ 1050-2013')
            print(f"  ТЕКСТ Б15в7: «ГОСТ 868» → «ГОСТ 1050-2013»")
            fixes_applied.append("Б15в7: ГОСТ 868 → ГОСТ 1050-2013")
            break

    # Б15в16 (Q296): «пересекающимися» → «перекрещивающимися»
    for q in questions:
        if q['ticket'] == 15 and q['num'] == 16:
            q['text'] = q['text'].replace('пересекающимися', 'перекрещивающимися')
            print(f"  ТЕКСТ Б15в16: «пересекающимися» → «перекрещивающимися»")
            fixes_applied.append("Б15в16: пересекающимися → перекрещивающимися")
            break

    print(f"\nВсего исправлений: {len(fixes_applied)}")
    return len(fixes_applied)


def rebalance(questions, target_per_letter=75):
    """Rebalance D-distribution after key changes. Need: C→B (1 swap)."""
    global_counts, per_ticket = compute_balance(questions)

    print(f"\nБаланс до ребалансировки: A={global_counts['A']} B={global_counts['B']} C={global_counts['C']} D={global_counts['D']}")

    over = {}
    under = {}
    for letter in 'ABCD':
        diff = global_counts[letter] - target_per_letter
        if diff > 0:
            over[letter] = diff
        elif diff < 0:
            under[letter] = -diff

    if not over or not under:
        print("Ребалансировка не нужна!")
        return 0

    print(f"Избыток: {dict(over)}, Дефицит: {dict(under)}")

    fixed_keys = {(6,20), (7,7), (9,4), (9,18), (9,20)}  # user-corrected questions
    used_indices = set()
    swaps = []

    def get_candidates(source_letter, target_letter):
        result = []
        for i, q in enumerate(questions):
            if i in used_indices:
                continue
            if (q['ticket'], q['num']) in fixed_keys:
                continue
            if q['answer'] != source_letter:
                continue
            if target_letter not in q['options']:
                continue
            if q['options'][source_letter] == q['options'][target_letter]:
                continue
            t = q['ticket']
            tc = per_ticket[t]
            if tc[source_letter] <= 4 or tc[target_letter] >= 6:
                continue
            result.append((i, q))
        return result

    total_swaps = 0
    max_iterations = 50

    for iteration in range(max_iterations):
        for letter in 'ABCD':
            diff = global_counts[letter] - target_per_letter
            if diff > 0:
                over[letter] = diff
            else:
                over.pop(letter, None)
        for letter in 'ABCD':
            diff = global_counts[letter] - target_per_letter
            if diff < 0:
                under[letter] = -diff
            else:
                under.pop(letter, None)

        if not over or not under:
            break

        made_swap = False
        for source_letter in sorted(over.keys(), key=lambda x: -over[x]):
            if made_swap:
                break
            for target_letter in sorted(under.keys(), key=lambda x: -under[x]):
                if made_swap:
                    break
                if over.get(source_letter, 0) <= 0 or under.get(target_letter, 0) <= 0:
                    continue

                candidates = get_candidates(source_letter, target_letter)
                if not candidates:
                    continue

                best = None
                best_score = 999
                for idx, q in candidates:
                    t = q['ticket']
                    tc = per_ticket[t]
                    src_after = tc[source_letter] - 1
                    tgt_after = tc[target_letter] + 1
                    score = 0
                    if src_after < 4: score += 100
                    if tgt_after > 6: score += 100
                    score += abs(src_after - 5) + abs(tgt_after - 5)
                    if score < best_score:
                        best_score = score
                        best = (idx, q)

                if best is None:
                    continue

                idx, q = best
                src_text = q['options'][source_letter]
                tgt_text = q['options'][target_letter]

                q['options'][source_letter] = tgt_text
                q['options'][target_letter] = src_text
                q['answer'] = target_letter
                used_indices.add(idx)

                global_counts[source_letter] -= 1
                global_counts[target_letter] += 1
                over[source_letter] = over.get(source_letter, 0) - 1
                under[target_letter] = under.get(target_letter, 0) - 1
                if over.get(source_letter, 0) <= 0:
                    over.pop(source_letter, None)
                if under.get(target_letter, 0) <= 0:
                    under.pop(target_letter, None)

                t = q['ticket']
                per_ticket[t][source_letter] -= 1
                per_ticket[t][target_letter] += 1

                total_swaps += 1
                made_swap = True
                swaps.append(f"  Б{q['ticket']}в{q['num']}: {source_letter}→{target_letter}")

    if swaps:
        print(f"\nСвопы для ребалансировки ({total_swaps}):")
        for s in swaps:
            print(s)

    print(f"\nБаланс после ребалансировки: A={global_counts['A']} B={global_counts['B']} C={global_counts['C']} D={global_counts['D']}")

    bad_tickets = []
    for t in sorted(per_ticket.keys()):
        for letter in 'ABCD':
            cnt = per_ticket[t][letter]
            if cnt < 4 or cnt > 6:
                bad_tickets.append(f"Билет {t}: {letter}={cnt}")
    if bad_tickets:
        print(f"⚠️ Нарушения локального баланса: {bad_tickets}")
    else:
        print("Локальный баланс ОК (4-6 в каждом билете)")

    return total_swaps


def fix_local_balance(questions, fixed_keys):
    """Fix per-ticket balance violations with cross-ticket swaps."""
    gc, pt = compute_balance(questions)

    print(f"\n[локальная ребалансировка]")
    print(f"  Глобальный баланс перед: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    def get_violations():
        viols = []
        for t in sorted(pt.keys()):
            for letter in 'ABCD':
                cnt = pt[t][letter]
                if cnt < 4:
                    viols.append((t, letter, 'deficit', cnt))
                elif cnt > 6:
                    viols.append((t, letter, 'excess', cnt))
        return viols

    violations = get_violations()
    if not violations:
        print("  Локальный баланс уже ОК")
        return 0

    print(f"  Нарушения: {[(f'Б{v[0]}:{v[1]}={v[3]}') for v in violations]}")

    swaps_done = 0
    for iteration in range(20):
        gc, pt = compute_balance(questions)
        violations = get_violations()
        if not violations:
            break

        made_swap = False
        for t1, l1, kind1, cnt1 in violations:
            if made_swap: break
            if kind1 != 'excess': continue

            for t1b, l2, kind2, cnt2 in violations:
                if made_swap: break
                if t1b != t1: continue
                if kind2 != 'deficit': continue

                for t2 in range(1, 16):
                    if t2 == t1: continue
                    tc2 = pt[t2]
                    if tc2[l1] >= 6 or tc2[l2] <= 4:
                        continue

                    q1 = None
                    for q in questions:
                        if q['ticket'] == t1 and q['answer'] == l1:
                            if (q['ticket'], q['num']) in fixed_keys: continue
                            if l2 in q['options'] and q['options'][l1] != q['options'][l2]:
                                if pt[t1][l2] < 6:
                                    q1 = q
                                    break

                    if q1 is None: continue

                    q2 = None
                    for q in questions:
                        if q['ticket'] == t2 and q['answer'] == l2:
                            if (q['ticket'], q['num']) in fixed_keys: continue
                            if l1 in q['options'] and q['options'][l2] != q['options'][l1]:
                                q2 = q
                                break

                    if q2 is None: continue

                    t1_src = q1['options'][l1]
                    t1_tgt = q1['options'][l2]
                    q1['options'][l1] = t1_tgt
                    q1['options'][l2] = t1_src
                    q1['answer'] = l2

                    t2_src = q2['options'][l2]
                    t2_tgt = q2['options'][l1]
                    q2['options'][l2] = t2_tgt
                    q2['options'][l1] = t2_src
                    q2['answer'] = l1

                    print(f"  Б{t1}в{q1['num']}: {l1}→{l2} | Б{t2}в{q2['num']}: {l2}→{l1}")
                    swaps_done += 2
                    made_swap = True
                    break

    gc, pt = compute_balance(questions)
    print(f"  Глобальный баланс после: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    bad = []
    for t in sorted(pt.keys()):
        for letter in 'ABCD':
            cnt = pt[t][letter]
            if cnt < 4 or cnt > 6:
                bad.append(f"Билет {t}: {letter}={cnt}")
    if bad:
        print(f"  ⚠️ Остались нарушения: {bad}")
    else:
        print(f"  Локальный баланс ОК")

    return swaps_done


def build_105_docx(header_paras, questions, output_path):
    """Build the 105 docx file with corrected content."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Фонд') or stripped.startswith('итоговой') or
            stripped.startswith('по основной') or stripped.startswith('(профессиональной') or
            stripped.startswith('по профессии') or stripped.startswith('«Оператор') or
            stripped.startswith('Тестовые') or stripped.startswith('(20 вопросов') or
            re.match(r'^Билет\s+№?\s*\d+$', stripped)):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14) if re.match(r'^Билет', stripped) else Pt(12)

    tickets = {}
    for q in questions:
        t = q['ticket']
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q)

    for t in sorted(tickets.keys()):
        tp = doc.add_paragraph(f'Билет {t}')
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(6)
        for run in tp.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        for q in tickets[t]:
            qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
            qp.paragraph_format.space_after = Pt(0)
            qp.paragraph_format.space_before = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)

            for letter in 'ABCD':
                if letter in q['options']:
                    op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                    op.paragraph_format.space_after = Pt(0)
                    op.paragraph_format.space_before = Pt(0)
                    op.paragraph_format.left_indent = Cm(1)

            ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
            ap.paragraph_format.space_after = Pt(4)
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.left_indent = Cm(0)
            for run in ap.runs:
                run.font.bold = True

    doc.save(output_path)
    print(f"\nФайл сохранён: {output_path}")


def main():
    print("=" * 60)
    print("ИСПРАВЛЕНИЕ 105 ПП: 11 исправлений (3 ключа + 8 текста) + ребаланс")
    print("=" * 60)

    print("\n[1] Парсинг...")
    header_paras, questions = parse_105(FILEPATH)
    print(f"  Вопросов: {len(questions)}")

    gc, pt = compute_balance(questions)
    print(f"  Исходный D-баланс: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    print("\n[2] Применение исправлений...")
    n_fixes = apply_fixes(questions)

    gc, pt = compute_balance(questions)
    print(f"  D-баланс после исправлений: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    print("\n[3] Глобальная ребалансировка D...")
    n_swaps = rebalance(questions)

    print("\n[3b] Локальная ребалансировка...")
    fixed_keys = {(6,20), (7,7), (9,4), (9,18), (9,20)}
    n_local = fix_local_balance(questions, fixed_keys)

    print("\n[4] Верификация исправленных вопросов...")
    checks = {
        (6, 9):   ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
        (6, 11):  ('text_contains', 'размеры', 'площадь'),
        (6, 20):  ('key', 'C'),
        (7, 7):   ('key', 'B'),
        (9, 4):   ('key', 'A'),
        (9, 18):  ('key', 'B'),  # Already fixed in prev commit
        (11, 10): ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
        (12, 4):  ('option_no_contains', '5272'),
        (12, 11): ('text_contains', 'плавающего резцедержателя', 'плавающего суппорта'),
        (15, 7):  ('text_contains', 'ГОСТ 1050-2013', 'ГОСТ 868'),
        (15, 16): ('text_contains', 'перекрещивающимися', 'пересекающимися'),
    }

    all_ok = True
    for q in questions:
        key = (q['ticket'], q['num'])
        if key not in checks:
            continue
        check = checks[key]
        check_type = check[0]

        if check_type == 'key':
            expected = check[1]
            ok = q['answer'] == expected
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: ключ={q['answer']} (ожидается {expected})")
            if not ok: all_ok = False

        elif check_type == 'text_contains':
            should_have = check[1]
            should_not = check[2]
            has_good = should_have in q['text']
            has_bad = should_not and should_not in q['text']
            ok = has_good and not has_bad
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: «{should_have}»={'да' if has_good else 'НЕТ'}, «{should_not}»={'да' if has_bad else 'нет'}")
            if not ok: all_ok = False

        elif check_type == 'option_no_contains':
            bad_str = check[1]
            found_bad = any(bad_str in v for v in q['options'].values())
            ok = not found_bad
            status = '✅' if ok else '❌'
            print(f"  {status} Б{q['ticket']}в{q['num']}: варианты {'не ' if ok else ''}содержат «{bad_str}»")
            if not ok: all_ok = False

    print("\n[5] Генерация файла...")
    build_105_docx(header_paras, questions, FILEPATH)

    print("\n[6] Повторная верификация из файла...")
    _, verify_qs = parse_105(FILEPATH)
    gc2, pt2 = compute_balance(verify_qs)
    print(f"  Вопросов из файла: {len(verify_qs)}")
    print(f"  D-баланс из файла: A={gc2['A']} B={gc2['B']} C={gc2['C']} D={gc2['D']}")

    no_answer = [q for q in verify_qs if not q['answer']]
    if no_answer:
        print(f"  ⚠️ Вопросов без ответа: {len(no_answer)}")
        all_ok = False
    else:
        print(f"  Все {len(verify_qs)} вопросов имеют ответы")

    if all_ok:
        print("\n  🟢 ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ")
    else:
        print("\n  🔴 ЕСТЬ ПРОБЛЕМЫ")

    print("\n" + "=" * 60)
    print("ГОТОВО")
    print("=" * 60)


if __name__ == '__main__':
    main()
