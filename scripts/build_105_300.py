#!/usr/bin/env python3
"""
Build 105: 15 tickets x 20 questions = 300 questions total.
150 existing + 150 new = 300 unique, D-balanced.
"""
import json, random, copy, sys
sys.path.insert(0, '/home/z/my-project/vol/scripts')
from new_150_questions import NEW_QUESTIONS

random.seed(42)

# ============================================================
# 1. Load existing 150 questions
# ============================================================
with open('/home/z/my-project/vol/scripts/existing_150_questions.json', 'r', encoding='utf-8') as f:
    existing = json.load(f)

print(f"Loaded {len(existing)} existing questions")
print(f"Loaded {len(NEW_QUESTIONS)} new questions")

# Convert existing to same format: (topic, q, a, b, c, d, ans)
all_questions = []
for q in existing:
    all_questions.append({
        'topic': 'existing',
        'question': q['question'],
        'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D'],
        'ans': q['ans']
    })

# Add new questions
for topic, q_text, a, b, c, d, ans in NEW_QUESTIONS:
    all_questions.append({
        'topic': topic,
        'question': q_text,
        'A': a, 'B': b, 'C': c, 'D': d,
        'ans': ans
    })

print(f"Total: {len(all_questions)} questions")

# Check for duplicate question texts
from collections import Counter
texts = [q['question'] for q in all_questions]
dupes = {t: c for t, c in Counter(texts).items() if c > 1}
if dupes:
    print(f"WARNING: {len(dupes)} duplicates!")
    for t, c in dupes.items():
        print(f"  [{c}x] {t[:60]}...")
    sys.exit(1)
else:
    print("No duplicate questions.")

# ============================================================
# 2. D-balance by swapping option positions
# ============================================================
target_per_option = len(all_questions) // 4  # 75

counts = Counter(q['ans'] for q in all_questions)
print(f"Before D-balance: {dict(sorted(counts.items()))}")

def swap_answer(q, target_letter):
    """Swap options so the correct answer moves to target_letter."""
    current = q['ans']
    if current == target_letter:
        return q
    
    opts = {'A': q['A'], 'B': q['B'], 'C': q['C'], 'D': q['D']}
    letters = ['A', 'B', 'C', 'D']
    
    # Find positions of current and target
    ci = letters.index(current)
    ti = letters.index(target_letter)
    
    # Swap the texts at these positions
    opts[letters[ci]], opts[letters[ti]] = opts[letters[ti]], opts[letters[ci]]
    
    q['A'] = opts['A']
    q['B'] = opts['B']
    q['C'] = opts['C']
    q['D'] = opts['D']
    q['ans'] = target_letter
    return q

# Shuffle to randomize which questions get reassigned
indices = list(range(len(all_questions)))
random.shuffle(indices)

# First pass: try to balance
need_more = {}
for letter in 'ABCD':
    diff = target_per_option - counts[letter]
    if diff > 0:
        need_more[letter] = diff

# Find questions from overrepresented options to reassign
idx = 0
for letter, need in need_more.items():
    reassigned = 0
    for i in indices:
        if reassigned >= need:
            break
        q = all_questions[i]
        if q['ans'] not in need_more and q['ans'] != letter:
            # Take from an overrepresented option
            all_questions[i] = swap_answer(copy.deepcopy(q), letter)
            reassigned += 1

# Second pass: if still not balanced, do another round
counts = Counter(q['ans'] for q in all_questions)
print(f"After D-balance pass 1: {dict(sorted(counts.items()))}")

need_more = {}
for letter in 'ABCD':
    diff = target_per_option - counts[letter]
    if diff > 0:
        need_more[letter] = diff

for letter, need in need_more.items():
    reassigned = 0
    for i in indices:
        if reassigned >= need:
            break
        q = all_questions[i]
        if q['ans'] not in need_more and q['ans'] != letter:
            all_questions[i] = swap_answer(copy.deepcopy(q), letter)
            reassigned += 1

counts = Counter(q['ans'] for q in all_questions)
print(f"After D-balance pass 2: {dict(sorted(counts.items()))}")

# Verify all answers are correct after swapping
errors = 0
for i, q in enumerate(all_questions):
    correct_text = q[q['ans']]
    # Check that the correct text is unique among options
    opts = [q['A'], q['B'], q['C'], q['D']]
    if len(opts) != len(set(opts)):
        print(f"ERROR Q{i+1}: duplicate options after swap!")
        errors += 1

if errors:
    print(f"{errors} errors found!")
    sys.exit(1)
else:
    print("All 300 answers verified (no duplicate options).")

# ============================================================
# 3. Distribute into 15 tickets x 20 questions with ticket-level D-balance
# ============================================================
# Strategy: sort by answer letter, then deal round-robin into tickets
# This ensures each ticket gets ~5 of each letter

TICKETS = 15
Q_PER_TICKET = 20

# Group questions by answer letter
by_letter = {letter: [] for letter in 'ABCD'}
for q in all_questions:
    by_letter[q['ans']].append(q)

# Shuffle within each group
for letter in 'ABCD':
    random.shuffle(by_letter[letter])

# Round-robin deal: for each letter group, distribute 1 question per ticket in rotation
# Each letter has 75 questions. 75 / 15 = 5 exactly per ticket.
tickets = [[] for _ in range(TICKETS)]

for letter in 'ABCD':
    group = by_letter[letter]
    for i, q in enumerate(group):
        ticket_idx = i % TICKETS
        tickets[ticket_idx].append(q)

# Shuffle within each ticket
for t in tickets:
    random.shuffle(t)

print(f"\nFormed {len(tickets)} tickets x {len(tickets[0])} questions")

# Ticket-level D-balance check
for t_idx, ticket in enumerate(tickets):
    tc = Counter(q['ans'] for q in ticket)
    # print(f"  Ticket {t_idx+1}: {dict(sorted(tc.items()))}")

# ============================================================
# 4. Generate DOCX
# ============================================================
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

# Title page
title_lines = [
    ("Фонд оценочных средств", True, 16),
    ("итоговой аттестации", True, 16),
    ("", False, 14),
    ("по основной программе профессионального обучения", False, 14),
    ("(профессиональной подготовки)", False, 14),
    ("", False, 14),
    ("по профессии", False, 14),
    ("\u00ABОператор автоматических и полуавтоматических станков и линий станков\u00BB", True, 14),
    ("", False, 14),
    ("Тестовые задания с выбором одного правильного ответа", False, 14),
    ("(20 вопросов)", False, 14),
]

for text, bold, size in title_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

# Page break before tickets
doc.add_page_break()

# Generate tickets
for t_idx, ticket in enumerate(tickets):
    # Ticket header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Билет {t_idx + 1}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    # Add empty line after ticket header
    doc.add_paragraph()
    
    for q_idx, q in enumerate(ticket, 1):
        # Question text
        p = doc.add_paragraph()
        run = p.add_run(f"{q_idx}. {q['question']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Options
        for letter in 'ABCD':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"{letter}) {q[letter]};")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        # Correct answer (inline)
        p = doc.add_paragraph()
        run = p.add_run(f"Правильный ответ под номером: {q['ans']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    
    # Add page break between tickets (except after last)
    if t_idx < len(tickets) - 1:
        doc.add_page_break()

# Save
output_path = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/профессиональной подготовки 15474 Оператор станков/105. ПО_П_ФОС_Оператор станков_2-3_разр.docx"
doc.save(output_path)
print(f"\nSaved: {output_path}")

# Final summary
print(f"\n=== FINAL SUMMARY ===")
print(f"Total questions: {len(all_questions)}")
print(f"Tickets: {len(tickets)} x {Q_PER_TICKET} questions")
print(f"D-balance: {dict(sorted(counts.items()))}")
for letter in 'ABCD':
    pct = counts[letter] / len(all_questions) * 100
    print(f"  {letter}: {counts[letter]} ({pct:.1f}%)")

# Topic distribution
from collections import Counter
topic_counts = Counter(q['topic'] for q in all_questions)
for topic, count in sorted(topic_counts.items()):
    print(f"  {topic}: {count}")
