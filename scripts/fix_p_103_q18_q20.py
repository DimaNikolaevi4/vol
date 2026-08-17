#!/usr/bin/env python3
"""
Замена вопросов 18 и 20 в 103. ПО_П_ФОС_Оператор станков_2-3_разр ОП.0.0.docx

Q18: уточнённая формулировка про припуск (ответ A остаётся)
Q20: новая формулировка по Приказу Минтруда № 903н (ответ D→B)

D-баланс: 5/5/5/5 → 5/6/5/4 (в пределах 4-6)
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"
FOLDER = "профессиональной подготовки 15474 Оператор станков"
FILEPATH = f"{BASE}/{FOLDER}/103. ПО_П_ФОС_Оператор станков_2-3_разр ОП.0.0.docx"


def parse_103(filepath):
    """Parse 103/104 file: header + 20 questions."""
    doc = Document(filepath)
    header_paras = []
    questions = []
    current_q = None
    in_questions = False

    for p in doc.paragraphs:
        txt = p.text
        txt_stripped = txt.strip()

        if not txt_stripped:
            header_paras.append(txt)
            continue

        q_match = re.match(r'^(\d+)\s*[.)]\s*(.*)', txt_stripped)
        if q_match:
            in_questions = True
            if current_q is not None:
                questions.append(current_q)
            current_q = {
                'num': int(q_match.group(1)),
                'text': q_match.group(2),
                'options': {},
                'answer': None
            }
            continue

        if not in_questions:
            header_paras.append(txt)
            continue

        opt_match = re.match(r'^([ABCD])\)\s*(.*)', txt_stripped)
        if opt_match and current_q is not None:
            current_q['options'][opt_match.group(1)] = opt_match.group(2)
            continue

        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', txt_stripped)
        if ans_match and current_q is not None:
            current_q['answer'] = ans_match.group(1)

    if current_q is not None:
        questions.append(current_q)

    return header_paras, questions


def build_103_docx(header_paras, questions, output_path):
    """Build 103/104 docx with standard formatting."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Header
    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        stripped = txt.strip()
        if not stripped:
            continue
        if (stripped.startswith('Приложение') or stripped.startswith('Фонд') or
            stripped.startswith('итоговой') or stripped.startswith('по основной') or
            stripped.startswith('(профессиональной') or stripped.startswith('по профессии') or
            stripped.startswith('«Оператор') or stripped.startswith('Тестовые') or
            stripped.startswith('(20 вопросов')):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True

    # Questions
    for q in questions:
        qp = doc.add_paragraph(f"{q['num']}. {q['text']}")
        qp.paragraph_format.space_after = Pt(0)
        qp.paragraph_format.space_before = Pt(3)

        for letter in 'ABCD':
            if letter in q['options']:
                op = doc.add_paragraph(f"{letter}) {q['options'][letter]}")
                op.paragraph_format.space_after = Pt(0)
                op.paragraph_format.space_before = Pt(0)
                op.paragraph_format.left_indent = Cm(1)

        ap = doc.add_paragraph(f"Правильный ответ под номером: {q['answer']}")
        ap.paragraph_format.space_after = Pt(4)
        ap.paragraph_format.space_before = Pt(0)
        for run in ap.runs:
            run.font.bold = True

    doc.save(output_path)
    print(f"Файл сохранён: {output_path}")


def main():
    print("=" * 60)
    print("ЗАМЕНА Q18 и Q20 в 103 П ОП.0.0")
    print("=" * 60)

    # Parse
    print("\n[1] Парсинг...")
    header_paras, questions = parse_103(FILEPATH)
    print(f"  Вопросов: {len(questions)}")

    gc = Counter(q['answer'] for q in questions)
    print(f"  D-баланс до: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    # Show old Q18 and Q20
    for q in questions:
        if q['num'] in (18, 20):
            print(f"\n  Старый Q{q['num']}: {q['text'][:100]}")
            print(f"    Ответ: {q['answer']}")

    # Replace Q18
    print("\n[2] Замена Q18...")
    for q in questions:
        if q['num'] == 18:
            q['text'] = 'Как называется слой материала, который необходимо удалить с поверхности заготовки для получения готовой детали?'
            q['options'] = {
                'A': 'Припуск',
                'B': 'Допуск',
                'C': 'Шероховатость',
                'D': 'Квалитет'
            }
            q['answer'] = 'A'
            print(f"  Новая формулировка (ответ A: Припуск)")
            break

    # Replace Q20
    print("\n[3] Замена Q20...")
    for q in questions:
        if q['num'] == 20:
            q['text'] = 'Согласно п. 2.3 Приказа Минтруда России от 15.12.2020 № 903н, работники, использующие в работе ручные электрические машины, переносной электроинструмент и светильники, а также электротехнологический персонал, должны иметь группу по электробезопасности не ниже:'
            q['options'] = {
                'A': 'I группы',
                'B': 'II группы',
                'C': 'III группы',
                'D': 'IV группы'
            }
            q['answer'] = 'B'
            print(f"  Новая формулировка (ответ B: II группы)")
            break

    # Verify balance
    gc = Counter(q['answer'] for q in questions)
    print(f"\n[4] D-баланс после: A={gc['A']} B={gc['B']} C={gc['C']} D={gc['D']}")

    # Verify numbering
    nums = [q['num'] for q in questions]
    print(f"  Нумерация: {nums}")
    assert nums == list(range(1, 21)), "Нумерация сбилась!"
    print("  Нумерация ОК (1-20)")

    # Verify all have answers
    no_ans = [q for q in questions if not q['answer']]
    if no_ans:
        print(f"  ⚠️ Вопросов без ответа: {no_ans}")
    else:
        print(f"  Все 20 вопросов имеют ответы")

    # Build
    print("\n[5] Генерация файла...")
    build_103_docx(header_paras, questions, FILEPATH)

    # Re-verify
    print("\n[6] Поверка из файла...")
    _, verify = parse_103(FILEPATH)
    gc2 = Counter(q['answer'] for q in verify)
    print(f"  Вопросов: {len(verify)}")
    print(f"  D-баланс: A={gc2['A']} B={gc2['B']} C={gc2['C']} D={gc2['D']}")

    # Show new Q18 and Q20
    for q in verify:
        if q['num'] in (18, 20):
            print(f"\n  Q{q['num']}: {q['text'][:120]}...")
            print(f"    A) {q['options']['A']}")
            print(f"    B) {q['options']['B']}")
            print(f"    C) {q['options']['C']}")
            print(f"    D) {q['options']['D']}")
            print(f"    Ответ: {q['answer']}")

    print("\n" + "=" * 60)
    print("ГОТОВО")
    print("=" * 60)


if __name__ == '__main__':
    main()
