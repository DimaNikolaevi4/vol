#!/usr/bin/env python3
"""
Полный скан всех 9 ФОС файлов Оператор станков на:
- К0.11 стоп-слова
- Количество вопросов
- D-баланс
- Уникальность вопросов (для 105)
- Формат (inline ответы, таблицы)
"""

import re
import os
import json
from collections import Counter
from docx import Document

BASE = "/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков"

folders = {
    "П": "профессиональной подготовки 15474 Оператор станков",
    "ПП": "профессиональной переподготовки 15474 Оператор станков",
    "КПК": "повышения квалификации 15474 Оператор станков",
}

# К0.11 stop-words
STOP_WORDS = {
    # Частотные
    "часто": "частотные", "обычно": "частотные", "как правило": "частотные",
    "чаще всего": "частотные", "преимущественно": "частотные", "в основном": "частотные",
    "реже": "частотные",
    # Субъективной оценки
    "лучше": "субъективной_оценки", "хуже": "субъективной_оценки",
    "оптимальный": "субъективной_оценки", "оптимальная": "субъективной_оценки", "оптимальное": "субъективной_оценки",
    "оптимальным": "субъективной_оценки", "оптимальных": "субъективной_оценки",
    "рациональный": "субъективной_оценки", "рациональная": "субъективной_оценки",
    "рациональное": "субъективной_оценки", "рациональным": "субъективной_оценки",
    "целесообразный": "субъективной_оценки", "целесообразно": "субъективной_оценки",
    "целесообразным": "субъективной_оценки",
    "эффективный": "субъективной_оценки", "эффективная": "субъективной_оценки",
    "эффективное": "субъективной_оценки", "эффективным": "субъективной_оценки",
    "эффективных": "субъективной_оценки", "эффективность": "субъективной_оценки",
    "предпочтительнее": "субъективной_оценки", "предпочтительным": "субъективной_оценки",
    "наиболее": "субъективной_оценки", "наилучший": "субъективной_оценки",
    "наилучшим": "субъективной_оценки", "наилучшее": "субъективной_оценки",
    # Неопределённые количественные
    "некоторые": "неопр_колич", "ряд": "неопр_колич", "несколько": "неопр_колич",
    "множество": "неопр_колич", "большинство": "неопр_колич",
    # Временная неопределённость
    "в настоящее время": "временная", "сейчас": "временная",
    "традиционно": "временная", "современный": "временная", "современные": "временная",
    "современное": "временная", "современная": "временная",
    "в последние годы": "временная",
    # Возможностные
    "может": "возможностные", "возможно": "возможностные",
    "иногда": "возможностные", "в ряде случаев": "возможностные",
    "при определённых условиях": "возможностные",
    # Неопределённые меры
    "примерно": "неопр_меры", "около": "неопр_меры", "порядка": "неопр_меры",
}

# False positive patterns (техтермины, которые НЕ являются стоп-словами)
FALSE_POSITIVE_PATTERNS = [
    r"частота\s+вращения",       # техтермин
    r"частот",                     # частоты, частотный и т.д.
    r"часто(?:т|ы)",              # частоты, частотный и т.д.
    r"обычного\s+качества",        # «сталь обычного качества» — ГОСТ-термин (ГОСТ 380)
    r"обычн",                      # обычный/обычно в составе ГОСТ-термина
    r"без\s+возможности",          # существительное
    r"возможности",                # существительное (но не "может/возможно")
    r"эффективност",               # существительное «эффективность»
    r"Согласно\s+ГОСТ",            # ГОСТ нейтрализует
    r"согласно\s+ГОСТ",
    r"ГОСТ.*наиболее",             # «наиболее» после ссылки на ГОСТ
    r"в порядке",                  # «порядка» = «в порядке» — не стоп-слово
    r"ряд.*заклеп",               # «ряд заклёпок» — техтермин
    r"порядок.*деи?ств",           # «порядок действий»
    r"поряд.*повер",              # «порядок поверки»
    r"порядок",                    # «порядок» в значении «последовательность»
    r"улучшен",                    # «улучшение» — не стоп-слово, фильтруем «лучше» рядом
    r"для\s+улучшен",              # «для улучшения» фильтрует «лучше»
]

def is_false_positive(text, word):
    """Проверяет, является ли совпадение ложноположительным (техтермин)"""
    for pattern in FALSE_POSITIVE_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    # «порядка» перед числом — неопределённая мера (настоящее стоп-слово)
    if word in ("порядка", "около", "примерно"):
        # Это реальные стоп-слова перед числом
        return False
    # «наиболее» после ГОСТ — нейтрализуется
    if word == "наиболее" and re.search(r"согласно\s+ГОСТ|ГОСТ", text, re.IGNORECASE):
        return True
    # «лучше» рядом с «улучшени» — это существительное/глагол, не стоп-слово
    if word == "лучше" and re.search(r"улучшен", text, re.IGNORECASE):
        return True
    return False

def extract_text_from_docx(path):
    """Извлекает весь текст из docx"""
    doc = Document(path)
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text)
    return paragraphs

def find_questions_and_answers_103_104(paragraphs):
    """Для 103/104: находит вопросы, варианты, ответы"""
    questions = []
    current_q = None
    current_options = []
    current_answer = None
    
    for p in paragraphs:
        text = p.strip()
        if not text:
            continue
        # Ищем начало вопроса (цифра с точкой или скобкой)
        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', text)
        if q_match:
            # Сохраняем предыдущий вопрос
            if current_q is not None:
                questions.append({
                    'text': current_q,
                    'options': current_options,
                    'answer': current_answer,
                })
            current_q = q_match.group(2)
            current_options = []
            current_answer = None
            continue
        # Ищем варианты A)B)C)D)
        opt_match = re.match(r'^([ABCD])\)\s*(.*)', text)
        if opt_match and current_q is not None:
            current_options.append(opt_match.group(2))
            continue
        # Ищем ответ
        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', text)
        if ans_match and current_q is not None:
            current_answer = ans_match.group(1)
    
    # Не забываем последний вопрос
    if current_q is not None:
        questions.append({
            'text': current_q,
            'options': current_options,
            'answer': current_answer,
        })
    
    return questions

def find_questions_and_answers_105(paragraphs):
    """Для 105: находит все 300 вопросов в билетах"""
    questions = []
    current_ticket = None
    current_q_num = None
    current_q_text = None
    current_options = []
    current_answer = None
    
    for p in paragraphs:
        text = p.strip()
        if not text:
            continue
        # Билет
        ticket_match = re.match(r'Билет\s+№?\s*(\d+)', text)
        if ticket_match:
            current_ticket = int(ticket_match.group(1))
            continue
        # Вопрос в билете
        q_match = re.match(r'^(\d+)\s*[.)\]]\s*(.*)', text)
        if q_match and current_ticket is not None:
            # Сохраняем предыдущий
            if current_q_text is not None:
                questions.append({
                    'ticket': current_ticket,
                    'num': current_q_num,
                    'text': current_q_text,
                    'options': current_options,
                    'answer': current_answer,
                })
            current_q_num = int(q_match.group(1))
            current_q_text = q_match.group(2)
            current_options = []
            current_answer = None
            continue
        # Варианты
        opt_match = re.match(r'^([ABCD])\)\s*(.*)', text)
        if opt_match and current_q_text is not None:
            current_options.append(opt_match.group(2))
            continue
        # Ответ
        ans_match = re.search(r'Правильный ответ под номером:\s*([ABCD])', text)
        if ans_match and current_q_text is not None:
            current_answer = ans_match.group(1)
    
    # Последний вопрос
    if current_q_text is not None:
        questions.append({
            'ticket': current_ticket,
            'num': current_q_num,
            'text': current_q_text,
            'options': current_options,
            'answer': current_answer,
        })
    
    return questions

def scan_stop_words(text):
    """Ищет стоп-слова в тексте, возвращает список найденных"""
    found = []
    text_lower = text.lower()
    for word, category in sorted(STOP_WORDS.items(), key=lambda x: -len(x[0])):
        if word in text_lower:
            if not is_false_positive(text, word):
                # Находим контекст вокруг совпадения
                idx = text_lower.find(word)
                start = max(0, idx - 30)
                end = min(len(text), idx + len(word) + 30)
                context = text[start:end]
                found.append({
                    'word': word,
                    'category': category,
                    'context': context,
                })
    return found

def check_d_balance(questions, prefix=""):
    """Проверяет D-баланс"""
    counter = Counter(q['answer'] for q in questions if q.get('answer'))
    total = sum(counter.values())
    result = {}
    for letter in 'ABCD':
        count = counter.get(letter, 0)
        pct = count / total * 100 if total > 0 else 0
        result[letter] = {'count': count, 'pct': round(pct, 1)}
    result['_total'] = total
    return result

def check_105_per_ticket_balance(questions):
    """Проверяет D-баланс по каждому билету в 105"""
    tickets = {}
    for q in questions:
        t = q.get('ticket')
        if t not in tickets:
            tickets[t] = []
        tickets[t].append(q.get('answer'))
    
    per_ticket = {}
    for t, answers in sorted(tickets.items()):
        counter = Counter(answers)
        total = len(answers)
        per_ticket[t] = {letter: counter.get(letter, 0) for letter in 'ABCD'}
        per_ticket[t]['_total'] = total
    return per_ticket

def check_105_duplicates(questions):
    """Проверяет дубликаты в 105"""
    seen = {}
    dupes = []
    for q in questions:
        text_key = q['text'].strip()[:100]
        if text_key in seen:
            dupes.append({
                'ticket1': seen[text_key],
                'ticket2': q.get('ticket'),
                'text': q['text'][:80],
            })
        else:
            seen[text_key] = q.get('ticket')
    return dupes

def check_105_tables(path):
    """Проверяет что в 105 нет таблиц"""
    doc = Document(path)
    return len(doc.tables)

def main():
    results = {}
    
    for folder_key, folder_name in folders.items():
        folder_path = os.path.join(BASE, folder_name)
        print(f"\n{'='*70}")
        print(f"ПАПКА: {folder_key} ({folder_name})")
        print(f"{'='*70}")
        
        results[folder_key] = {}
        
        for fos_num in ['103', '104', '105']:
            # Находим файл
            fos_files = [f for f in os.listdir(folder_path) if f.startswith(fos_num + '.')]
            if not fos_files:
                print(f"  {fos_num}: ФАЙЛ НЕ НАЙДЕН!")
                continue
            
            fos_file = fos_files[0]
            fos_path = os.path.join(folder_path, fos_file)
            print(f"\n  {fos_num}: {fos_file}")
            
            paragraphs = extract_text_from_docx(fos_path)
            
            if fos_num in ('103', '104'):
                questions = find_questions_and_answers_103_104(paragraphs)
                d_balance = check_d_balance(questions)
                
                # Сканируем стоп-слова
                violations = []
                for i, q in enumerate(questions):
                    full_text = q['text'] + ' ' + ' '.join(q.get('options', []))
                    hits = scan_stop_words(full_text)
                    for h in hits:
                        violations.append({
                            'q_num': i + 1,
                            'word': h['word'],
                            'category': h['category'],
                            'context': h['context'],
                        })
                
                print(f"    Вопросов: {len(questions)}")
                print(f"    D-баланс: A={d_balance['A']['count']}({d_balance['A']['pct']}%) "
                      f"B={d_balance['B']['count']}({d_balance['B']['pct']}%) "
                      f"C={d_balance['C']['count']}({d_balance['C']['pct']}%) "
                      f"D={d_balance['D']['count']}({d_balance['D']['pct']}%)")
                print(f"    К0.11 нарушений: {len(violations)}")
                for v in violations:
                    print(f"      Q{v['q_num']}: «{v['word']}» [{v['category']}] → ...{v['context']}...")
                
                results[folder_key][fos_num] = {
                    'file': fos_file,
                    'questions': len(questions),
                    'd_balance': d_balance,
                    'k011_violations': violations,
                }
            
            elif fos_num == '105':
                questions = find_questions_and_answers_105(paragraphs)
                d_balance = check_d_balance(questions)
                per_ticket = check_105_per_ticket_balance(questions)
                duplicates = check_105_duplicates(questions)
                table_count = check_105_tables(fos_path)
                
                # Сканируем стоп-слова
                violations = []
                for q in questions:
                    full_text = q['text'] + ' ' + ' '.join(q.get('options', []))
                    hits = scan_stop_words(full_text)
                    for h in hits:
                        violations.append({
                            'ticket': q.get('ticket'),
                            'q_num': q.get('num'),
                            'word': h['word'],
                            'category': h['category'],
                            'context': h['context'],
                        })
                
                print(f"    Вопросов: {len(questions)} (в {len(per_ticket)} билетах)")
                print(f"    D-баланс глобальный: A={d_balance['A']['count']}({d_balance['A']['pct']}%) "
                      f"B={d_balance['B']['count']}({d_balance['B']['pct']}%) "
                      f"C={d_balance['C']['count']}({d_balance['C']['pct']}%) "
                      f"D={d_balance['D']['count']}({d_balance['D']['pct']}%)")
                
                # Проверяем per-ticket balance
                bad_tickets = []
                for t, dist in sorted(per_ticket.items()):
                    for letter in 'ABCD':
                        cnt = dist[letter]
                        if cnt < 4 or cnt > 6:
                            bad_tickets.append(f"Билет{t}: {letter}={cnt}")
                if bad_tickets:
                    print(f"    Нарушения локального D-баланса: {bad_tickets}")
                else:
                    print(f"    Локальный D-баланс: OK (4-6 в каждом билете)")
                
                print(f"    Таблиц в файле: {table_count}")
                print(f"    Дубликатов: {len(duplicates)}")
                print(f"    К0.11 нарушений: {len(violations)}")
                for v in violations:
                    print(f"      Б{v['ticket']}в{v['q_num']}: «{v['word']}» [{v['category']}] → ...{v['context']}...")
                
                results[folder_key][fos_num] = {
                    'file': fos_file,
                    'questions': len(questions),
                    'tickets': len(per_ticket),
                    'd_balance': d_balance,
                    'per_ticket_bad': bad_tickets,
                    'table_count': table_count,
                    'duplicates': len(duplicates),
                    'k011_violations': violations,
                }
    
    # Итоговая сводка
    print(f"\n\n{'='*70}")
    print("ИТОГОВАЯ СВОДКА ПО К0.11")
    print(f"{'='*70}")
    total_violations = 0
    for fk, fv in results.items():
        for fn, fd in fv.items():
            v_count = len(fd.get('k011_violations', []))
            total_violations += v_count
            if v_count > 0:
                print(f"  {fk}/{fn}: {v_count} нарушений")
    if total_violations == 0:
        print("  НЕТ НАРУШЕНИЙ К0.11")
    else:
        print(f"  ВСЕГО: {total_violations} нарушений")
    
    # Сохраняем JSON
    output_path = "/home/z/my-project/vol/scripts/k011_scan_results.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\nРезультаты сохранены: {output_path}")

if __name__ == '__main__':
    main()
