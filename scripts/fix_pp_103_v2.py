#!/usr/bin/env python3
"""
ПП 103: замена 2 спорных вопросов на новые (ключи сохранены)
- Q15 (C): микрометр/штангенциркуль двусмысленность → ГОСТ 19265 группа ВК
- Q18 (B): упрощённый припуск → модуль упругости (модуль Юнга)
D-баланс 5/5/5/5 без изменений.
"""

import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков'
FOLDER = 'профессиональной переподготовки 15474 Оператор станков'
FILEPATH = BASE + '/' + FOLDER + '/103. ПО_ПП_ФОС_Оператор станков_2-3_разр ОП.0.0.docx'

REPLACEMENTS = {
    15: {
        'text': 'Согласно ГОСТ 19265, какие твёрдые сплавы обозначаются буквами «ВК»?',
        'options': {
            'A': 'Титановольфрамокобальтовые сплавы (группа ТК);',
            'B': 'Безвольфрамовые сплавы (группа ТН);',
            'C': 'Вольфрамокобальтовые сплавы (группа ВК);',
            'D': 'Титанотанталовые сплавы (группа ТТ);',
        },
        'answer': 'C',
    },
    18: {
        'text': 'Что характеризует модуль упругости (модуль Юнга) материала?',
        'options': {
            'A': 'Твёрдость поверхности детали;',
            'B': 'Способность материала сопротивляться упругой деформации;',
            'C': 'Способность материала к пластической деформации;',
            'D': 'Ударную вязкость при динамических нагрузках;',
        },
        'answer': 'B',
    },
}

Q_RE = re.compile(r'^(\d+)\s*[.)]\s*(.*)')
OPT_RE = re.compile(r'^([ABCD])\)\s*(.*)')
ANS_RE = re.compile(r'Правильный ответ под номером:\s*([ABCD])')
HEADER_WORDS = ('Тестовые', 'Фонд', 'итоговой')


def parse_103(filepath):
    doc = Document(filepath)
    header_paras = []
    questions = []
    in_questions = False
    cur_q = None
    cur_text = None
    cur_opts = {}
    cur_ans = None

    for p in doc.paragraphs:
        txt_s = p.text.strip()
        if not txt_s:
            header_paras.append('')
            continue
        if not in_questions:
            if any(w in txt_s for w in HEADER_WORDS) or txt_s.startswith('Тестовые'):
                header_paras.append(p.text)
                continue
            m = Q_RE.match(txt_s)
            if m:
                in_questions = True
                cur_q = int(m.group(1))
                cur_text = m.group(2)
                cur_opts = {}
                cur_ans = None
                continue
            header_paras.append(p.text)
            continue
        # in_questions == True
        m = Q_RE.match(txt_s)
        if m:
            if cur_text is not None:
                questions.append({'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
            cur_q = int(m.group(1))
            cur_text = m.group(2)
            cur_opts = {}
            cur_ans = None
            continue
        m2 = OPT_RE.match(txt_s)
        if m2 and cur_text is not None:
            cur_opts[m2.group(1)] = m2.group(2)
            continue
        m3 = ANS_RE.search(txt_s)
        if m3 and cur_text is not None:
            cur_ans = m3.group(1)

    if cur_text is not None:
        questions.append({'num': cur_q, 'text': cur_text, 'options': cur_opts, 'answer': cur_ans})
    return header_paras, questions


def apply_replacements(questions):
    n = 0
    for q in questions:
        if q['num'] in REPLACEMENTS:
            new = REPLACEMENTS[q['num']]
            print('  ЗАМЕНА В' + str(q['num']) + ': ' + q['text'][:60] + '...')
            q['text'] = new['text']
            q['options'] = dict(new['options'])
            q['answer'] = new['answer']
            print('    -> ' + new['text'][:70] + '...')
            n += 1
    return n


def build_docx(header_paras, questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for txt in header_paras:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        s = txt.strip()
        if s and ('Фонд' in s or 'итоговой' in s or 'по основной' in s or
                  'по профессии' in s or s.startswith('«Оператор') or
                  'Тестовые' in s):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True

    for q in questions:
        qp = doc.add_paragraph(str(q['num']) + '. ' + q['text'])
        qp.paragraph_format.space_after = Pt(0)
        qp.paragraph_format.space_before = Pt(2)
        for letter in 'ABCD':
            if letter in q['options']:
                op = doc.add_paragraph(letter + ') ' + q['options'][letter])
                op.paragraph_format.space_after = Pt(0)
                op.paragraph_format.space_before = Pt(0)
                op.paragraph_format.left_indent = Cm(1)
        ap = doc.add_paragraph('Правильный ответ под номером: ' + q['answer'])
        ap.paragraph_format.space_after = Pt(4)
        ap.paragraph_format.space_before = Pt(0)
        for run in ap.runs:
            run.font.bold = True

    doc.save(output_path)


def main():
    print('=' * 60)
    print('ПП 103: замена 2 спорных вопросов')
    print('=' * 60)

    print('\n[1] Парсинг...')
    header_paras, questions = parse_103(FILEPATH)
    print('  Вопросов:', len(questions))
    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс до: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))

    print('\n[2] Замены...')
    n = apply_replacements(questions)
    print('  Заменено:', n)

    gc = Counter(q['answer'] for q in questions)
    print('  D-баланс после: A=' + str(gc['A']) + ' B=' + str(gc['B']) + ' C=' + str(gc['C']) + ' D=' + str(gc['D']))
    if gc['A'] == gc['B'] == gc['C'] == gc['D'] == 5:
        print('  D-баланс ОК')
    else:
        print('  D-БАЛАНС НАРУШЕН!')
        return

    print('\n[3] Генерация файла...')
    build_docx(header_paras, questions, FILEPATH)

    print('\n[4] Верификация из файла...')
    _, vq = parse_103(FILEPATH)
    gc2 = Counter(q['answer'] for q in vq)
    print('  Вопросов: ' + str(len(vq)) + ', D-баланс: A=' + str(gc2['A']) + ' B=' + str(gc2['B']) + ' C=' + str(gc2['C']) + ' D=' + str(gc2['D']))

    ok = True
    for q in vq:
        if q['num'] in REPLACEMENTS:
            exp = REPLACEMENTS[q['num']]
            if q['text'] == exp['text'] and q['answer'] == exp['answer']:
                print('  OK В' + str(q['num']) + ': ' + q['text'][:60] + '... -> ' + q['answer'])
            else:
                print('  FAIL В' + str(q['num']))
                ok = False

    if ok and gc2['A'] == gc2['B'] == gc2['C'] == gc2['D'] == 5:
        print('\n  ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ')
    else:
        print('\n  ЕСТЬ ПРОБЛЕМЫ')

    print('\n' + '=' * 60)


if __name__ == '__main__':
    main()
