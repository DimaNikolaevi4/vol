#!/usr/bin/env python3
"""Extract all questions from КПК 05 (old format) and deduplicate."""
import re, json
from collections import Counter
from docx import Document

FOLDER = '/home/z/my-project/vol/Оператор автоматических и полуавтоматических станков и линий станков/повышения квалификации 15474 Оператор станков'
PATH = f'{FOLDER}/05. ПО_КПК_ФОС_Оператор станков_4_разр.docx'

doc = Document(PATH)
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {sum(1 for el in doc.element.body if el.tag.endswith("tbl"))}')

# All questions are in single paragraphs with \n inside (same format as 03/04)
questions = []
seen_texts = set()
dupes = 0

for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'^(\d+)\.[\s]+(.+)', text, re.DOTALL)
    if not m:
        continue
    q_num = int(m.group(1))
    full_text = m.group(2)
    
    lines = full_text.split('\n')
    q_text = lines[0].strip()
    
    options = {}
    for line in lines[1:]:
        line = line.strip()
        opt_m = re.match(r'^([а-гА-Г])\)\s*(.+)', line)
        if opt_m:
            letter_cyr = opt_m.group(1).lower()
            letter_map = {'а': 'A', 'б': 'B', 'в': 'C', 'г': 'D'}
            options[letter_map[letter_cyr]] = opt_m.group(2).strip().rstrip(' ;.')
    
    if len(options) != 4:
        print(f'  Q{q_num}: {len(options)} options - SKIP')
        continue
    
    # Check for duplicate question text
    if q_text in seen_texts:
        dupes += 1
        continue
    seen_texts.add(q_text)
    
    questions.append({
        'num': q_num,
        'question': q_text,
        'A': options.get('A', ''),
        'B': options.get('B', ''),
        'C': options.get('C', ''),
        'D': options.get('D', ''),
        'ans': 'X'  # Will be assigned during D-balance
    })

print(f'\nTotal question paragraphs: {len([p for p in doc.paragraphs if re.match(r"^\d+\\.", p.text.strip())])}')
print(f'Parsed with 4 options: {len(questions) + dupes}')
print(f'Unique: {len(questions)}')
print(f'Duplicates removed: {dupes}')

# Save unique questions
OUT = '/home/z/my-project/vol/scripts/kpk_existing_unique.json'
with open(OUT, 'w', encoding='utf-8') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)
print(f'Saved {len(questions)} unique questions to {OUT}')

# Need: 300 - existing = how many new?
needed = 300 - len(questions)
print(f'\nNeed to generate: {needed} new questions to reach 300')
